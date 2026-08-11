#!/usr/bin/env python3
"""
Master Resume Generator
========================
Builds an exhaustive, "everything" master resume .docx from
src/data/john_profile.json — every professional-experience highlight, every
major achievement, all 7 AI/software projects in full detail, and all 16
LinkedIn recommendations verbatim.

This is deliberately NOT a send-ready 1-2 page resume. It is a single-source
knowledgebase document that tailored per-JD resumes (built by
scripts/jd_scorecard_resume_v2.py) can be checked against / mined from, so
that no piece of career history is missed when matching a new JD.

The .txt -> .docx conversion logic (style_docx_document / add_docx_text_block
/ add_runs_with_markup / convert_text_file_to_docx) below is copied — not
imported — from the proven converter in jd_scorecard_resume_v2.py, so this
produces the exact same visual style as every tailored resume and cover
letter. It is copied rather than imported because jd_scorecard_resume_v2.py
is a flat top-to-bottom script with no `if __name__ == "__main__"` guard —
importing it executes its entire live JD-processing pipeline (LLM calls,
file writes) as a side effect of the import statement. Copying the pure
helper functions avoids that entirely, and jd_scorecard_resume_v2.py itself
is not touched (soul.md golden-rule: proven scripts stay immutable).

Usage:
  python scripts/generate_master_resume.py

Outputs:
  data_processed/MasterResume/resume/txt/JohnHauResume2026_MASTER_FULL_<DATE>.txt
  data_processed/MasterResume/resume/docx/JohnHauResume2026_MASTER_FULL_<DATE>.docx
  JohnHauResume2026_MASTER_FULL_<DATE>.docx   (copied to project root)

Source: src/data/john_profile.json (single source of truth — no hallucination,
nothing here is written that isn't already in the profile JSON).
"""

import json
import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent.parent

PROFILE_PATH = ROOT / "src/data/john_profile.json"
DATE_STAMP = datetime.now().strftime("%d%b%Y").upper()  # e.g. 11AUG2026

# ── .txt -> .docx converter, copied from jd_scorecard_resume_v2.py ─────────

BOLD_MARKUP_RE = re.compile(r"\*\*(.+?)\*\*")


def add_runs_with_markup(paragraph, text):
    """Add text to a paragraph, rendering **marked** spans as bold runs."""
    if text.count("**") % 2 != 0:
        paragraph.add_run(text.replace("**", ""))
        return
    pos = 0
    for match in BOLD_MARKUP_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        bold_run = paragraph.add_run(match.group(1))
        bold_run.bold = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def style_docx_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)


def add_docx_text_block(doc, line):
    stripped = line.strip()
    if not stripped:
        doc.add_paragraph()
        return

    if set(stripped) <= {"=", "_", "-"} and len(stripped) >= 8:
        return

    if stripped.startswith(("•", "* ", "- ")):
        bullet_text = re.sub(r"^[•*-]\s+", "", stripped, count=1).strip()
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.left_indent = Inches(0.25)
        add_runs_with_markup(para, bullet_text)
        return

    if stripped.isupper() and len(stripped) <= 80:
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = para.add_run(stripped)
        run.bold = True
        run.font.size = Pt(12)
        return

    if " — " in stripped and not stripped.endswith(":"):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(3)
        left, right = stripped.split(" — ", 1)
        run1 = para.add_run(left + " — ")
        run1.bold = True
        add_runs_with_markup(para, right)
        return

    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    add_runs_with_markup(para, stripped)


def _is_effectively_blank(line):
    stripped = line.strip()
    if not stripped:
        return True
    return set(stripped) <= {"=", "_", "-"} and len(stripped) >= 8


_COMPANY_HEADER_RE = re.compile(r"^\S.*\s—\s\S")
_DATE_RANGE_RE = re.compile(
    r"^[A-Za-z]{3,9}\.?\s+\d{4}\s*[–—-]\s*(?:[A-Za-z]{3,9}\.?\s+\d{4}|Present)$"
)


def convert_text_file_to_docx(txt_path, out_path):
    doc = Document()
    style_docx_document(doc)

    collapsed_lines = []
    blank_pending = False
    for line in txt_path.read_text(encoding="utf-8").splitlines():
        if _is_effectively_blank(line):
            blank_pending = True
            continue
        if blank_pending and collapsed_lines:
            collapsed_lines.append("")
        blank_pending = False
        collapsed_lines.append(line)

    normalized_lines = []
    i = 0
    while i < len(collapsed_lines):
        line = collapsed_lines[i]
        if (
            line == ""
            and 0 < i < len(collapsed_lines) - 1
            and _COMPANY_HEADER_RE.match(collapsed_lines[i - 1].strip())
            and _DATE_RANGE_RE.match(collapsed_lines[i + 1].strip())
        ):
            i += 1
            continue
        normalized_lines.append(line)
        i += 1

    for line in normalized_lines:
        add_docx_text_block(doc, line)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path

OUT_BASE = ROOT / "data_processed/MasterResume/resume"
OUT_TXT_DIR = OUT_BASE / "txt"
OUT_DOCX_DIR = OUT_BASE / "docx"
OUT_NAME = f"JohnHauResume2026_MASTER_FULL_{DATE_STAMP}"

FIELD_LABELS = {
    "scale": "Scale",
    "impact": "Impact",
    "metric": "Metric",
    "cost_savings": "Cost Savings",
    "improvement": "Improvement",
    "tech": "Tech",
    "duration": "Duration",
    "count": "Count",
    "adoption": "Adoption",
    "security": "Security",
    "sla_compliance": "SLA Compliance",
    "coverage": "Coverage",
    "savings": "Savings",
    "fte_saved": "FTE Saved",
    "hours_saved": "Hours Saved",
    "method": "Method",
}


def sep(lines):
    lines.append("")


def header(lines, text):
    lines.append(text.upper())
    lines.append("=" * max(8, min(len(text), 78)))
    lines.append("")


def bullet(lines, text):
    lines.append(f"- {text}")


def build_contact_block(lines, meta):
    lines.append(meta["name"])
    lines.append(meta["title"])
    lines.append(
        f"{meta['location']} | {meta['phone']} | {meta['email']} | {meta['linkedin']}"
    )
    lines.append(f"Availability: {meta.get('availability', '')}")
    lines.append(
        f"Years of experience: {meta.get('years_experience')} "
        f"(career span {meta.get('experience_timeline', {}).get('total_span_years')} years, "
        f"including a {meta.get('experience_timeline', {}).get('career_break_years')}-year career break)"
    )
    sep(lines)
    lines.append("-" * 40)
    sep(lines)


def build_summary(lines, profile):
    header(lines, "Professional Summary")
    lines.append(profile["summary"])
    sep(lines)


def build_core_competencies(lines, cc):
    header(lines, "Core Competencies")
    for category, items in cc.items():
        label = category.replace("_", " ").title()
        lines.append(f"{label} — {', '.join(items)}")
    sep(lines)


def build_technical_skills(lines, skills):
    header(lines, "Technical Skills")
    for line in skills:
        bullet(lines, line)
    sep(lines)


def build_ai_highlights(lines, hl):
    header(lines, "AI & Automation Highlights")
    lines.append(hl.get("positioning", ""))
    sep(lines)
    for b in hl.get("bullets", []):
        bullet(lines, b)
    sep(lines)


def build_ai_projects(lines, projects):
    header(lines, "AI / Software Projects — Full Detail")
    for p in projects:
        lines.append(f"**{p['title']}**")
        meta_bits = [p.get("category", ""), p.get("status", ""), p.get("period", "")]
        lines.append(" | ".join(b for b in meta_bits if b))
        if p.get("url"):
            lines.append(p["url"])
        if p.get("github_url"):
            lines.append(p["github_url"])
        sep(lines)
        lines.append(p.get("description", ""))
        sep(lines)
        if p.get("tech_stack"):
            lines.append(f"Tech Stack — {', '.join(p['tech_stack'])}")
        if p.get("key_features"):
            lines.append("Key Features:")
            for f in p["key_features"]:
                bullet(lines, f)
        if p.get("impact"):
            lines.append(f"Impact — {p['impact']}")
        if p.get("transferable_skills"):
            lines.append(f"Transferable Skills — {', '.join(p['transferable_skills'])}")
        if p.get("resume_bullet"):
            sep(lines)
            lines.append(f"Resume bullet form: {p['resume_bullet']}")
        sep(lines)
        lines.append("-" * 40)
        sep(lines)


def build_professional_experience(lines, jobs):
    header(lines, "Professional Experience — Full Detail")
    for job in jobs:
        lines.append(f"{job['company']} — {job['title']}")
        lines.append(job["period"])
        if job.get("scope"):
            lines.append(job["scope"])
        sep(lines)
        for h in job.get("highlights", []):
            bullet(lines, h)
        sep(lines)


def build_major_achievements(lines, achievements):
    header(lines, "Major Achievements — Full List By Company")
    by_company = {}
    for a in achievements:
        by_company.setdefault(a.get("company", "Unspecified"), []).append(a)

    for company, items in by_company.items():
        lines.append(company.upper())
        lines.append("-" * min(len(company), 60))
        for a in items:
            details = []
            for key, label in FIELD_LABELS.items():
                if a.get(key):
                    details.append(f"{label}: {a[key]}")
            detail_text = "; ".join(details)
            text = f"**{a['achievement']}**" + (f" — {detail_text}" if detail_text else "")
            bullet(lines, text)
        sep(lines)


def build_education(lines, items):
    header(lines, "Education & Certifications")
    for e in items:
        bits = [e.get("issuer", "")]
        if e.get("year"):
            bits.append(str(e["year"]))
        if e.get("type"):
            bits.append(e["type"])
        line = f"**{e['credential']}** — " + ", ".join(b for b in bits if b)
        if e.get("details"):
            line += f" — {e['details']}"
        bullet(lines, line)
    sep(lines)


def build_languages(lines, langs):
    header(lines, "Languages")
    for l in langs:
        bullet(lines, f"{l['language']} — {l['proficiency']}")
    sep(lines)


def build_soft_skills(lines, skills):
    header(lines, "Soft Skills")
    for s in skills:
        bullet(lines, s)
    sep(lines)


def build_key_topics(lines, topics):
    header(lines, "Key Topics Index (JD-Matching Keyword Reference)")
    lines.append(", ".join(topics))
    sep(lines)


def build_recommendations(lines, recs):
    header(lines, f"LinkedIn Recommendations ({len(recs)})")
    for r in recs:
        # No ** markup here: a line containing " — " is already bolded on its
        # left side by add_docx_text_block's em-dash rule, which bolds the
        # raw left-hand text as-is rather than running it through
        # add_runs_with_markup — wrapping the name in ** as well would leak
        # literal asterisks into the rendered docx.
        lines.append(
            f"{r['recommender_name']} — {r['recommender_title']}, "
            f"{r['recommender_company']} ({r.get('location', '')})"
        )
        rel_bits = [f"Relationship: {r.get('relationship', '')}", f"Date: {r.get('date', '')}"]
        if r.get("credentials"):
            rel_bits.append(f"Credentials: {r['credentials']}")
        lines.append(" | ".join(rel_bits))
        sep(lines)
        lines.append(r["recommendation"])
        sep(lines)
        lines.append("-" * 40)
        sep(lines)


def build_master_resume_text(data):
    profile = data["profile"]
    lines = []

    build_contact_block(lines, profile["metadata"])
    build_summary(lines, profile)
    build_core_competencies(lines, profile["core_competencies"])
    build_technical_skills(lines, profile["technical_skills"])
    build_ai_highlights(lines, profile["ai_automation_highlights"])
    build_ai_projects(lines, profile["ai_projects"])
    build_professional_experience(lines, profile["professional_experience"])
    build_major_achievements(lines, profile["major_achievements"])
    build_education(lines, profile["education_certifications"])
    build_languages(lines, profile["languages"])
    build_soft_skills(lines, profile["soft_skills"])
    build_key_topics(lines, profile["key_topics_for_qa"])
    build_recommendations(lines, profile["linkedin_recommendations"])

    return "\n".join(lines)


def main():
    if not PROFILE_PATH.exists():
        sys.exit(f"ERROR: profile not found: {PROFILE_PATH}")

    data = json.loads(PROFILE_PATH.read_text(encoding="utf-8"))
    text = build_master_resume_text(data)

    OUT_TXT_DIR.mkdir(parents=True, exist_ok=True)
    OUT_DOCX_DIR.mkdir(parents=True, exist_ok=True)

    txt_path = OUT_TXT_DIR / f"{OUT_NAME}.txt"
    txt_path.write_text(text, encoding="utf-8")
    print(f"Wrote {txt_path}")

    docx_path = OUT_DOCX_DIR / f"{OUT_NAME}.docx"
    convert_text_file_to_docx(txt_path, docx_path)
    print(f"Wrote {docx_path}")

    root_copy = ROOT / f"{OUT_NAME}.docx"
    shutil.copyfile(docx_path, root_copy)
    print(f"Copied to {root_copy}")


if __name__ == "__main__":
    main()
