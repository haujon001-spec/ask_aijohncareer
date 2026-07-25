#!/usr/bin/env python3
"""
All-in-One JD Application Generator — v2
=========================================
v2 of scripts/jd_scorecard_resume.py (kept alongside, not replacing it, per
soul.md's golden-rule: branch a new dated/versioned script rather than
editing a proven one in place).

Generates the same three documents from a Job Description + john_profile.json:
  1. JD Match Scorecard  — unchanged from v1
  2. Tailored Resume     — v2 formatting changes (see below)
  3. Cover Letter        — v2 formatting changes (see below)

v2 changes vs. v1 (requested 21 Jul 2026, based on a diff of the user's
manual edits to the McDonalds resume/cover-letter docx outputs):
  Resume:
    - Title no longer prefixed with "TAILORED RESUME —" (just "EMPLOYER | ROLE")
    - "Generated :" / "Profile :" boilerplate lines removed from the output
    - Achievement bullets are written in SMART form with the key metric or
      standout achievement phrase wrapped in **double asterisks**, rendered
      as bold text in the .docx
    - Never states an exact "27 years" / "27+ years" figure — reworded to
      "extensive years" / "extensive experience"
    - Blank-line spacing tightened (no double-blank runs, no blank line
      directly under a section header)
  Cover Letter:
    - Same Generated/Profile line removal, same bold-achievement markup,
      same "extensive years" wording, same blank-line tightening
    - Title keeps its "COVER LETTER —" prefix (unlike the resume) — the
      user's own manual edit did not remove it
  Scorecard: unchanged — out of scope for this pass.

Usage: identical to v1, just invoke this file instead.
  python scripts/jd_scorecard_resume_v2.py
  python scripts/jd_scorecard_resume_v2.py "path/to/AnotherJD.txt"
  python scripts/jd_scorecard_resume_v2.py --scorecard-only
  python scripts/jd_scorecard_resume_v2.py --resume-only
  python scripts/jd_scorecard_resume_v2.py --coverletter-only
  python scripts/jd_scorecard_resume_v2.py --refresh-blueprint
  python scripts/jd_scorecard_resume_v2.py --batch
  python scripts/jd_scorecard_resume_v2.py --batch --force
  python scripts/jd_scorecard_resume_v2.py --llm=sonnet|deepseek|gemini
  python scripts/jd_scorecard_resume_v2.py --ResumeAdjustment
  python scripts/jd_scorecard_resume_v2.py --llm=custom --model=anthropic/claude-opus-4-8 --provider=openrouter
  python scripts/jd_scorecard_resume_v2.py --llm=sonnet --api-key=sk-or-v1-...

--llm=custom / --model= / --provider= (added 23 Jul 2026):
  Run any OpenRouter or DeepSeek model id without editing LLM_CONFIGS —
  --model=<model id> plus --provider=openrouter|deepseek together select the
  model and endpoint/key-env, same as the portal's "Custom" LLM option.

--api-key=<key> (added 23 Jul 2026):
  Overrides the resolved API key for this run only (works with any --llm
  choice, preset or custom) — bring your own key without touching
  .env.local/.env.vps. Same override mechanism as the portal's "API Keys"
  settings panel, which stores a personal key server-side and injects it into
  just the spawned process's environment.

--ResumeAdjustment (added 22 Jul 2026):
  Pulls the "6a) Resume Adjustments" recommendations out of this JD's own Match
  Scorecard (auto-detects the latest scorecard already on disk for this
  employer/JD; generates one first if none exists yet) and feeds them to the
  Resume and Cover Letter prompts as extra tailoring guidance. It only shapes
  wording/emphasis/section framing in the existing output — it never adds a
  visible "Resume Adjustments" heading, and it can never introduce a fact,
  figure, or claim that isn't already in src/data/john_profile.json (same
  single-source-of-truth / no-hallucination rule as everything else in this
  script).

Outputs go to the same data_processed/<Employer>/ tree as v1 — filenames are
date-stamped, so re-running v2 against a JD already processed by v1 on a
different day will not collide or overwrite the v1 output.

Source: src/data/john_profile.json  (single source of truth — no hallucination)
"""

import json
import os
import re
import subprocess
import sys
import time
import requests
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Inches, Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ── Project root (one level up from scripts/) ─────────────────────────────────
ROOT = Path(__file__).resolve().parent.parent

# ── Fixed paths ────────────────────────────────────────────────────────────────
PROFILE_PATH       = ROOT / "src/data/john_profile.json"
DEFAULT_JD_DIR     = ROOT / "data_raw/jd/txt"
JD_BLUEPRINT_DIR   = ROOT / "src/data/jd"
TEMPLATE_PATH      = ROOT / "data_raw/resume/txt/JohnHauResume2026_MorganStanley.md"

DATE_STAMP    = datetime.now().strftime("%d%b%Y").upper()   # e.g. 31MAR2026
RESUME_YEAR   = datetime.now().strftime("%Y")               # e.g. 2026

# ── v2-only: bold-markup + text-normalization helpers ─────────────────────────
BOLD_MARKUP_RE = re.compile(r"\*\*(.+?)\*\*")

# Matches "27 years", "27+ years", "27.33 years experience" etc. — case-insensitive
YEARS_FIGURE_RE = re.compile(r"\b27(?:\.\d+)?\+?\s+years?\b", re.IGNORECASE)


def soften_experience_years(text):
    """Replace any exact '27(+) years' figure with softer 'extensive years' wording."""
    return YEARS_FIGURE_RE.sub("extensive years", text)


def tighten_contact_header(text):
    """Remove ALL blank lines from the resume's top contact block (name,
    address, LinkedIn, website) — up to but excluding the first section
    separator line — leaving the rest of the document's spacing untouched."""
    lines = text.splitlines()
    sep_idx = next(
        (i for i, l in enumerate(lines) if set(l.strip()) <= {"=", "_", "-"} and len(l.strip()) >= 8),
        None,
    )
    if sep_idx is None:
        return text
    head = [l for l in lines[:sep_idx] if l.strip() != ""]
    return "\n".join(head + lines[sep_idx:])


DATE_LINE_RE = re.compile(r"^\d{1,2}\s+\w+\s+\d{4}$")


def clean_coverletter_header(text, employer_display):
    """Safety net: strip any date / 'Hiring Manager' / company-name lines and
    blank lines from the letter header block, in case the LLM didn't fully
    follow the prompt instructions. Stops at the salutation line."""
    lines = text.splitlines()
    sal_idx = next(
        (i for i, l in enumerate(lines) if l.strip().lower().startswith("dear ")),
        None,
    )
    if sal_idx is None:
        return text

    head = []
    for l in lines[:sal_idx]:
        stripped = l.strip()
        if not stripped:
            continue
        if stripped.lower() == "hiring manager":
            continue
        if stripped == employer_display:
            continue
        if DATE_LINE_RE.match(stripped):
            continue
        head.append(l)

    return "\n".join(head + lines[sal_idx:])


def normalize_blank_lines(text):
    """Collapse runs of 2+ blank lines to exactly 1, and drop a blank line that
    immediately follows an ALL-CAPS section header (tighter spacing per the
    user's manual V2 edits)."""
    lines = text.splitlines()
    out = []
    prev_blank = False
    prev_was_header = False
    for line in lines:
        stripped = line.strip()
        is_blank = not stripped
        is_header = bool(stripped) and stripped.isupper() and len(stripped) <= 80

        if is_blank:
            if prev_blank or prev_was_header:
                continue  # collapse double-blanks and header-then-blank
            out.append(line)
            prev_blank = True
            prev_was_header = False
            continue

        out.append(line)
        prev_blank = False
        prev_was_header = is_header

    return "\n".join(out)


def add_runs_with_markup(paragraph, text):
    """Add text to a paragraph, rendering **marked** spans as bold runs."""
    if text.count("**") % 2 != 0:
        # Unbalanced markup from the LLM — fall back to plain text rather
        # than leaking a stray "**" into the visible output.
        paragraph.add_run(text.replace("**", ""))
        return
    pos = 0
    for match in BOLD_MARKUP_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        bold_run = paragraph.add_run(match.group(1))
        bold_run.bold = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def resolve_resume_template():
    """Use the Markdown master resume template as the single authoritative source."""
    if TEMPLATE_PATH.exists():
        return TEMPLATE_PATH
    sys.exit(f"ERROR: Resume template not found: {TEMPLATE_PATH}")

TEMPLATE_PATH = resolve_resume_template()

def resolve_default_jd():
    """Pick the JD file explicitly provided via env, otherwise the most recently updated JD in data_raw/jd/txt/."""
    env_jd = os.environ.get("JD_DEFAULT_PATH", "").strip()
    if env_jd:
        env_path = Path(env_jd)
        if not env_path.is_absolute():
            env_path = ROOT / env_path
        if env_path.exists():
            return env_path

    jd_files = sorted(DEFAULT_JD_DIR.glob("JD_*.txt"), key=lambda p: p.stat().st_mtime, reverse=True)
    if jd_files:
        return jd_files[0]

    fallback_files = sorted(DEFAULT_JD_DIR.glob("*.txt"), key=lambda p: p.stat().st_mtime, reverse=True)
    if fallback_files:
        return fallback_files[0]

    sys.exit(f"ERROR: No JD text files found under {DEFAULT_JD_DIR}")

def load_jd_blueprint(jd_stem):
    """Optionally load a structured JD blueprint from src/data/jd/<JD_STEM>.json."""
    blueprint_path = JD_BLUEPRINT_DIR / f"{jd_stem}.json"
    if not blueprint_path.exists():
        return None, None

    try:
        blueprint = json.loads(blueprint_path.read_text(encoding="utf-8"))
        return blueprint_path, blueprint
    except Exception as exc:
        print(f"⚠️  Warning: Could not parse JD blueprint {blueprint_path.name}: {exc}")
        return blueprint_path, None


def derive_jd_metadata(path):
    jd_stem_local = path.stem
    jd_parts_local = jd_stem_local.replace("JD_", "").split("_")
    employer_local = jd_parts_local[0] if jd_parts_local else "Employer"
    role_slug_local = "_".join(jd_parts_local[1:]) if len(jd_parts_local) > 1 else "Role"
    role_tag_local = f"_{role_slug_local}" if role_slug_local and role_slug_local != "Role" else ""
    return jd_stem_local, employer_local, role_slug_local, role_tag_local


def build_output_targets(path):
    jd_stem_local, employer_local, role_slug_local, role_tag_local = derive_jd_metadata(path)
    out_base = ROOT / "data_processed" / employer_local
    return {
        "blueprint": JD_BLUEPRINT_DIR / f"{jd_stem_local}.json",
        "scorecard_dir": out_base / "ScoreCard" / "txt",
        "resume_dir": out_base / "resume" / "txt",
        "cover_dir": out_base / "CoverLetter" / "txt",
        "scorecard_pattern": f"JD_SCORECARD_{employer_local}{role_tag_local}_*.txt",
        "resume_pattern": f"JohnHauResume*_{employer_local}{role_tag_local}_*.txt",
        "cover_pattern": f"JohnHauCoverLetter_{employer_local}{role_tag_local}_*.txt",
    }


def requested_outputs_exist(path, include_scorecard=True, include_resume=True, include_coverletter=True):
    targets = build_output_targets(path)
    checks = []
    if include_scorecard:
        checks.append(any(targets["scorecard_dir"].glob(targets["scorecard_pattern"])))
    if include_resume:
        checks.append(any(targets["resume_dir"].glob(targets["resume_pattern"])))
    if include_coverletter:
        checks.append(any(targets["cover_dir"].glob(targets["cover_pattern"])))
    return bool(checks) and all(checks)


# ── v2-only: --ResumeAdjustment support ────────────────────────────────────
# Extracts "6a) Resume Adjustments" out of a scorecard's own TAILORING
# RECOMMENDATIONS section, stopping at "b)"/"c)" or end of text.
RESUME_ADJUSTMENTS_RE = re.compile(
    r"\**\s*(?:a\)\s*)?Resume Adjustments\b[^\n]*\n(.*?)"
    r"(?=\n\s*\**\s*(?:[bc]\)\s*)?(?:Interview Preparation|Certifications)|\n\s*#{1,4}\s*\d|\Z)",
    re.IGNORECASE | re.DOTALL,
)


def find_latest_scorecard_txt(path):
    """Most recently modified existing scorecard .txt for this JD, if any."""
    targets = build_output_targets(path)
    candidates = sorted(
        targets["scorecard_dir"].glob(targets["scorecard_pattern"]),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )
    return candidates[0] if candidates else None


def extract_resume_adjustments(scorecard_text):
    """Pull the '6a) Resume Adjustments' block out of a JD Match Scorecard body."""
    match = RESUME_ADJUSTMENTS_RE.search(scorecard_text)
    return match.group(1).strip() if match else None

# ── Parse CLI arguments ────────────────────────────────────────────────────────
args = [a for a in sys.argv[1:] if not a.startswith("-")]
flags = [a for a in sys.argv[1:] if a.startswith("-")]

jd_path = Path(args[0]) if args else resolve_default_jd()
if not jd_path.is_absolute():
    jd_path = ROOT / jd_path
if not jd_path.exists():
    sys.exit(f"ERROR: JD file not found: {jd_path}")

run_scorecard     = "--resume-only" not in flags and "--coverletter-only" not in flags
run_resume        = "--scorecard-only" not in flags and "--coverletter-only" not in flags
run_coverletter   = "--scorecard-only" not in flags and "--resume-only" not in flags
refresh_blueprint = "--refresh-blueprint" in flags
batch_mode        = "--batch" in flags
force_run         = "--force" in flags
generate_docx     = "--no-docx" not in flags
resume_adjustment = any(f.lower() == "--resumeadjustment" for f in flags)

# ── LLM selection via --llm=<name> flag ───────────────────────────────────────
llm_flag = next((f for f in flags if f.startswith("--llm=")), "--llm=sonnet")
llm_choice = llm_flag.split("=", 1)[1].lower()   # sonnet | deepseek | gemini | custom

# ── Bring-your-own-key / dynamic model overrides (added 23 Jul 2026) ──────────
# --model= and --provider= together select an arbitrary OpenRouter/DeepSeek
# model id when --llm=custom, instead of one of the three fixed presets below.
# --api-key= overrides the resolved API key for ANY --llm choice (preset or
# custom) — lets this script be run with a personal key without touching
# .env.local, same override the portal's "API Keys" settings panel uses.
model_flag    = next((f for f in flags if f.startswith("--model=")), None)
provider_flag = next((f for f in flags if f.startswith("--provider=")), None)
api_key_flag  = next((f for f in flags if f.startswith("--api-key=")), None)
custom_model_id   = model_flag.split("=", 1)[1] if model_flag else None
custom_provider    = provider_flag.split("=", 1)[1].lower() if provider_flag else None
cli_api_key_override = api_key_flag.split("=", 1)[1] if api_key_flag else None

# ── Derive employer slug from JD filename for output naming ───────────────────
jd_stem    = jd_path.stem                               # e.g. JD_MandarinOriental_ClusterDirectorOfIT
jd_parts   = jd_stem.replace("JD_", "").split("_")
employer   = jd_parts[0] if jd_parts else "Employer"   # e.g. MandarinOriental
role_slug  = "_".join(jd_parts[1:]) if len(jd_parts) > 1 else "Role"
jd_blueprint_path, jd_blueprint = load_jd_blueprint(jd_stem)

# ── Output directories: data_processed/<Employer>/<Type>/txt/ ─────────────────
OUT_BASE          = ROOT / "data_processed" / employer
OUT_DIR_SCORECARD = OUT_BASE / "ScoreCard" / "txt"
OUT_DIR_RESUME    = OUT_BASE / "resume"    / "txt"
OUT_DIR_COVER     = OUT_BASE / "CoverLetter" / "txt"
_role_tag         = f"_{role_slug}" if role_slug and role_slug != "Role" else ""
OUT_SCORECARD     = OUT_DIR_SCORECARD / f"JD_SCORECARD_{employer}{_role_tag}_{DATE_STAMP}.txt"
OUT_RESUME        = OUT_DIR_RESUME    / f"JohnHauResume{RESUME_YEAR}_{employer}{_role_tag}_{DATE_STAMP}.txt"
OUT_COVERLETTER   = OUT_DIR_COVER     / f"JohnHauCoverLetter_{employer}{_role_tag}_{DATE_STAMP}.txt"

# ── Load env (.env.local → .env.vps → .env) — file always wins over stale shell ──
def load_env(path):
    if not path.exists():
        return False
    for line in path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if line and not line.startswith("#") and "=" in line:
            k, _, v = line.partition("=")
            os.environ[k.strip()] = v.strip()
    return True

for env_file in [ROOT / ".env.local", ROOT / ".env.vps", ROOT / ".env"]:
    if load_env(env_file):
        break

# ── Resolve API key and model based on LLM selection ─────────────────────────
LLM_CONFIGS = {
    # name          model_id                                     api_key_env            base_url
    "sonnet":  ("anthropic/claude-sonnet-5",                "OPENROUTER_API_KEY",  "https://openrouter.ai/api/v1/chat/completions"),
    "deepseek":("deepseek-reasoner",                        "DEEPSEEK_API_KEY",    "https://api.deepseek.com/chat/completions"),
    "gemini":  ("google/gemini-3.1-flash-lite-preview",     "OPENROUTER_API_KEY",  "https://openrouter.ai/api/v1/chat/completions"),
}

# provider name -> (api_key_env, base_url), used only when --llm=custom
PROVIDER_ENDPOINTS = {
    "openrouter": ("OPENROUTER_API_KEY", "https://openrouter.ai/api/v1/chat/completions"),
    "deepseek":   ("DEEPSEEK_API_KEY",   "https://api.deepseek.com/chat/completions"),
}

if llm_choice == "custom":
    if not custom_model_id or not custom_provider:
        sys.exit("ERROR: --llm=custom requires both --model=<model id> and --provider=openrouter|deepseek")
    if custom_provider not in PROVIDER_ENDPOINTS:
        sys.exit(f"ERROR: Unknown --provider value '{custom_provider}'. Valid: openrouter, deepseek")
    MODEL = custom_model_id
    api_key_env, LLM_ENDPOINT = PROVIDER_ENDPOINTS[custom_provider]
elif llm_choice in LLM_CONFIGS:
    MODEL, api_key_env, LLM_ENDPOINT = LLM_CONFIGS[llm_choice]
else:
    sys.exit(f"ERROR: Unknown --llm value '{llm_choice}'. Valid: sonnet, deepseek, gemini, custom")

API_KEY = cli_api_key_override or os.environ.get(api_key_env, "")
if not API_KEY:
    sys.exit(f"ERROR: {api_key_env} not found in .env.local / .env.vps / .env (or pass --api-key=<key>)")

HEADERS = {
    "Authorization": f"Bearer {API_KEY}",
    "Content-Type": "application/json",
    "HTTP-Referer": "https://www.askcareer-ai.com",
    "X-Title": "JD Application Generator - John Hau",
}

# Retried on: connection/timeout errors, and status codes observed in practice
# to be transient on OpenRouter's multi-provider routing (a request can 403
# against one upstream route and succeed seconds later against another for
# the identical key/model — confirmed 25 Jul 2026, not a real auth failure).
LLM_MAX_ATTEMPTS = 3
LLM_RETRY_BACKOFF_SEC = 3
LLM_RETRYABLE_STATUS_CODES = {403, 408, 425, 429, 500, 502, 503, 504}

if batch_mode:
    jd_files = sorted(DEFAULT_JD_DIR.glob("JD_*.txt"))
    if not jd_files:
        sys.exit(f"ERROR: No JD text files found under {DEFAULT_JD_DIR}")

    print(f"\n{'='*60}")
    print("  JD APPLICATION GENERATOR v2 — BATCH MODE")
    print(f"{'='*60}")
    print(f"  JD Folder   : {DEFAULT_JD_DIR.relative_to(ROOT)}")
    print(f"  File Count  : {len(jd_files)}")
    print(f"  LLM         : {llm_choice} ({MODEL})")
    print(f"  Skip Existing: {'Yes' if not force_run else 'No (force enabled)'}")
    print(f"{'='*60}\n")

    processed = 0
    skipped = 0
    failed = 0
    batch_results = []
    forwarded_flags = [f for f in flags if f not in {"--batch", "--force"}]

    for batch_jd in jd_files:
        print(f"\n▶ Processing {batch_jd.name}")
        targets = build_output_targets(batch_jd)
        blueprint_exists = targets["blueprint"].exists()

        if not force_run and requested_outputs_exist(batch_jd, run_scorecard, run_resume, run_coverletter):
            print("  ↷ Skipped — requested outputs already exist")
            skipped += 1
            batch_results.append({
                "jd": batch_jd.name,
                "status": "SKIPPED",
                "blueprint": "EXISTS" if blueprint_exists else "MISSING",
                "note": "Outputs already existed",
            })
            continue

        cmd = [sys.executable, str(Path(__file__).resolve()), str(batch_jd)] + forwarded_flags
        result = subprocess.run(cmd, cwd=str(ROOT))
        if result.returncode == 0:
            processed += 1
            batch_results.append({
                "jd": batch_jd.name,
                "status": "DONE",
                "blueprint": "READY",
                "note": f"Output -> data_processed/{derive_jd_metadata(batch_jd)[1]}/",
            })
        else:
            failed += 1
            print(f"  ✗ Failed — exit code {result.returncode}")
            batch_results.append({
                "jd": batch_jd.name,
                "status": "FAILED",
                "blueprint": "CHECK",
                "note": f"Exit code {result.returncode}",
            })

    print(f"\n{'='*60}")
    print("  BATCH SUMMARY")
    print(f"{'='*60}")
    print(f"  Processed : {processed}")
    print(f"  Skipped   : {skipped}")
    print(f"  Failed    : {failed}")

    print(f"\n{'-'*112}")
    print(f"  {'JD FILE':48} {'STATUS':10} {'BLUEPRINT':10} NOTE")
    print(f"{'-'*112}")
    for item in batch_results:
        jd_name = (item['jd'][:45] + '...') if len(item['jd']) > 48 else item['jd']
        print(f"  {jd_name:48} {item['status']:10} {item['blueprint']:10} {item['note']}")
    print(f"{'-'*112}")

    sys.exit(0 if failed == 0 else 1)

# ── v2-only: resolve --ResumeAdjustment source before anything prints ─────────
# Auto-detects the latest existing scorecard for this JD; if none exists yet,
# forces the scorecard to be generated this run so the guidance is available.
resume_adjustments_text = None
if resume_adjustment and (run_resume or run_coverletter):
    existing_scorecard_path = find_latest_scorecard_txt(jd_path)
    if existing_scorecard_path:
        resume_adjustments_text = extract_resume_adjustments(
            existing_scorecard_path.read_text(encoding="utf-8")
        )
        _scorecard_source_note = f"existing → {existing_scorecard_path.relative_to(ROOT)}"
    elif not run_scorecard:
        run_scorecard = True
        _scorecard_source_note = "none found — generating one first"
    else:
        _scorecard_source_note = "will extract from this run's freshly generated scorecard"
elif resume_adjustment:
    _scorecard_source_note = "ignored — no resume/cover-letter requested this run"

# ── Read source files ──────────────────────────────────────────────────────────
print(f"\n{'='*60}")
print("  JD APPLICATION GENERATOR v2 — John Hau")
print(f"{'='*60}")
print(f"  JD File     : {jd_path.name}")
print(f"  JD Blueprint: {jd_blueprint_path.relative_to(ROOT) if jd_blueprint_path and jd_blueprint else 'dynamic from JD text'}")
print(f"  Template    : {TEMPLATE_PATH.relative_to(ROOT)}")
print(f"  Profile     : src/data/john_profile.json")
print(f"  Model       : {MODEL}")
print(f"  API Key     : {API_KEY[:25]}...")
print(f"  Date        : {DATE_STAMP}")
print(f"  LLM         : {llm_choice} ({MODEL})")
print(f"  Generating  : {'Scorecard ' if run_scorecard else ''}{'Resume ' if run_resume else ''}{'CoverLetter' if run_coverletter else ''}")
if resume_adjustment:
    print(f"  ResumeAdj   : {_scorecard_source_note}")
print(f"{'='*60}\n")

print("📂  Reading source files...")
jd_text      = jd_path.read_text(encoding="utf-8")
profile_raw  = json.loads(PROFILE_PATH.read_text(encoding="utf-8"))
template_txt = TEMPLATE_PATH.read_text(encoding="utf-8")

# Unwrap nested "profile" key if present
profile = profile_raw.get("profile", profile_raw)

# Build a compact profile context — metadata + summary + experience + achievements
# Stays within ~12K chars to avoid token overrun while keeping all key facts
def build_profile_context(profile):
    sections = {}
    for key in ["metadata", "summary", "professional_experience", "major_achievements",
                "ai_projects", "core_competencies", "technical_skills",
                "education_certifications", "languages_spoken"]:
        if key in profile:
            sections[key] = profile[key]
    return json.dumps(sections, indent=2, ensure_ascii=False)[:14000]

profile_context = build_profile_context(profile)

def build_blueprint_context(blueprint):
    return (
        json.dumps(blueprint, indent=2, ensure_ascii=False)[:6000]
        if blueprint
        else "No JD JSON blueprint found. Infer the priority criteria, must-have skills, soft skills, and weighting dynamically from the raw JD text."
    )

jd_blueprint_context = build_blueprint_context(jd_blueprint)

# ── LLM call helper ────────────────────────────────────────────────────────────
def call_llm(system_prompt, user_prompt, max_tokens=6000, label="", response_format=None):
    if label:
        print(f"  ↳ Calling OpenRouter ({MODEL}) — {label}")
    payload = {
        "model": MODEL,
        "max_tokens": max_tokens,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_prompt},
        ],
    }
    if response_format:
        payload["response_format"] = response_format

    resp = None
    for attempt in range(1, LLM_MAX_ATTEMPTS + 1):
        try:
            resp = requests.post(
                LLM_ENDPOINT,
                headers=HEADERS,
                json=payload,
                timeout=180,
            )
        except requests.exceptions.RequestException as exc:
            # Connection/timeout failures — always worth a retry until the last attempt.
            if attempt < LLM_MAX_ATTEMPTS:
                wait = LLM_RETRY_BACKOFF_SEC * attempt
                print(
                    f"  ⚠️  {label or 'unlabeled'} call failed ({exc}) "
                    f"(attempt {attempt}/{LLM_MAX_ATTEMPTS}) — retrying in {wait}s..."
                )
                time.sleep(wait)
                continue
            raise
        if resp.status_code in LLM_RETRYABLE_STATUS_CODES and attempt < LLM_MAX_ATTEMPTS:
            wait = LLM_RETRY_BACKOFF_SEC * attempt
            print(
                f"  ⚠️  {resp.status_code} from {label or 'unlabeled'} call "
                f"(attempt {attempt}/{LLM_MAX_ATTEMPTS}) — retrying in {wait}s..."
            )
            time.sleep(wait)
            continue
        resp.raise_for_status()  # raises immediately for non-retryable or attempts-exhausted errors
        break
    data = resp.json()
    choice = data["choices"][0]
    message = choice["message"]
    content = message.get("content")
    finish_reason = choice.get("finish_reason")
    if not content:
        refusal = message.get("refusal")
        raise RuntimeError(
            f"OpenRouter returned empty content for '{label or 'unlabeled'}' call "
            f"(model={MODEL}, finish_reason={finish_reason!r}, refusal={refusal!r}). "
            "This usually means the model ran out of max_tokens before producing an "
            "answer, or refused the request — check the values above."
        )
    if finish_reason == "length":
        usage = data.get("usage", {})
        raise RuntimeError(
            f"OpenRouter truncated the '{label or 'unlabeled'}' response — it ran out of "
            f"max_tokens ({max_tokens}) before finishing (model={MODEL}, "
            f"completion_tokens={usage.get('completion_tokens')}, "
            f"reasoning_tokens={usage.get('completion_tokens_details', {}).get('reasoning_tokens')}). "
            "The partial output was discarded rather than silently shipped incomplete; raise "
            "max_tokens for this call and retry."
        )
    return content

def extract_json_object(raw_text):
    """Extract a JSON object from an LLM response, allowing for code fences and minor cleanup."""
    cleaned = raw_text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
        cleaned = re.sub(r"\s*```$", "", cleaned)

    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("LLM response did not contain a valid JSON object")

    candidate = cleaned[start:end + 1]
    candidate = (
        candidate
        .replace("“", '"')
        .replace("”", '"')
        .replace("‘", "'")
        .replace("’", "'")
    )

    attempts = [
        candidate,
        re.sub(r",\s*([}\]])", r"\1", candidate),
    ]

    last_error = None
    for attempt in attempts:
        try:
            return json.loads(attempt)
        except json.JSONDecodeError as exc:
            last_error = exc

    raise last_error


def repair_json_with_llm(raw_text, label="JSON Repair"):
    """Ask the model to repair malformed JSON and return valid JSON only."""
    REPAIR_SYS = (
        "You repair malformed JSON for production automation. "
        "Return VALID JSON ONLY. Do not add commentary, markdown, or explanations."
    )
    REPAIR_USER = f"""
Repair the following malformed JSON so it parses successfully.
Preserve the original meaning and keys as closely as possible.
Return only valid JSON.

=== MALFORMED JSON ===
{raw_text}
"""
    repaired = call_llm(
        REPAIR_SYS,
        REPAIR_USER,
        max_tokens=6000,
        label=label,
        response_format={"type": "json_object"},
    )
    return extract_json_object(repaired)


def style_docx_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def add_docx_text_block(doc, line):
    stripped = line.strip()
    if not stripped:
        doc.add_paragraph()
        return

    if set(stripped) <= {"=", "_", "-"} and len(stripped) >= 8:
        return

    if stripped.startswith(("•", "* ", "- ")):
        bullet_text = re.sub(r"^[•*-]\s+", "", stripped, count=1).strip()
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.left_indent = Inches(0.25)
        add_runs_with_markup(para, bullet_text)
        return

    if stripped.isupper() and len(stripped) <= 80:
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = para.add_run(stripped)
        run.bold = True
        run.font.size = Pt(12)
        return

    if " — " in stripped and not stripped.endswith(":"):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(3)
        left, right = stripped.split(" — ", 1)
        run1 = para.add_run(left + " — ")
        run1.bold = True
        add_runs_with_markup(para, right)
        return

    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    add_runs_with_markup(para, stripped)


def convert_text_file_to_docx(txt_path):
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed")

    doc = Document()
    style_docx_document(doc)
    for line in txt_path.read_text(encoding="utf-8").splitlines():
        add_docx_text_block(doc, line)

    target_dir = txt_path.parent.parent / "docx"
    target_dir.mkdir(parents=True, exist_ok=True)
    out_path = target_dir / f"{txt_path.stem}.docx"
    try:
        doc.save(out_path)
        return out_path, False
    except PermissionError:
        fallback_name = f"{txt_path.stem}_{datetime.now().strftime('%H%M%S')}.docx"
        fallback_path = target_dir / fallback_name
        doc.save(fallback_path)
        return fallback_path, True


def maybe_convert_to_docx(txt_path, label):
    if not generate_docx:
        return None
    if not DOCX_AVAILABLE:
        print(f"  ⚠️  {label} DOCX skipped — python-docx not installed")
        return None
    try:
        out_path, used_fallback = convert_text_file_to_docx(txt_path)
        if used_fallback:
            print(f"  ⚠️  {label} DOCX original file was locked; wrote fallback → {out_path.relative_to(ROOT)}")
        else:
            print(f"  ✅  {label} DOCX → {out_path.relative_to(ROOT)}")
        return out_path
    except Exception as exc:
        print(f"  ⚠️  {label} DOCX conversion failed: {exc}")
        return None


def generate_jd_blueprint(jd_text, jd_path, employer, role_slug):
    """Auto-create a structured JD blueprint for any new JD text file."""
    blueprint_path = JD_BLUEPRINT_DIR / f"{jd_stem}.json"
    JD_BLUEPRINT_DIR.mkdir(parents=True, exist_ok=True)

    BLUEPRINT_SYS = (
        "You are an expert job architecture analyst and recruiter. "
        "Read the JD and return VALID JSON ONLY for downstream automation. "
        "Do not use markdown code fences. Do not include commentary. "
        "Make the criteria practical, recruiter-style, and specific to this JD."
    )

    BLUEPRINT_USER = f"""
=== JOB DESCRIPTION ===
{jd_text}

=== INSTRUCTIONS ===
Create a JSON blueprint for this JD using the exact top-level keys below:
{{
  "jd_file": "...",
  "company": "...",
  "company_slug": "...",
  "job_title": "...",
  "role_slug": "...",
  "industry": "...",
  "role_summary": "...",
  "must_have_requirements": ["..."],
  "nice_to_have_requirements": ["..."],
  "soft_skills_priority": ["..."],
  "semantic_match_guidance": {{
    "category_name": ["related resume/profile evidence terms"]
  }},
  "scoring_criteria": [
    {{
      "name": "...",
      "weight": 0,
      "why_it_matters": "...",
      "evidence_type_to_look_for": ["..."],
      "gap_expected": "optional"
    }}
  ],
  "resume_tailoring_guidance": {{
    "mandatory_sections": ["PROFESSIONAL SUMMARY", "AI & AUTOMATION HIGHLIGHTS", "CORE COMPETENCIES", "SOFT SKILLS"],
    "summary_focus": ["..."],
    "do_not_overclaim": ["..."]
  }},
  "cover_letter_guidance": {{
    "tone": "...",
    "positioning": "...",
    "bridge_gaps_honestly": ["..."]
  }}
}}

Rules:
- Base the blueprint on the JD text only
- Include 8-10 scoring criteria
- Make all scoring weights numeric and total exactly 100
- Keep wording concise and usable for resume / scorecard / cover letter generation
- Include semantic matching hints so adjacent evidence can be matched truthfully without hallucination
"""

    raw = call_llm(
        BLUEPRINT_SYS,
        BLUEPRINT_USER,
        max_tokens=6000,
        label="JD Blueprint",
        response_format={"type": "json_object"},
    )
    try:
        blueprint = extract_json_object(raw)
    except Exception as exc:
        print(f"  ↳ Repairing malformed JD blueprint JSON ({exc})")
        blueprint = repair_json_with_llm(raw, label="JD Blueprint Repair")

    blueprint.setdefault("jd_file", str(jd_path.relative_to(ROOT)).replace("\\", "/"))
    blueprint.setdefault("company", employer)
    blueprint.setdefault("company_slug", employer)
    blueprint.setdefault("job_title", role_slug.replace("_", " "))
    blueprint.setdefault("role_slug", role_slug)
    blueprint.setdefault("resume_tailoring_guidance", {})
    blueprint["resume_tailoring_guidance"].setdefault(
        "mandatory_sections",
        ["PROFESSIONAL SUMMARY", "AI & AUTOMATION HIGHLIGHTS", "CORE COMPETENCIES", "SOFT SKILLS"],
    )

    blueprint_path.write_text(json.dumps(blueprint, indent=2, ensure_ascii=False), encoding="utf-8")
    return blueprint_path, blueprint

if refresh_blueprint or not jd_blueprint:
    action = "Refreshing" if refresh_blueprint and jd_blueprint_path else "Generating"
    print(f"🧭  [0/3] {action} JD blueprint...")
    try:
        jd_blueprint_path, jd_blueprint = generate_jd_blueprint(jd_text, jd_path, employer, role_slug)
        jd_blueprint_context = build_blueprint_context(jd_blueprint)
        print(f"  ✓  JD Blueprint ready -> {jd_blueprint_path.relative_to(ROOT)}")
    except Exception as exc:
        print(f"  ⚠️  JD blueprint generation skipped: {exc}")
        jd_blueprint_context = build_blueprint_context(jd_blueprint)

# ══════════════════════════════════════════════════════════════════════════════
# TASK 1 — JD SCORECARD (unchanged from v1)
# ══════════════════════════════════════════════════════════════════════════════
scorecard_text = ""
if run_scorecard:
    print("🔍  [1/3] Generating JD Scorecard...")

    SCORECARD_SYS = (
        "You are an expert executive recruiter and talent analyst. "
        "Produce a detailed, honest, structured scorecard comparing the candidate to the JD. "
        "Be evidence-based and professional. Never invent facts not present in the profile data. "
        "Use plain text with clear section headers. Derive the scoring criteria dynamically from the specific JD, "
        "focus only on materially relevant requirements, and compute a weighted overall score. "
        "When evidence is adjacent or transferable rather than exact, say so explicitly and do not present it as direct hands-on experience."
    )

    SCORECARD_USER = f"""
=== JOB DESCRIPTION ===
{jd_text}

=== OPTIONAL JD BLUEPRINT (use this if present, otherwise infer from JD text) ===
{jd_blueprint_context}

=== CANDIDATE PROFILE (source of truth — do not invent facts outside this data) ===
{profile_context}

=== INSTRUCTIONS ===
Produce a comprehensive JD Match Scorecard with ALL of these sections:

1. ROLE & COMPANY SUMMARY
   Brief overview of the role and what the employer is seeking (3-4 sentences).

2. OVERALL MATCH SCORE
   Format:  XX/100 — [STRONG MATCH / GOOD MATCH / PARTIAL MATCH / WEAK MATCH]
   Followed by one honest paragraph verdict.
   IMPORTANT: compute this as a weighted result from the JD-specific criteria below.
   Do NOT penalise the candidate for criteria that are not actually important in this JD.

3. DYNAMIC SCORING CRITERIA
   - First identify the 8-12 MOST IMPORTANT criteria directly from THIS JD.
   - These criteria MUST change depending on the role. Do not reuse a fixed checklist.
   - Use role-relevant categories only, for example when applicable:
     enterprise architecture, solution/integration design, stakeholder engagement,
     AI innovation, hospitality systems, cyber/compliance, vendor leadership,
     business partnership, programme delivery, people leadership, communication.
   - If the optional JD JSON blueprint exists, use its priorities/criteria/weights as the starting point.
   - If no JSON exists, infer the criteria directly from the raw JD.
   - Use semantic matching, not just exact keyword matching. Example: if the JD asks for cybersecurity,
     and the profile shows IT security, risk, audit, DR, patching, compliance, or resilience work,
     connect those dots clearly as relevant evidence.
   - Reward strong transferable evidence when it is meaningfully adjacent to the JD, but do NOT overstate or invent direct experience.
   - For each criterion provide:
     * Criterion Name
     * Weight (%) based on JD importance — all weights must total 100
     * Score: X/10
     * Evidence: specific proof from john_profile.json
     * Gap: what is missing or weaker versus the JD

4. KEY STRENGTHS
   Bullet list — what clearly matches or exceeds the JD requirements.

5. GAPS & RISKS
   Honest, specific gaps between the candidate profile and JD expectations.

6. TAILORING RECOMMENDATIONS
   a) Resume adjustments — specific wording and framing suggestions
   b) Interview preparation — topics to research, answers to rehearse
   c) Certifications to pursue — quick wins before applying

7. VERDICT
   Should the candidate apply? Competitive positioning. Recommended approach.
"""

    scorecard_text = call_llm(SCORECARD_SYS, SCORECARD_USER, max_tokens=12000, label="Scorecard")
    print("  ✓  Scorecard complete")

    if resume_adjustment and resume_adjustments_text is None:
        resume_adjustments_text = extract_resume_adjustments(scorecard_text)
        if not resume_adjustments_text:
            print("  ⚠️  --ResumeAdjustment: could not find a '6a) Resume Adjustments' section in this scorecard")

# ══════════════════════════════════════════════════════════════════════════════
# TASK 2 — TAILORED RESUME (v2: SMART + bold-markup + no "27 years")
# ══════════════════════════════════════════════════════════════════════════════
resume_text = ""
if run_resume:
    print("📝  [2/3] Generating Tailored Resume...")

    RESUME_SYS = (
        "You are a professional executive resume writer specialising in senior IT leadership roles. "
        "Rules you MUST follow:\n"
        "1. ALL facts, figures, dates, and claims come ONLY from the candidate profile — never invent or embellish.\n"
        "2. Mirror the reference template layout EXACTLY: same section headers, same ___ separators, same bullet style.\n"
        "3. Tailor bullet wording to echo the JD language without distorting facts.\n"
        "4. Use semantic alignment for related skills, but NEVER claim direct hands-on experience with named systems, standards, tools, or industries unless they appear explicitly in the profile data.\n"
        "5. If the JD asks for hospitality/PMS/POS/PCI or other domain-specific items not explicitly in the profile, position John's background as transferable or adjacent experience only.\n"
        "6. The `AI & AUTOMATION HIGHLIGHTS` section is mandatory in every resume output. Preserve it and populate it with vivid, concrete examples from the profile/template only.\n"
        "7. Make the resume balanced: show both technical depth and leadership/soft skills supported by the profile.\n"
        "8. Every achievement bullet must be written in SMART form (Specific, Measurable, Achievable, Relevant, Time-bound). Wrap EVERY distinct quantifiable figure in the bullet in double asterisks — percentages, dollar/HK$ amounts, headcounts, device/user counts, time savings, etc. — no cap on how many per bullet; bold every one that appears, but never wrap overlapping or adjacent text as a single run. The reference template shows this convention in a few of its bullets; follow that pattern for every bullet you write.\n"
        "9. Never state an exact number of years of experience (e.g. do not write '27 years' or '27+ years'). Use wording like 'extensive years' or 'extensive experience' instead.\n"
        "10. Weave visible people-management evidence into the most recent 2-3 roles specifically (not just a generic soft-skills line) — team leadership, mentoring/coaching, hiring or onboarding, performance management, career development, org design — using only what the profile actually documents (e.g. leading a team of 50+, coaching teams across multiple countries, mentorship and team development).\n"
        "11. Output ONLY the resume text — no preamble, no explanation, no markdown code fences (the ** bold markers from rule 8 are the one exception — those are expected).\n"
        + ("12. Recruiter resume-adjustment guidance (below) may be supplied — it comes from this JD's own Match Scorecard. Apply it ONLY to wording, emphasis, section framing, and which existing facts get foregrounded. It must NEVER be used to introduce a fact, figure, project, or claim that is not already present in the candidate profile data — the guidance changes how real facts are presented, never what the facts are. Never print the guidance verbatim or add a visible \"Resume Adjustments\" heading.\n" if resume_adjustment else "")
    )

    resume_adjustment_block = (
        f"""
=== RECRUITER RESUME-ADJUSTMENT GUIDANCE (from this JD's own Match Scorecard — apply per system rule 12; never invent a fact to satisfy it) ===
{resume_adjustments_text}
"""
        if resume_adjustments_text else ""
    )

    RESUME_USER = f"""
=== REFERENCE RESUME TEMPLATE (replicate this layout exactly) ===
{template_txt}

=== JOB DESCRIPTION (tailor content for this role) ===
{jd_text}

=== OPTIONAL JD BLUEPRINT (use this if present, otherwise infer from JD text) ===
{jd_blueprint_context}

=== CANDIDATE PROFILE (single source of truth — all facts from here only) ===
{profile_context}
{resume_adjustment_block}
=== INSTRUCTIONS ===
Write a complete tailored resume for John Hau targeting the role above.

Layout rules:
- Keep the identical header format, ________ separators, and bullet style from the reference template
- Respect and preserve any existing soft-skills section already present in the reference template, while tailoring it to this JD using only profile facts
- The `AI & AUTOMATION HIGHLIGHTS` section from the template is mandatory and must never be omitted, because it demonstrates current innovation capability with vivid examples
- First infer the 5-8 most important themes from THIS JD (technical, leadership, business, industry, and interpersonal)
- Replace the Professional Summary to foreground those JD-specific themes using only evidence from the profile
- Add a dynamic bridging section immediately before PROFESSIONAL EXPERIENCE. The section title and content must fit the JD nature
  (for example: "HOSPITALITY RELEVANCE", "ROLE RELEVANCE", or "ENTERPRISE ARCHITECTURE RELEVANCE") and should never be hardcoded
- BULLET COUNT RULES (non-negotiable):
    * 1st company (most recent):  exactly 12 bullets
    * 2nd company:                exactly 12 bullets
    * 3rd company:                exactly 10 bullets
    * 4th company:                exactly 10 bullets
    * 5th company:                exactly 8 bullets
    * Any earlier roles:          3-4 bullets combined
- Choose the bullets most relevant to the JD — reword to echo JD language without distorting facts
- Use semantic matching to connect related evidence from `john_profile.json` to the JD. Example: cybersecurity ↔ IT security / audit / risk / DR / compliance;
  stakeholder management ↔ executive presentations / business partnering; operational excellence ↔ SLA, uptime, automation, service quality improvements
- Do NOT convert adjacent evidence into unsupported exact claims. Example: do not say John has direct `PCI-DSS`, `Opera PMS`, or hotel `POS/CRM` experience unless the profile explicitly says so
- Soft skills must be visible across the resume, especially in the Professional Summary, the dynamic relevance bridge,
  and selected experience bullets: stakeholder management, executive communication, cross-functional collaboration,
  mentoring, vendor management, team leadership, problem-solving, and change leadership — only where supported by profile facts
- The goal is to maximise JD match through truthful semantic alignment, not keyword stuffing and not hallucination
- If the template includes a competencies/skills section, ensure it reflects a mix of technical and soft skills
- Keep EDUCATION & CERTIFICATIONS, LANGUAGES, and AVAILABILITY sections unchanged from profile data
- Total length: equivalent to the reference template (~2 pages of content)
- Every achievement bullet: SMART form, every distinct quantifiable figure **bolded** (see system rule 8)
- Weave people-management evidence into the most recent roles (see system rule 10)
- Do not state an exact years-of-experience number anywhere (see system rule 9)
- Keep spacing tight: a single blank line between sections is enough, no blank line directly under a section header before its content, and NO blank lines at all between the name/address/LinkedIn/website lines at the very top of the resume
{"- If recruiter resume-adjustment guidance was supplied above, weave it in naturally (headline framing, which bullets/section lead) — do not quote it or add a visible heading for it (see system rule 12)" if resume_adjustments_text else ""}
"""

    resume_text = call_llm(RESUME_SYS, RESUME_USER, max_tokens=20000, label="Resume")
    resume_text = soften_experience_years(resume_text)
    resume_text = tighten_contact_header(resume_text)
    resume_text = normalize_blank_lines(resume_text)
    print("  ✓  Resume complete")

# ══════════════════════════════════════════════════════════════════════════════
# TASK 3 — COVER LETTER (v2: bold-markup + no "27 years")
# ══════════════════════════════════════════════════════════════════════════════
coverletter_text = ""
if run_coverletter:
    print("✉️   [3/3] Generating Cover Letter...")

    # Extract contact details from profile metadata for the letter header
    meta     = profile.get("metadata", {})
    name     = meta.get("name", "John Hau")
    location = meta.get("location", "Hong Kong SAR")
    phone    = meta.get("phone", "+852 5722 2007")
    email    = meta.get("email", "haujon001@gmail.com")
    linkedin = meta.get("linkedin", "linkedin.com/in/johnhau")
    website  = "https://askcareer-ai.com"
    today    = datetime.now().strftime("%d %B %Y")

    COVERLETTER_SYS = (
        "You are a professional executive cover letter writer with deep expertise in senior technology leadership roles. "
        "Adapt the tone and emphasis to the specific JD and company context provided. "
        "Rules:\n"
        "1. ALL claims and achievements must come ONLY from the candidate profile data provided — no fabrication.\n"
        "2. Write in first person, confident executive tone, not generic or templated.\n"
        "3. The letter must be 4-5 substantive paragraphs:\n"
        "   Para 1 — Why this role at this company based on the JD/company context provided.\n"
        "   Para 2 — Most relevant leadership & operational experience (quantified from profile).\n"
        "   Para 3 — How the candidate's enterprise background translates to this specific role context.\n"
        "   Para 4 — Innovation, transformation, and people leadership as differentiators.\n"
        "   Para 5 — Closing: call to action, enthusiasm, availability.\n"
        "4. Use semantic matching to connect adjacent experience from the profile to the JD without overstating direct domain experience.\n"
        "5. Never claim direct experience with named domain-specific tools, standards, or industries unless they are explicitly present in the profile data.\n"
        "6. Do NOT use generic phrases like 'I am writing to apply for'. Open with impact.\n"
        "7. Wrap EVERY distinct quantifiable figure mentioned anywhere in the letter in double asterisks — percentages, dollar amounts, headcounts, user/device counts, time savings, etc., e.g. '**delivering HK$3.5M in savings**' or '**supporting 80,000 users**' — no cap on how many, bold every one that appears, but never bold a whole sentence.\n"
        "8. Whenever a specific job title is named together with a company (e.g. 'Associate Director of Infrastructure Services at AIA', 'VP, Asia Manager at Morgan Stanley'), wrap the title phrase itself in double asterisks too — do this consistently for every company/title mention in the letter, not just one.\n"
        "9. Never state an exact number of years of experience (e.g. do not write '27 years' or '27+ years'). Use wording like 'extensive years' or 'extensive experience' instead.\n"
        "10. The header block is ONLY: candidate name, location, phone | email, LinkedIn, website — one line each, no blank lines between them. Do NOT include a date line, a recipient name line ('Hiring Manager'), or a company name line anywhere before the salutation. Go directly from the header block to 'Dear Hiring Manager,'.\n"
        "11. Output ONLY the cover letter text — no preamble, no commentary (the ** bold markers from rules 7-8 are the one exception — those are expected).\n"
        + ("12. Recruiter resume-adjustment guidance (below) may be supplied — it comes from this JD's own Match Scorecard. Apply it ONLY to tone, emphasis, and which existing achievements get foregrounded. It must NEVER be used to introduce a fact, figure, project, or claim that is not already present in the candidate profile data. Never print the guidance verbatim or add a visible heading for it.\n" if resume_adjustment else "")
    )

    # Derive employer display name from employer slug (e.g. MandarinOriental -> Mandarin Oriental)
    import re as _re
    employer_display = _re.sub(r'(?<=[a-z])(?=[A-Z])', ' ', employer)  # CamelCase -> spaced
    role_display     = role_slug.replace('_', ' ')                      # slug -> readable

    coverletter_adjustment_block = (
        f"""
=== RECRUITER RESUME-ADJUSTMENT GUIDANCE (from this JD's own Match Scorecard — apply per system rule 12; never invent a fact to satisfy it) ===
{resume_adjustments_text}
"""
        if resume_adjustments_text else ""
    )

    COVERLETTER_USER = f"""
=== JOB DESCRIPTION ===
{jd_text}

=== OPTIONAL JD BLUEPRINT (use this if present, otherwise infer from JD text) ===
{jd_blueprint_context}

=== CANDIDATE PROFILE (single source of truth — all facts from here only) ===
{profile_context}
{coverletter_adjustment_block}
=== LETTER HEADER DETAILS ===
Name     : {name}
Location : {location}
Phone    : {phone}
Email    : {email}
LinkedIn : {linkedin}
Website  : {website}
(Context only, do not print as header lines — Date: {today}, Hiring Company: {employer_display})

=== INSTRUCTIONS ===
Write a formal executive cover letter for John Hau applying for the {role_display} role
at {employer_display}.

Structure:
- Letter header block, exactly these lines with no blank lines between them: Name / Location / Phone | Email / LinkedIn / Website
- No date line, no recipient name line, no company name line
- Salutation: Dear Hiring Manager,
- 4-5 focused paragraphs as described in the system instructions
- Professional sign-off

Key themes to hit (using only profile facts):
- The highest-priority business and technology themes from this specific JD
- Strong transferable leadership, operational, cybersecurity, transformation, vendor, and stakeholder evidence from the profile where relevant
- Budget governance and measurable savings where they support the business case
- Team building, coaching, and cross-functional partnership where relevant to the JD
- AI/automation innovation as a differentiator when it genuinely supports the role
- Immediate availability, Hong Kong-based, and Cantonese-speaking only if helpful and supported by the profile
- Use semantic alignment: connect adjacent evidence honestly rather than waiting for exact keyword matches
- Bold every distinct quantified achievement and every company/title mention (see system rules 7-8)
- Do not state an exact years-of-experience number anywhere (see system rule 9)
- No date/recipient-name/company lines before the salutation (see system rule 10)
{"- If recruiter resume-adjustment guidance was supplied above, weave it into tone/emphasis only — do not quote it or add a visible heading for it (see system rule 12)" if resume_adjustments_text else ""}
"""

    coverletter_text = call_llm(COVERLETTER_SYS, COVERLETTER_USER, max_tokens=8000, label="Cover Letter")
    coverletter_text = soften_experience_years(coverletter_text)
    coverletter_text = clean_coverletter_header(coverletter_text, employer_display)
    coverletter_text = normalize_blank_lines(coverletter_text)
    print("  ✓  Cover Letter complete")

# ══════════════════════════════════════════════════════════════════════════════
# WRITE OUTPUT FILES
# ══════════════════════════════════════════════════════════════════════════════
# Create all output dirs (txt + docx/pdf companion folders)
for d in [OUT_DIR_SCORECARD, OUT_DIR_RESUME, OUT_DIR_COVER]:
    d.mkdir(parents=True, exist_ok=True)
# Ensure sibling docx/pdf dirs also exist for future use
for subtype in ["ScoreCard", "resume", "CoverLetter"]:
    for fmt in ["docx", "pdf"]:
        (OUT_BASE / subtype / fmt).mkdir(parents=True, exist_ok=True)

generated_at = datetime.now().strftime("%Y-%m-%d %H:%M")
print("\n💾  Writing output files...")

if run_scorecard and scorecard_text:
    # Scorecard header format is unchanged from v1 (out of scope for this pass)
    header = (
        f"JD MATCH SCORECARD — {employer.upper()} | {role_slug.upper()}\n"
        f"Generated : {generated_at}  |  Model: {MODEL}\n"
        f"Source JD : {jd_path.name}\n"
        f"Profile   : src/data/john_profile.json\n"
        f"Candidate : John Hau\n"
        + "=" * 72 + "\n\n"
    )
    OUT_SCORECARD.write_text(header + scorecard_text, encoding="utf-8")
    print(f"  ✅  Scorecard   → {OUT_SCORECARD.relative_to(ROOT)}")
    maybe_convert_to_docx(OUT_SCORECARD, "Scorecard")

if run_resume and resume_text:
    # v2: no "TAILORED RESUME —" prefix, no Generated/Profile lines
    resume_header = (
        f"{employer.upper()} | {role_slug.upper()}\n"
        + "=" * 72 + "\n\n"
    )
    OUT_RESUME.write_text(resume_header + resume_text, encoding="utf-8")
    print(f"  ✅  Resume      → {OUT_RESUME.relative_to(ROOT)}")
    maybe_convert_to_docx(OUT_RESUME, "Resume")

if run_coverletter and coverletter_text:
    # v2: title prefix kept, Generated/Profile lines removed
    cl_header = (
        f"COVER LETTER — {employer.upper()} | {role_slug.upper()}\n"
        + "=" * 72 + "\n\n"
    )
    OUT_COVERLETTER.write_text(cl_header + coverletter_text, encoding="utf-8")
    print(f"  ✅  Cover Letter → {OUT_COVERLETTER.relative_to(ROOT)}")
    maybe_convert_to_docx(OUT_COVERLETTER, "Cover Letter")

print(f"\n  Output folder: data_processed/{employer}/")

print(f"\n{'='*60}")
print("  ALL DONE")
print(f"{'='*60}")
print(f"\n  How to run this script:")
print(f"  python scripts/jd_scorecard_resume_v2.py")
print(f"  python scripts/jd_scorecard_resume_v2.py \"data_raw/jd/txt/AnotherJD.txt\"")
print(f"  python scripts/jd_scorecard_resume_v2.py --scorecard-only")
print(f"  python scripts/jd_scorecard_resume_v2.py --resume-only")
print(f"  python scripts/jd_scorecard_resume_v2.py --coverletter-only")
print(f"  python scripts/jd_scorecard_resume_v2.py --batch")
print(f"  python scripts/jd_scorecard_resume_v2.py --batch --force")
print(f"  python scripts/jd_scorecard_resume_v2.py --refresh-blueprint")
print(f"  python scripts/jd_scorecard_resume_v2.py --ResumeAdjustment  (apply this JD's scorecard 6a guidance)")
print(f"  python scripts/jd_scorecard_resume_v2.py --llm=sonnet    (default, Claude Sonnet 5)")
print(f"  python scripts/jd_scorecard_resume_v2.py --llm=deepseek  (DeepSeek R1)")
print(f"  python scripts/jd_scorecard_resume_v2.py --llm=gemini    (Gemini Flash Lite - cheaper)")
print(f"  python scripts/jd_scorecard_resume_v2.py --llm=custom --model=<id> --provider=openrouter|deepseek")
print(f"  python scripts/jd_scorecard_resume_v2.py --llm=sonnet --api-key=<key>  (bring your own key)")
