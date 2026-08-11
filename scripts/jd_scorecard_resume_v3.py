#!/usr/bin/env python3
"""
All-in-One JD Application Generator — v3 (multi-profile)
=========================================================
v3 of scripts/jd_scorecard_resume_v2.py (kept alongside, not replacing it, per
soul.md's golden-rule: branch a new dated/versioned script rather than
editing a proven one in place). v2 is completely unaffected by this file and
remains the exclusive John Hau pipeline.

v3 adds multi-profile support (11 Aug 2026, Phase 1 of the multi-tenant JD
Portal epic — CLI foundation only, no auth/portal-UI/self-service-signup
work yet, see docs/guides/JDSCORECARDRESUMEV3_MULTIPROFILE_FOUNDATION_11AUG2026.md):
  - New REQUIRED `--profile=<Name>` flag resolves to src/data/<Name>/profile.json
    instead of the hardcoded john_profile.json. No default, no fallback —
    missing/invalid values fail loudly rather than silently touching anyone's
    data by accident.
  - Every path is namespaced per profile: src/data/<Name>/profile.json,
    src/data/<Name>/jd/ (blueprints), data_raw/<Name>/jd/txt/ (JD corpus),
    data_processed/<Name>/<Employer>/... (outputs). JD text and blueprints are
    forked per profile (not shared) so one profile's --batch run can never
    pick up another profile's job postings.
  - The reference resume template is now a generic, placeholder-only layout
    exemplar (data_raw/resume/txt/GenericStructuralResumeTemplate.md) instead
    of John's real personal resume content — avoids sending John's real PII
    through the LLM on every other profile's generation call.
  - All "John Hau" literal content (system/task prompts, API request title,
    banner, output headers, cover-letter website field) now derives from the
    loaded profile's own metadata.name at runtime.
  - Profile metadata (name/location/phone/email/linkedin) is loaded once,
    early, and validated — a profile missing any of those fields fails loudly
    instead of silently falling back to John's real contact details (v2's
    cover-letter fallback behavior, safe only because v2 only ever loads
    John's own profile).

Otherwise generates the same three documents as v2 from a Job Description +
a named profile's profile.json:
  1. JD Match Scorecard
  2. Tailored Resume
  3. Cover Letter

Usage:
  python scripts/jd_scorecard_resume_v3.py --profile=<Name>
  python scripts/jd_scorecard_resume_v3.py --profile=<Name> "data_raw/<Name>/jd/txt/AnotherJD.txt"
  python scripts/jd_scorecard_resume_v3.py --profile=<Name> --scorecard-only
  python scripts/jd_scorecard_resume_v3.py --profile=<Name> --resume-only
  python scripts/jd_scorecard_resume_v3.py --profile=<Name> --coverletter-only
  python scripts/jd_scorecard_resume_v3.py --profile=<Name> --refresh-blueprint
  python scripts/jd_scorecard_resume_v3.py --profile=<Name> --batch
  python scripts/jd_scorecard_resume_v3.py --profile=<Name> --batch --force
  python scripts/jd_scorecard_resume_v3.py --profile=<Name> --llm=sonnet|deepseek|gemini
  python scripts/jd_scorecard_resume_v3.py --profile=<Name> --ResumeAdjustment
  python scripts/jd_scorecard_resume_v3.py --profile=<Name> --llm=custom --model=anthropic/claude-opus-4-8 --provider=openrouter
  python scripts/jd_scorecard_resume_v3.py --profile=<Name> --llm=sonnet --api-key=sk-or-v1-...

--llm=custom / --model= / --provider= (carried over from v2, 23 Jul 2026):
  Run any OpenRouter or DeepSeek model id without editing LLM_CONFIGS —
  --model=<model id> plus --provider=openrouter|deepseek together select the
  model and endpoint/key-env, same as the portal's "Custom" LLM option.

--api-key=<key> (carried over from v2, 23 Jul 2026):
  Overrides the resolved API key for this run only (works with any --llm
  choice, preset or custom) — bring your own key without touching
  .env.local/.env.vps. Same override mechanism as the portal's "API Keys"
  settings panel, which stores a personal key server-side and injects it into
  just the spawned process's environment.

--ResumeAdjustment (carried over from v2, 22 Jul 2026):
  Pulls the "6a) Resume Adjustments" recommendations out of this JD's own Match
  Scorecard (auto-detects the latest scorecard already on disk for this
  employer/JD; generates one first if none exists yet) and feeds them to the
  Resume and Cover Letter prompts as extra tailoring guidance. It only shapes
  wording/emphasis/section framing in the existing output — it never adds a
  visible "Resume Adjustments" heading, and it can never introduce a fact,
  figure, or claim that isn't already in the profile's own profile.json (same
  single-source-of-truth / no-hallucination rule as everything else in this
  script).

Outputs go to data_processed/<Name>/<Employer>/ — filenames are date-stamped,
so re-running v3 against a JD already processed on a different day will not
collide or overwrite the earlier output.

Source: src/data/<Name>/profile.json (single source of truth — no hallucination)
"""

import json
import os
import re
import subprocess
import sys
import time
import requests
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Inches, Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ── Project root (one level up from scripts/) ─────────────────────────────────
ROOT = Path(__file__).resolve().parent.parent

# ── v3-only: resolve --profile=<Name> before anything else, since every other
# path in this script is namespaced under it. No default — a missing/invalid
# value fails loudly rather than silently touching any one profile's data.
_argv_flags = [a for a in sys.argv[1:] if a.startswith("-")]
_profile_flag = next((f for f in _argv_flags if f.startswith("--profile=")), None)
if not _profile_flag:
    sys.exit("ERROR: --profile=<Name> is required, e.g. --profile=MatinaFung (folder name under src/data/)")
PROFILE_NAME = _profile_flag.split("=", 1)[1]
if not re.fullmatch(r"[A-Za-z0-9_-]+", PROFILE_NAME):
    sys.exit(
        f"ERROR: --profile value '{PROFILE_NAME}' contains invalid characters — letters, digits, "
        f"hyphen, underscore only (no path separators or dots) to keep profile folders isolated."
    )

# ── Fixed paths ────────────────────────────────────────────────────────────────
# JD text and blueprints are namespaced per profile (not shared) so one
# profile's --batch run can never pick up another profile's job postings —
# blueprints are keyed only by JD filename, so a shared JD folder would mean
# every profile sees every other profile's JDs.
PROFILE_DIR        = ROOT / "src" / "data" / PROFILE_NAME
PROFILE_PATH       = PROFILE_DIR / "profile.json"
if PROFILE_DIR.resolve().parent != (ROOT / "src" / "data").resolve():
    sys.exit(f"ERROR: resolved profile path escapes src/data/ — refusing to proceed: {PROFILE_DIR}")
if not PROFILE_PATH.exists():
    sys.exit(
        f"ERROR: No profile found at {PROFILE_PATH.relative_to(ROOT)} — create it manually "
        f"(same top-level shape as src/data/john_profile.json: {{timestamp, sourceFile, profile: {{...}}}})."
    )
DEFAULT_JD_DIR     = ROOT / "data_raw" / PROFILE_NAME / "jd" / "txt"
JD_BLUEPRINT_DIR   = ROOT / "src" / "data" / PROFILE_NAME / "jd"
# Generic, placeholder-only layout exemplar — shared across all profiles (it's
# not the candidate's own data, so it doesn't need to be per-profile). Not
# John's real resume (see v3 docstring) — avoids sending his real personal
# content through the LLM on every other profile's generation call.
TEMPLATE_PATH      = ROOT / "data_raw/resume/txt/GenericStructuralResumeTemplate.md"

DATE_STAMP    = datetime.now().strftime("%d%b%Y").upper()   # e.g. 31MAR2026
RESUME_YEAR   = datetime.now().strftime("%Y")               # e.g. 2026

# ── v2-only: bold-markup + text-normalization helpers ─────────────────────────
BOLD_MARKUP_RE = re.compile(r"\*\*(.+?)\*\*")

# Matches "27 years", "27+ years", "27.33 years experience" etc. — case-insensitive
YEARS_FIGURE_RE = re.compile(r"\b27(?:\.\d+)?\+?\s+years?\b", re.IGNORECASE)


def soften_experience_years(text):
    """Replace any exact '27(+) years' figure with softer 'extensive years' wording."""
    return YEARS_FIGURE_RE.sub("extensive years", text)


def tighten_contact_header(text):
    """Remove ALL blank lines from the resume's top contact block (name,
    address, LinkedIn, website) — up to but excluding the first section
    separator line — leaving the rest of the document's spacing untouched."""
    lines = text.splitlines()
    sep_idx = next(
        (i for i, l in enumerate(lines) if set(l.strip()) <= {"=", "_", "-"} and len(l.strip()) >= 8),
        None,
    )
    if sep_idx is None:
        return text
    head = [l for l in lines[:sep_idx] if l.strip() != ""]
    return "\n".join(head + lines[sep_idx:])


DATE_LINE_RE = re.compile(r"^\d{1,2}\s+\w+\s+\d{4}$")


def clean_coverletter_header(text, employer_display):
    """Safety net: strip any date / 'Hiring Manager' / company-name lines and
    blank lines from the letter header block, in case the LLM didn't fully
    follow the prompt instructions. Stops at the salutation line."""
    lines = text.splitlines()
    sal_idx = next(
        (i for i, l in enumerate(lines) if l.strip().lower().startswith("dear ")),
        None,
    )
    if sal_idx is None:
        return text

    head = []
    for l in lines[:sal_idx]:
        stripped = l.strip()
        if not stripped:
            continue
        if stripped.lower() == "hiring manager":
            continue
        if stripped == employer_display:
            continue
        if DATE_LINE_RE.match(stripped):
            continue
        head.append(l)

    return "\n".join(head + lines[sal_idx:])


def normalize_blank_lines(text):
    """Collapse runs of 2+ blank lines to exactly 1, and drop a blank line that
    immediately follows an ALL-CAPS section header (tighter spacing per the
    user's manual V2 edits)."""
    lines = text.splitlines()
    out = []
    prev_blank = False
    prev_was_header = False
    for line in lines:
        stripped = line.strip()
        is_blank = not stripped
        is_header = bool(stripped) and stripped.isupper() and len(stripped) <= 80

        if is_blank:
            if prev_blank or prev_was_header:
                continue  # collapse double-blanks and header-then-blank
            out.append(line)
            prev_blank = True
            prev_was_header = False
            continue

        out.append(line)
        prev_blank = False
        prev_was_header = is_header

    return "\n".join(out)


def add_runs_with_markup(paragraph, text):
    """Add text to a paragraph, rendering **marked** spans as bold runs."""
    if text.count("**") % 2 != 0:
        # Unbalanced markup from the LLM — fall back to plain text rather
        # than leaking a stray "**" into the visible output.
        paragraph.add_run(text.replace("**", ""))
        return
    pos = 0
    for match in BOLD_MARKUP_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        bold_run = paragraph.add_run(match.group(1))
        bold_run.bold = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def resolve_resume_template():
    """Use the Markdown master resume template as the single authoritative source."""
    if TEMPLATE_PATH.exists():
        return TEMPLATE_PATH
    sys.exit(f"ERROR: Resume template not found: {TEMPLATE_PATH}")

TEMPLATE_PATH = resolve_resume_template()

def resolve_default_jd():
    """Pick the JD file explicitly provided via env, otherwise the most recently updated JD in data_raw/jd/txt/."""
    env_jd = os.environ.get("JD_DEFAULT_PATH", "").strip()
    if env_jd:
        env_path = Path(env_jd)
        if not env_path.is_absolute():
            env_path = ROOT / env_path
        if env_path.exists():
            return env_path

    jd_files = sorted(DEFAULT_JD_DIR.glob("JD_*.txt"), key=lambda p: p.stat().st_mtime, reverse=True)
    if jd_files:
        return jd_files[0]

    fallback_files = sorted(DEFAULT_JD_DIR.glob("*.txt"), key=lambda p: p.stat().st_mtime, reverse=True)
    if fallback_files:
        return fallback_files[0]

    sys.exit(f"ERROR: No JD text files found under {DEFAULT_JD_DIR}")

def load_jd_blueprint(jd_stem):
    """Optionally load a structured JD blueprint from src/data/jd/<JD_STEM>.json."""
    blueprint_path = JD_BLUEPRINT_DIR / f"{jd_stem}.json"
    if not blueprint_path.exists():
        return None, None

    try:
        blueprint = json.loads(blueprint_path.read_text(encoding="utf-8"))
        return blueprint_path, blueprint
    except Exception as exc:
        print(f"⚠️  Warning: Could not parse JD blueprint {blueprint_path.name}: {exc}")
        return blueprint_path, None


def derive_jd_metadata(path):
    jd_stem_local = path.stem
    jd_parts_local = jd_stem_local.replace("JD_", "").split("_")
    employer_local = jd_parts_local[0] if jd_parts_local else "Employer"
    role_slug_local = "_".join(jd_parts_local[1:]) if len(jd_parts_local) > 1 else "Role"
    role_tag_local = f"_{role_slug_local}" if role_slug_local and role_slug_local != "Role" else ""
    return jd_stem_local, employer_local, role_slug_local, role_tag_local


def build_output_targets(path):
    jd_stem_local, employer_local, role_slug_local, role_tag_local = derive_jd_metadata(path)
    out_base = ROOT / "data_processed" / PROFILE_NAME / employer_local
    return {
        "blueprint": JD_BLUEPRINT_DIR / f"{jd_stem_local}.json",
        "scorecard_dir": out_base / "ScoreCard" / "txt",
        "resume_dir": out_base / "resume" / "txt",
        "cover_dir": out_base / "CoverLetter" / "txt",
        "scorecard_pattern": f"JD_SCORECARD_{employer_local}{role_tag_local}_*.txt",
        "resume_pattern": f"{PROFILE_NAME}Resume*_{employer_local}{role_tag_local}_*.txt",
        "cover_pattern": f"{PROFILE_NAME}CoverLetter_{employer_local}{role_tag_local}_*.txt",
    }


def requested_outputs_exist(path, include_scorecard=True, include_resume=True, include_coverletter=True):
    targets = build_output_targets(path)
    checks = []
    if include_scorecard:
        checks.append(any(targets["scorecard_dir"].glob(targets["scorecard_pattern"])))
    if include_resume:
        checks.append(any(targets["resume_dir"].glob(targets["resume_pattern"])))
    if include_coverletter:
        checks.append(any(targets["cover_dir"].glob(targets["cover_pattern"])))
    return bool(checks) and all(checks)


# ── v2-only: --ResumeAdjustment support ────────────────────────────────────
# Extracts "6a) Resume Adjustments" out of a scorecard's own TAILORING
# RECOMMENDATIONS section, stopping at "b)"/"c)" or end of text.
RESUME_ADJUSTMENTS_RE = re.compile(
    r"\**\s*(?:a\)\s*)?Resume Adjustments\b[^\n]*\n(.*?)"
    r"(?=\n\s*\**\s*(?:[bc]\)\s*)?(?:Interview Preparation|Certifications)|\n\s*#{1,4}\s*\d|\Z)",
    re.IGNORECASE | re.DOTALL,
)


def find_latest_scorecard_txt(path):
    """Most recently modified existing scorecard .txt for this JD, if any."""
    targets = build_output_targets(path)
    candidates = sorted(
        targets["scorecard_dir"].glob(targets["scorecard_pattern"]),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )
    return candidates[0] if candidates else None


def extract_resume_adjustments(scorecard_text):
    """Pull the '6a) Resume Adjustments' block out of a JD Match Scorecard body."""
    match = RESUME_ADJUSTMENTS_RE.search(scorecard_text)
    return match.group(1).strip() if match else None

# ── Parse CLI arguments ────────────────────────────────────────────────────────
args = [a for a in sys.argv[1:] if not a.startswith("-")]
flags = [a for a in sys.argv[1:] if a.startswith("-")]

jd_path = Path(args[0]) if args else resolve_default_jd()
if not jd_path.is_absolute():
    jd_path = ROOT / jd_path
if not jd_path.exists():
    sys.exit(f"ERROR: JD file not found: {jd_path}")

run_scorecard     = "--resume-only" not in flags and "--coverletter-only" not in flags
run_resume        = "--scorecard-only" not in flags and "--coverletter-only" not in flags
run_coverletter   = "--scorecard-only" not in flags and "--resume-only" not in flags
refresh_blueprint = "--refresh-blueprint" in flags
batch_mode        = "--batch" in flags
force_run         = "--force" in flags
generate_docx     = "--no-docx" not in flags
resume_adjustment = any(f.lower() == "--resumeadjustment" for f in flags)

# ── LLM selection via --llm=<name> flag ───────────────────────────────────────
llm_flag = next((f for f in flags if f.startswith("--llm=")), "--llm=sonnet")
llm_choice = llm_flag.split("=", 1)[1].lower()   # sonnet | deepseek | gemini | custom

# ── Bring-your-own-key / dynamic model overrides (added 23 Jul 2026) ──────────
# --model= and --provider= together select an arbitrary OpenRouter/DeepSeek
# model id when --llm=custom, instead of one of the three fixed presets below.
# --api-key= overrides the resolved API key for ANY --llm choice (preset or
# custom) — lets this script be run with a personal key without touching
# .env.local, same override the portal's "API Keys" settings panel uses.
model_flag    = next((f for f in flags if f.startswith("--model=")), None)
provider_flag = next((f for f in flags if f.startswith("--provider=")), None)
api_key_flag  = next((f for f in flags if f.startswith("--api-key=")), None)
custom_model_id   = model_flag.split("=", 1)[1] if model_flag else None
custom_provider    = provider_flag.split("=", 1)[1].lower() if provider_flag else None
cli_api_key_override = api_key_flag.split("=", 1)[1] if api_key_flag else None

# ── Derive employer slug from JD filename for output naming ───────────────────
jd_stem    = jd_path.stem                               # e.g. JD_MandarinOriental_ClusterDirectorOfIT
jd_parts   = jd_stem.replace("JD_", "").split("_")
employer   = jd_parts[0] if jd_parts else "Employer"   # e.g. MandarinOriental
role_slug  = "_".join(jd_parts[1:]) if len(jd_parts) > 1 else "Role"
jd_blueprint_path, jd_blueprint = load_jd_blueprint(jd_stem)

# ── Output directories: data_processed/<Name>/<Employer>/<Type>/txt/ ──────────
OUT_BASE          = ROOT / "data_processed" / PROFILE_NAME / employer
OUT_DIR_SCORECARD = OUT_BASE / "ScoreCard" / "txt"
OUT_DIR_RESUME    = OUT_BASE / "resume"    / "txt"
OUT_DIR_COVER     = OUT_BASE / "CoverLetter" / "txt"
_role_tag         = f"_{role_slug}" if role_slug and role_slug != "Role" else ""
OUT_SCORECARD     = OUT_DIR_SCORECARD / f"JD_SCORECARD_{employer}{_role_tag}_{DATE_STAMP}.txt"
OUT_RESUME        = OUT_DIR_RESUME    / f"{PROFILE_NAME}Resume{RESUME_YEAR}_{employer}{_role_tag}_{DATE_STAMP}.txt"
OUT_COVERLETTER   = OUT_DIR_COVER     / f"{PROFILE_NAME}CoverLetter_{employer}{_role_tag}_{DATE_STAMP}.txt"

# ── v3-only: load the profile and validate/extract its metadata early, so the
# candidate's own name/contact details are available to the banner, the API
# request title, and every prompt below — not just deep inside the
# cover-letter block like in v2. Fails loudly on missing fields rather than
# silently falling back to anyone else's real contact details (v2's
# cover-letter fallback was only ever safe because v2 only loads John's own
# profile).
profile_raw = json.loads(PROFILE_PATH.read_text(encoding="utf-8"))
profile     = profile_raw.get("profile", profile_raw)

REQUIRED_METADATA_FIELDS = ["name", "location", "phone", "email", "linkedin"]
meta = profile.get("metadata", {})
_missing_meta = [f for f in REQUIRED_METADATA_FIELDS if not meta.get(f)]
if _missing_meta:
    sys.exit(
        f"ERROR: profile '{PROFILE_NAME}' metadata is missing required field(s): "
        f"{', '.join(_missing_meta)} — refusing to substitute placeholder/fallback contact "
        f"details. Add them under profile.metadata in {PROFILE_PATH.relative_to(ROOT)}."
    )
name, location, phone, email, linkedin = (meta[f] for f in REQUIRED_METADATA_FIELDS)
website = meta.get("website", "")

# ── Load env (.env.local → .env.vps → .env) — file always wins over stale shell ──
def load_env(path):
    if not path.exists():
        return False
    for line in path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if line and not line.startswith("#") and "=" in line:
            k, _, v = line.partition("=")
            os.environ[k.strip()] = v.strip()
    return True

for env_file in [ROOT / ".env.local", ROOT / ".env.vps", ROOT / ".env"]:
    if load_env(env_file):
        break

# ── Resolve API key and model based on LLM selection ─────────────────────────
# "deepseek-reasoner" was a legacy alias DeepSeek retired 24 Jul 2026 (it kept
# routing transparently to deepseek-v4-flash's thinking mode, but reliance on a
# retired alias risked breaking without notice) — migrated to the current
# official model id 3 Aug 2026. See LLM_DEEPSEEK_REASONING_EFFORT below for why
# runs were also timing out/truncating under this model's default high-effort
# thinking mode.
LLM_CONFIGS = {
    # name          model_id                                     api_key_env            base_url
    "sonnet":  ("anthropic/claude-sonnet-5",                "OPENROUTER_API_KEY",  "https://openrouter.ai/api/v1/chat/completions"),
    "deepseek":("deepseek-v4-flash",                        "DEEPSEEK_API_KEY",    "https://api.deepseek.com/chat/completions"),
    "gemini":  ("google/gemini-3.1-flash-lite-preview",     "OPENROUTER_API_KEY",  "https://openrouter.ai/api/v1/chat/completions"),
}

# provider name -> (api_key_env, base_url), used only when --llm=custom
PROVIDER_ENDPOINTS = {
    "openrouter": ("OPENROUTER_API_KEY", "https://openrouter.ai/api/v1/chat/completions"),
    "deepseek":   ("DEEPSEEK_API_KEY",   "https://api.deepseek.com/chat/completions"),
}

if llm_choice == "custom":
    if not custom_model_id or not custom_provider:
        sys.exit("ERROR: --llm=custom requires both --model=<model id> and --provider=openrouter|deepseek")
    if custom_provider not in PROVIDER_ENDPOINTS:
        sys.exit(f"ERROR: Unknown --provider value '{custom_provider}'. Valid: openrouter, deepseek")
    MODEL = custom_model_id
    api_key_env, LLM_ENDPOINT = PROVIDER_ENDPOINTS[custom_provider]
elif llm_choice in LLM_CONFIGS:
    MODEL, api_key_env, LLM_ENDPOINT = LLM_CONFIGS[llm_choice]
else:
    sys.exit(f"ERROR: Unknown --llm value '{llm_choice}'. Valid: sonnet, deepseek, gemini, custom")

API_KEY = cli_api_key_override or os.environ.get(api_key_env, "")
if not API_KEY:
    sys.exit(f"ERROR: {api_key_env} not found in .env.local / .env.vps / .env (or pass --api-key=<key>)")

HEADERS = {
    "Authorization": f"Bearer {API_KEY}",
    "Content-Type": "application/json",
    "HTTP-Referer": "https://www.askcareer-ai.com",
    "X-Title": f"JD Application Generator - {name}",
}

# Retried on: connection/timeout errors, and status codes observed in practice
# to be transient on OpenRouter's multi-provider routing (a request can 403
# against one upstream route and succeed seconds later against another for
# the identical key/model — confirmed 25 Jul 2026, not a real auth failure).
LLM_MAX_ATTEMPTS = 3
LLM_RETRY_BACKOFF_SEC = 3
LLM_RETRYABLE_STATUS_CODES = {403, 408, 425, 429, 500, 502, 503, 504}

# Reasoning models (e.g. DeepSeek's "deepseek-reasoner") spend part of the
# max_tokens budget on hidden chain-of-thought before writing the visible
# answer, so a call can hit finish_reason="length" with empty content even
# at a generous-looking max_tokens — the whole budget went to reasoning.
# Observed live 3 Aug 2026 on a Resume call (max_tokens=20000, deepseek-reasoner,
# portal ran out of tokens and surfaced the raw RuntimeError to the user).
# Instead of failing immediately, retry with a larger budget first.
LLM_LENGTH_RETRY_ATTEMPTS = 2
LLM_LENGTH_RETRY_MULTIPLIER = 2
LLM_MAX_TOKENS_CEILING = 64000

# DeepSeek's v4-flash/v4-pro default to thinking mode enabled with
# reasoning_effort="high" when the request doesn't specify otherwise — this is
# what made a --refresh-blueprint run (4 sequential DeepSeek calls) exceed the
# portal's 900s timeout on 3 Aug 2026, and was a contributing factor in the
# same day's max_tokens truncation bug. "low" keeps thinking mode on (still a
# reasoning model, unlike Sonnet/Gemini) while cutting reasoning-token spend
# enough to bring per-call latency back in line with the other providers.
LLM_DEEPSEEK_REASONING_EFFORT = "low"

if batch_mode:
    jd_files = sorted(DEFAULT_JD_DIR.glob("JD_*.txt"))
    if not jd_files:
        sys.exit(f"ERROR: No JD text files found under {DEFAULT_JD_DIR}")

    print(f"\n{'='*60}")
    print(f"  JD APPLICATION GENERATOR v3 — BATCH MODE — {PROFILE_NAME}")
    print(f"{'='*60}")
    print(f"  JD Folder   : {DEFAULT_JD_DIR.relative_to(ROOT)}")
    print(f"  File Count  : {len(jd_files)}")
    print(f"  LLM         : {llm_choice} ({MODEL})")
    print(f"  Skip Existing: {'Yes' if not force_run else 'No (force enabled)'}")
    print(f"{'='*60}\n")

    processed = 0
    skipped = 0
    failed = 0
    batch_results = []
    forwarded_flags = [f for f in flags if f not in {"--batch", "--force"}]

    for batch_jd in jd_files:
        print(f"\n▶ Processing {batch_jd.name}")
        targets = build_output_targets(batch_jd)
        blueprint_exists = targets["blueprint"].exists()

        if not force_run and requested_outputs_exist(batch_jd, run_scorecard, run_resume, run_coverletter):
            print("  ↷ Skipped — requested outputs already exist")
            skipped += 1
            batch_results.append({
                "jd": batch_jd.name,
                "status": "SKIPPED",
                "blueprint": "EXISTS" if blueprint_exists else "MISSING",
                "note": "Outputs already existed",
            })
            continue

        cmd = [sys.executable, str(Path(__file__).resolve()), str(batch_jd)] + forwarded_flags
        result = subprocess.run(cmd, cwd=str(ROOT))
        if result.returncode == 0:
            processed += 1
            batch_results.append({
                "jd": batch_jd.name,
                "status": "DONE",
                "blueprint": "READY",
                "note": f"Output -> data_processed/{derive_jd_metadata(batch_jd)[1]}/",
            })
        else:
            failed += 1
            print(f"  ✗ Failed — exit code {result.returncode}")
            batch_results.append({
                "jd": batch_jd.name,
                "status": "FAILED",
                "blueprint": "CHECK",
                "note": f"Exit code {result.returncode}",
            })

    print(f"\n{'='*60}")
    print("  BATCH SUMMARY")
    print(f"{'='*60}")
    print(f"  Processed : {processed}")
    print(f"  Skipped   : {skipped}")
    print(f"  Failed    : {failed}")

    print(f"\n{'-'*112}")
    print(f"  {'JD FILE':48} {'STATUS':10} {'BLUEPRINT':10} NOTE")
    print(f"{'-'*112}")
    for item in batch_results:
        jd_name = (item['jd'][:45] + '...') if len(item['jd']) > 48 else item['jd']
        print(f"  {jd_name:48} {item['status']:10} {item['blueprint']:10} {item['note']}")
    print(f"{'-'*112}")

    sys.exit(0 if failed == 0 else 1)

# ── v2-only: resolve --ResumeAdjustment source before anything prints ─────────
# Auto-detects the latest existing scorecard for this JD; if none exists yet,
# forces the scorecard to be generated this run so the guidance is available.
resume_adjustments_text = None
if resume_adjustment and (run_resume or run_coverletter):
    existing_scorecard_path = find_latest_scorecard_txt(jd_path)
    if existing_scorecard_path:
        resume_adjustments_text = extract_resume_adjustments(
            existing_scorecard_path.read_text(encoding="utf-8")
        )
        _scorecard_source_note = f"existing → {existing_scorecard_path.relative_to(ROOT)}"
    elif not run_scorecard:
        run_scorecard = True
        _scorecard_source_note = "none found — generating one first"
    else:
        _scorecard_source_note = "will extract from this run's freshly generated scorecard"
elif resume_adjustment:
    _scorecard_source_note = "ignored — no resume/cover-letter requested this run"

# ── Read source files ──────────────────────────────────────────────────────────
print(f"\n{'='*60}")
print(f"  JD APPLICATION GENERATOR v3 — {name}")
print(f"{'='*60}")
print(f"  JD File     : {jd_path.name}")
print(f"  JD Blueprint: {jd_blueprint_path.relative_to(ROOT) if jd_blueprint_path and jd_blueprint else 'dynamic from JD text'}")
print(f"  Template    : {TEMPLATE_PATH.relative_to(ROOT)}")
print(f"  Profile     : {PROFILE_PATH.relative_to(ROOT)}")
print(f"  Model       : {MODEL}")
print(f"  API Key     : {API_KEY[:25]}...")
print(f"  Date        : {DATE_STAMP}")
print(f"  LLM         : {llm_choice} ({MODEL})")
print(f"  Generating  : {'Scorecard ' if run_scorecard else ''}{'Resume ' if run_resume else ''}{'CoverLetter' if run_coverletter else ''}")
if resume_adjustment:
    print(f"  ResumeAdj   : {_scorecard_source_note}")
print(f"{'='*60}\n")

print("📂  Reading source files...")
jd_text      = jd_path.read_text(encoding="utf-8")
template_txt = TEMPLATE_PATH.read_text(encoding="utf-8")
# profile / profile_raw already loaded in the hoisted profile-load block above.

# Build the full profile context — all 13 profile.* sections.
# Cap is generous (model context window is 1M tokens) — it exists only to guard
# against a runaway/corrupt profile file, not to trim real content.
def build_profile_context(profile):
    sections = {}
    for key in ["metadata", "summary", "professional_experience", "major_achievements",
                "ai_projects", "core_competencies", "technical_skills",
                "education_certifications", "languages_spoken",
                "linkedin_recommendations", "soft_skills", "languages",
                "key_topics_for_qa"]:
        if key in profile:
            sections[key] = profile[key]
    return json.dumps(sections, indent=2, ensure_ascii=False)[:100000]

profile_context = build_profile_context(profile)

def build_blueprint_context(blueprint):
    return (
        json.dumps(blueprint, indent=2, ensure_ascii=False)[:6000]
        if blueprint
        else "No JD JSON blueprint found. Infer the priority criteria, must-have skills, soft skills, and weighting dynamically from the raw JD text."
    )

jd_blueprint_context = build_blueprint_context(jd_blueprint)

# ── LLM call helper ────────────────────────────────────────────────────────────
# Anthropic models support explicit prompt caching (via OpenRouter's pass-through of
# Anthropic's `cache_control` field). DeepSeek and Gemini already cache repeated
# prefixes automatically/implicitly on the provider side, so no explicit marker is
# needed (or supported the same way) for those.
CACHEABLE_MODEL_PREFIXES = ("anthropic/",)


def call_llm(system_prompt, user_prompt, max_tokens=6000, label="", response_format=None, cacheable_prefix=None):
    if label:
        print(f"  ↳ Calling OpenRouter ({MODEL}) — {label}")

    if cacheable_prefix and MODEL.startswith(CACHEABLE_MODEL_PREFIXES):
        # Mark the large, byte-identical context block (JD + blueprint + profile) as an
        # ephemeral cache breakpoint. Scorecard/Resume/Cover Letter all send this exact
        # same block in one run — after the first call it's served from cache instead
        # of being reprocessed from scratch, cutting repeated cost/latency.
        user_content = [
            {"type": "text", "text": cacheable_prefix, "cache_control": {"type": "ephemeral"}},
            {"type": "text", "text": user_prompt},
        ]
    else:
        user_content = (cacheable_prefix or "") + user_prompt

    payload = {
        "model": MODEL,
        "max_tokens": max_tokens,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_content},
        ],
    }
    if response_format:
        payload["response_format"] = response_format
    if "deepseek.com" in LLM_ENDPOINT:
        # Safe for both v4-flash and v4-pro: v4-pro currently only supports
        # high/max and treats "low" as "high" rather than rejecting it.
        payload["reasoning_effort"] = LLM_DEEPSEEK_REASONING_EFFORT

    resp = None
    message = {}
    content = None
    finish_reason = None
    usage = {}
    length_retries = 0
    for attempt in range(1, LLM_MAX_ATTEMPTS + 1):
        try:
            resp = requests.post(
                LLM_ENDPOINT,
                headers=HEADERS,
                json=payload,
                timeout=180,
            )
        except requests.exceptions.RequestException as exc:
            # Connection/timeout failures — always worth a retry until the last attempt.
            if attempt < LLM_MAX_ATTEMPTS:
                wait = LLM_RETRY_BACKOFF_SEC * attempt
                print(
                    f"  ⚠️  {label or 'unlabeled'} call failed ({exc}) "
                    f"(attempt {attempt}/{LLM_MAX_ATTEMPTS}) — retrying in {wait}s..."
                )
                time.sleep(wait)
                continue
            raise
        if resp.status_code in LLM_RETRYABLE_STATUS_CODES and attempt < LLM_MAX_ATTEMPTS:
            wait = LLM_RETRY_BACKOFF_SEC * attempt
            print(
                f"  ⚠️  {resp.status_code} from {label or 'unlabeled'} call "
                f"(attempt {attempt}/{LLM_MAX_ATTEMPTS}) — retrying in {wait}s..."
            )
            time.sleep(wait)
            continue
        resp.raise_for_status()  # raises immediately for non-retryable or attempts-exhausted errors

        data = resp.json()
        choice = data["choices"][0]
        message = choice["message"]
        content = message.get("content")
        finish_reason = choice.get("finish_reason")
        usage = data.get("usage", {})
        # OpenRouter reports Anthropic prompt-cache stats under prompt_tokens_details
        # (cached_tokens = served from cache, cache_write_tokens = newly cached this call).
        prompt_token_details = usage.get("prompt_tokens_details", {}) or {}
        cache_read = prompt_token_details.get("cached_tokens")
        cache_write = prompt_token_details.get("cache_write_tokens")
        if cache_read or cache_write:
            print(f"  ↳ Cache — read: {cache_read or 0} tokens, written: {cache_write or 0} tokens (cost: ${usage.get('cost', 0):.4f})")

        # Truncated on reasoning-token overrun: bump the budget and retry
        # rather than failing the whole pipeline run on the first hit.
        if finish_reason == "length" and length_retries < LLM_LENGTH_RETRY_ATTEMPTS:
            old_max = payload["max_tokens"]
            new_max = min(old_max * LLM_LENGTH_RETRY_MULTIPLIER, LLM_MAX_TOKENS_CEILING)
            if new_max > old_max and attempt < LLM_MAX_ATTEMPTS:
                length_retries += 1
                payload["max_tokens"] = new_max
                reasoning_tokens = usage.get("completion_tokens_details", {}).get("reasoning_tokens")
                print(
                    f"  ⚠️  {label or 'unlabeled'} call truncated (finish_reason='length', "
                    f"max_tokens={old_max}, reasoning_tokens={reasoning_tokens}) — retrying with "
                    f"max_tokens={new_max} (length retry {length_retries}/{LLM_LENGTH_RETRY_ATTEMPTS})..."
                )
                continue
        break

    if not content:
        refusal = message.get("refusal")
        raise RuntimeError(
            f"{MODEL} returned empty content for '{label or 'unlabeled'}' call "
            f"(finish_reason={finish_reason!r}, refusal={refusal!r}, "
            f"max_tokens={payload['max_tokens']}). This usually means the model ran out of "
            "max_tokens before producing an answer (reasoning models can spend the whole "
            "budget on hidden reasoning before writing a visible answer), or refused the "
            "request — check the values above."
        )
    if finish_reason == "length":
        raise RuntimeError(
            f"{MODEL} truncated the '{label or 'unlabeled'}' response — it ran out of "
            f"max_tokens ({payload['max_tokens']}) before finishing even after "
            f"{length_retries} automatic retry/retries at a higher budget "
            f"(completion_tokens={usage.get('completion_tokens')}, "
            f"reasoning_tokens={usage.get('completion_tokens_details', {}).get('reasoning_tokens')}). "
            "The partial output was discarded rather than silently shipped incomplete; raise "
            "the base max_tokens for this call and retry."
        )
    return content

def extract_json_object(raw_text):
    """Extract a JSON object from an LLM response, allowing for code fences and minor cleanup."""
    cleaned = raw_text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
        cleaned = re.sub(r"\s*```$", "", cleaned)

    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("LLM response did not contain a valid JSON object")

    candidate = cleaned[start:end + 1]
    candidate = (
        candidate
        .replace("“", '"')
        .replace("”", '"')
        .replace("‘", "'")
        .replace("’", "'")
    )

    attempts = [
        candidate,
        re.sub(r",\s*([}\]])", r"\1", candidate),
    ]

    last_error = None
    for attempt in attempts:
        try:
            return json.loads(attempt)
        except json.JSONDecodeError as exc:
            last_error = exc

    raise last_error


def repair_json_with_llm(raw_text, label="JSON Repair"):
    """Ask the model to repair malformed JSON and return valid JSON only."""
    REPAIR_SYS = (
        "You repair malformed JSON for production automation. "
        "Return VALID JSON ONLY. Do not add commentary, markdown, or explanations."
    )
    REPAIR_USER = f"""
Repair the following malformed JSON so it parses successfully.
Preserve the original meaning and keys as closely as possible.
Return only valid JSON.

=== MALFORMED JSON ===
{raw_text}
"""
    repaired = call_llm(
        REPAIR_SYS,
        REPAIR_USER,
        max_tokens=6000,
        label=label,
        response_format={"type": "json_object"},
    )
    return extract_json_object(repaired)


def style_docx_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    # python-docx's default template applies a document-wide 10pt "space after"
    # (w:docDefaults/w:pPrDefault) to every paragraph that doesn't explicitly
    # override it — including every blank spacer paragraph, section-title
    # line, and bullet. That silently stacked extra air on top of the blank
    # rows this converter deliberately inserts. Zero it at the Normal-style
    # level (List Bullet is based on Normal) so only the explicit Pt(3)/Pt(4)
    # overrides below ever add spacing.
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)


def add_docx_text_block(doc, line):
    stripped = line.strip()
    if not stripped:
        doc.add_paragraph()
        return

    if set(stripped) <= {"=", "_", "-"} and len(stripped) >= 8:
        return

    if stripped.startswith(("•", "* ", "- ")):
        bullet_text = re.sub(r"^[•*-]\s+", "", stripped, count=1).strip()
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.left_indent = Inches(0.25)
        add_runs_with_markup(para, bullet_text)
        return

    if stripped.isupper() and len(stripped) <= 80:
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = para.add_run(stripped)
        run.bold = True
        run.font.size = Pt(12)
        return

    if " — " in stripped and not stripped.endswith(":"):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(3)
        left, right = stripped.split(" — ", 1)
        run1 = para.add_run(left + " — ")
        run1.bold = True
        add_runs_with_markup(para, right)
        return

    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    add_runs_with_markup(para, stripped)


def _is_effectively_blank(line):
    """Blank line OR a ---/===/___ separator — both collapse to one gap in the docx."""
    stripped = line.strip()
    if not stripped:
        return True
    return set(stripped) <= {"=", "_", "-"} and len(stripped) >= 8


# "Company — Title" header line, as used for each Professional Experience entry.
_COMPANY_HEADER_RE = re.compile(r"^\S.*\s—\s\S")

# "Mon YYYY – Mon YYYY" / "Mon YYYY – Present" date-range line that follows
# each company header line.
_DATE_RANGE_RE = re.compile(
    r"^[A-Za-z]{3,9}\.?\s+\d{4}\s*[–—-]\s*(?:[A-Za-z]{3,9}\.?\s+\d{4}|Present)$"
)


def convert_text_file_to_docx(txt_path):
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed")

    doc = Document()
    style_docx_document(doc)

    # Pass 1: a section boundary is "blank line, separator, blank line" in the
    # .txt — each of those independently would add its own empty paragraph,
    # doubling the visual gap in Word. Collapse any run of blank/separator
    # lines into a single blank paragraph between real content lines.
    collapsed_lines = []
    blank_pending = False
    for line in txt_path.read_text(encoding="utf-8").splitlines():
        if _is_effectively_blank(line):
            blank_pending = True
            continue
        if blank_pending and collapsed_lines:
            collapsed_lines.append("")
        blank_pending = False
        collapsed_lines.append(line)

    # Pass 2: within a Professional Experience entry, drop the blank-line gap
    # between the "Company — Title" header and its date-range line right
    # below it, so those two rows sit adjacent with no spacing.
    normalized_lines = []
    i = 0
    while i < len(collapsed_lines):
        line = collapsed_lines[i]
        if (
            line == ""
            and 0 < i < len(collapsed_lines) - 1
            and _COMPANY_HEADER_RE.match(collapsed_lines[i - 1].strip())
            and _DATE_RANGE_RE.match(collapsed_lines[i + 1].strip())
        ):
            i += 1
            continue
        normalized_lines.append(line)
        i += 1

    for line in normalized_lines:
        add_docx_text_block(doc, line)

    target_dir = txt_path.parent.parent / "docx"
    target_dir.mkdir(parents=True, exist_ok=True)
    out_path = target_dir / f"{txt_path.stem}.docx"
    try:
        doc.save(out_path)
        return out_path, False
    except PermissionError:
        fallback_name = f"{txt_path.stem}_{datetime.now().strftime('%H%M%S')}.docx"
        fallback_path = target_dir / fallback_name
        doc.save(fallback_path)
        return fallback_path, True


def maybe_convert_to_docx(txt_path, label):
    if not generate_docx:
        return None
    if not DOCX_AVAILABLE:
        print(f"  ⚠️  {label} DOCX skipped — python-docx not installed")
        return None
    try:
        out_path, used_fallback = convert_text_file_to_docx(txt_path)
        if used_fallback:
            print(f"  ⚠️  {label} DOCX original file was locked; wrote fallback → {out_path.relative_to(ROOT)}")
        else:
            print(f"  ✅  {label} DOCX → {out_path.relative_to(ROOT)}")
        return out_path
    except Exception as exc:
        print(f"  ⚠️  {label} DOCX conversion failed: {exc}")
        return None


def generate_jd_blueprint(jd_text, jd_path, employer, role_slug):
    """Auto-create a structured JD blueprint for any new JD text file."""
    blueprint_path = JD_BLUEPRINT_DIR / f"{jd_stem}.json"
    JD_BLUEPRINT_DIR.mkdir(parents=True, exist_ok=True)

    BLUEPRINT_SYS = (
        "You are an expert job architecture analyst and recruiter. "
        "Read the JD and return VALID JSON ONLY for downstream automation. "
        "Do not use markdown code fences. Do not include commentary. "
        "Make the criteria practical, recruiter-style, and specific to this JD."
    )

    BLUEPRINT_USER = f"""
=== JOB DESCRIPTION ===
{jd_text}

=== INSTRUCTIONS ===
Create a JSON blueprint for this JD using the exact top-level keys below:
{{
  "jd_file": "...",
  "company": "...",
  "company_slug": "...",
  "job_title": "...",
  "role_slug": "...",
  "industry": "...",
  "role_summary": "...",
  "must_have_requirements": ["..."],
  "nice_to_have_requirements": ["..."],
  "soft_skills_priority": ["..."],
  "semantic_match_guidance": {{
    "category_name": ["related resume/profile evidence terms"]
  }},
  "scoring_criteria": [
    {{
      "name": "...",
      "weight": 0,
      "why_it_matters": "...",
      "evidence_type_to_look_for": ["..."],
      "gap_expected": "optional"
    }}
  ],
  "resume_tailoring_guidance": {{
    "mandatory_sections": ["PROFESSIONAL SUMMARY", "AI & AUTOMATION HIGHLIGHTS", "CORE COMPETENCIES", "SOFT SKILLS"],
    "summary_focus": ["..."],
    "do_not_overclaim": ["..."]
  }},
  "cover_letter_guidance": {{
    "tone": "...",
    "positioning": "...",
    "bridge_gaps_honestly": ["..."]
  }}
}}

Rules:
- Base the blueprint on the JD text only
- Include 8-10 scoring criteria
- Make all scoring weights numeric and total exactly 100
- Keep wording concise and usable for resume / scorecard / cover letter generation
- Include semantic matching hints so adjacent evidence can be matched truthfully without hallucination
"""

    raw = call_llm(
        BLUEPRINT_SYS,
        BLUEPRINT_USER,
        max_tokens=6000,
        label="JD Blueprint",
        response_format={"type": "json_object"},
    )
    try:
        blueprint = extract_json_object(raw)
    except Exception as exc:
        print(f"  ↳ Repairing malformed JD blueprint JSON ({exc})")
        blueprint = repair_json_with_llm(raw, label="JD Blueprint Repair")

    blueprint.setdefault("jd_file", str(jd_path.relative_to(ROOT)).replace("\\", "/"))
    blueprint.setdefault("company", employer)
    blueprint.setdefault("company_slug", employer)
    blueprint.setdefault("job_title", role_slug.replace("_", " "))
    blueprint.setdefault("role_slug", role_slug)
    blueprint.setdefault("resume_tailoring_guidance", {})
    blueprint["resume_tailoring_guidance"].setdefault(
        "mandatory_sections",
        ["PROFESSIONAL SUMMARY", "AI & AUTOMATION HIGHLIGHTS", "CORE COMPETENCIES", "SOFT SKILLS"],
    )

    blueprint_path.write_text(json.dumps(blueprint, indent=2, ensure_ascii=False), encoding="utf-8")
    return blueprint_path, blueprint

if refresh_blueprint or not jd_blueprint:
    action = "Refreshing" if refresh_blueprint and jd_blueprint_path else "Generating"
    print(f"🧭  [0/3] {action} JD blueprint...")
    try:
        jd_blueprint_path, jd_blueprint = generate_jd_blueprint(jd_text, jd_path, employer, role_slug)
        jd_blueprint_context = build_blueprint_context(jd_blueprint)
        print(f"  ✓  JD Blueprint ready -> {jd_blueprint_path.relative_to(ROOT)}")
    except Exception as exc:
        print(f"  ⚠️  JD blueprint generation skipped: {exc}")
        jd_blueprint_context = build_blueprint_context(jd_blueprint)

# Byte-identical across the Scorecard/Resume/Cover Letter calls below in this run — built
# once (after any blueprint refresh above) so it can be passed as `cacheable_prefix` to
# call_llm() and cache-hit on the 2nd/3rd/4th call instead of being reprocessed from scratch.
def build_shared_context_block(jd_text, jd_blueprint_context, profile_context):
    return f"""=== JOB DESCRIPTION ===
{jd_text}

=== OPTIONAL JD BLUEPRINT (use this if present, otherwise infer from JD text) ===
{jd_blueprint_context}

=== CANDIDATE PROFILE (single source of truth — all facts from here only, never invent) ===
{profile_context}
"""

SHARED_CONTEXT_BLOCK = build_shared_context_block(jd_text, jd_blueprint_context, profile_context)

# Anthropic prompt caching hashes the FULL prefix up to a cache_control breakpoint —
# system prompt included — so a byte-identical SHARED_CONTEXT_BLOCK only cache-hits
# across calls if the system message is ALSO byte-identical. Scorecard/Resume/Cover
# Letter each need different task rules, so those rules move into each call's own
# user-turn content (after the cache breakpoint) instead of living in the system
# message, which is now this one shared, generic string for all three calls.
SHARED_SYSTEM_PROMPT = (
    f"You are an expert executive career-document assistant for candidate {name} — "
    "acting as recruiter/analyst, resume writer, or cover-letter writer depending on the "
    "task rules given in the user message. Follow those task-specific rules exactly. "
    "ALL facts, figures, dates, and claims must come ONLY from the candidate profile data "
    "provided in the user message — never invent or embellish."
)

# ── v2-only: pre-computed recommendations summary (added 5 Aug 2026, moved out of the
# Cover Letter section 5 Aug 2026 so the Resume can use it too) ────────────────────
# Ground the "social proof" line in real, computed data instead of letting the LLM
# eyeball/re-derive it from 16 raw paragraphs each run — count and theme frequency are
# computed here in Python, not by the model.
RECOMMENDATION_THEME_KEYWORDS = {
    "client-focused":     ["client-focus", "client focus", "client first", "client-first", "put client", "clients first"],
    "detail-oriented":    ["detail-mind", "detailed-mind", "detail oriented", "attention to detail", "detailed,"],
    "dedicated / driven": ["dedicat"],
    "collaborative / strong communicator": ["communicat", "collaborat", "coordinat", "facilitat"],
    "innovative / creative": ["innovat", "creativ"],
    "result-oriented":    ["result-oriented", "result oriented"],
    "trusted / knowledgeable expert": ["knowledgeable", "invaluable", "expertise", " expert"],
}


def build_recommendations_summary(profile):
    recs = profile.get("linkedin_recommendations", [])
    if not recs:
        return None
    count = len(recs)

    def _bucket(relationship):
        rel_l = (relationship or "").lower()
        if "vendor" in rel_l:
            return "vendor"
        if "client" in rel_l:
            return "client"
        return "colleague"

    buckets = {"colleague": 0, "client": 0, "vendor": 0}
    for r in recs:
        buckets[_bucket(r.get("relationship", ""))] += 1

    theme_counts = {}
    for r in recs:
        text_l = (r.get("recommendation") or "").lower()
        for theme, keywords in RECOMMENDATION_THEME_KEYWORDS.items():
            if any(kw in text_l for kw in keywords):
                theme_counts[theme] = theme_counts.get(theme, 0) + 1
    ranked_themes = sorted(theme_counts.items(), key=lambda kv: kv[1], reverse=True)

    relationship_parts = []
    if buckets["colleague"]:
        relationship_parts.append(f"{buckets['colleague']} colleagues/team members")
    if buckets["client"]:
        relationship_parts.append(f"{buckets['client']} clients")
    if buckets["vendor"]:
        relationship_parts.append(f"{buckets['vendor']} vendors")

    theme_lines = "\n".join(f"- {name}: mentioned in {n} of {count}" for name, n in ranked_themes)
    return (
        f"Total recommendations: {count}\n"
        f"Relationship mix: {', '.join(relationship_parts)}\n"
        f"Recurring themes, ranked by how many of the {count} recommendations mention them "
        f"(computed from the actual recommendation text — pick whichever 2-3 are most relevant "
        f"to THIS JD, not always the same fixed set):\n"
        f"{theme_lines}\n"
    )

recommendations_summary = build_recommendations_summary(profile)

# ══════════════════════════════════════════════════════════════════════════════
# TASK 1 — JD SCORECARD (unchanged from v1)
# ══════════════════════════════════════════════════════════════════════════════
scorecard_text = ""
if run_scorecard:
    print("🔍  [1/3] Generating JD Scorecard...")

    SCORECARD_USER = f"""
=== ROLE & TASK RULES (moved out of the system message so it stays a shared, cacheable
    prefix across the Scorecard/Resume/Cover Letter calls in this run — see SHARED_SYSTEM_PROMPT) ===
You are an expert executive recruiter and talent analyst for this task.
Produce a detailed, honest, structured scorecard comparing the candidate to the JD.
Be evidence-based and professional. Never invent facts not present in the profile data.
Use plain text with clear section headers. Derive the scoring criteria dynamically from the specific JD,
focus only on materially relevant requirements, and compute a weighted overall score.
When evidence is adjacent or transferable rather than exact, say so explicitly and do not present it as direct hands-on experience.

=== INSTRUCTIONS ===
Produce a comprehensive JD Match Scorecard with ALL of these sections:

1. ROLE & COMPANY SUMMARY
   Brief overview of the role and what the employer is seeking (3-4 sentences).

2. OVERALL MATCH SCORE
   Format:  XX/100 — [STRONG MATCH / GOOD MATCH / PARTIAL MATCH / WEAK MATCH]
   Followed by one honest paragraph verdict.
   IMPORTANT: compute this as a weighted result from the JD-specific criteria below.
   Do NOT penalise the candidate for criteria that are not actually important in this JD.

3. DYNAMIC SCORING CRITERIA
   - First identify the 8-12 MOST IMPORTANT criteria directly from THIS JD.
   - These criteria MUST change depending on the role. Do not reuse a fixed checklist.
   - Use role-relevant categories only, for example when applicable:
     enterprise architecture, solution/integration design, stakeholder engagement,
     AI innovation, hospitality systems, cyber/compliance, vendor leadership,
     business partnership, programme delivery, people leadership, communication.
   - If the optional JD JSON blueprint exists, use its priorities/criteria/weights as the starting point.
   - If no JSON exists, infer the criteria directly from the raw JD.
   - Use semantic matching, not just exact keyword matching. Example: if the JD asks for cybersecurity,
     and the profile shows IT security, risk, audit, DR, patching, compliance, or resilience work,
     connect those dots clearly as relevant evidence.
   - Reward strong transferable evidence when it is meaningfully adjacent to the JD, but do NOT overstate or invent direct experience.
   - For each criterion provide:
     * Criterion Name
     * Weight (%) based on JD importance — all weights must total 100
     * Score: X/10
     * Evidence: specific proof from the candidate profile data
     * Gap: what is missing or weaker versus the JD

4. KEY STRENGTHS
   Bullet list — what clearly matches or exceeds the JD requirements.

5. GAPS & RISKS
   Honest, specific gaps between the candidate profile and JD expectations.

6. TAILORING RECOMMENDATIONS
   a) Resume adjustments — specific wording and framing suggestions
   b) Interview preparation — topics to research, answers to rehearse
   c) Certifications to pursue — quick wins before applying

7. VERDICT
   Should the candidate apply? Competitive positioning. Recommended approach.
"""

    scorecard_text = call_llm(
        SHARED_SYSTEM_PROMPT, SCORECARD_USER, max_tokens=12000, label="Scorecard",
        cacheable_prefix=SHARED_CONTEXT_BLOCK,
    )
    print("  ✓  Scorecard complete")

    if resume_adjustment and resume_adjustments_text is None:
        resume_adjustments_text = extract_resume_adjustments(scorecard_text)
        if not resume_adjustments_text:
            print("  ⚠️  --ResumeAdjustment: could not find a '6a) Resume Adjustments' section in this scorecard")

# ══════════════════════════════════════════════════════════════════════════════
# TASK 2 — TAILORED RESUME (v2: SMART + bold-markup + no "27 years")
# ══════════════════════════════════════════════════════════════════════════════
resume_text = ""
if run_resume:
    print("📝  [2/3] Generating Tailored Resume...")

    RESUME_TASK_RULES = (
        "=== ROLE & TASK RULES (moved out of the system message so it stays a shared, cacheable\n"
        "    prefix across the Scorecard/Resume/Cover Letter calls in this run — see SHARED_SYSTEM_PROMPT) ===\n"
        "You are a professional executive resume writer specialising in senior IT leadership roles for this task. "
        "Rules you MUST follow:\n"
        "1. ALL facts, figures, dates, and claims come ONLY from the candidate profile — never invent or embellish.\n"
        "2. Mirror the reference template layout EXACTLY: same section headers, same ___ separators, same bullet style.\n"
        "3. Tailor bullet wording to echo the JD language without distorting facts.\n"
        "4. Use semantic alignment for related skills, but NEVER claim direct hands-on experience with named systems, standards, tools, or industries unless they appear explicitly in the profile data.\n"
        f"5. If the JD asks for hospitality/PMS/POS/PCI or other domain-specific items not explicitly in the profile, position {name}'s background as transferable or adjacent experience only.\n"
        "6. The `AI & AUTOMATION HIGHLIGHTS` section is mandatory in every resume output. Populate it ONLY from `profile.ai_projects` — never from the reference template's own AI/Automation bullets, which are a LAYOUT example only, not a content source. Select and order the `ai_projects` entries most relevant to THIS JD's themes (e.g. a trading-automation JD should favor the trading/ML projects; an EUC/VDI-heavy JD should favor the automation/dashboard projects), same JD-relevance-first principle as rule 10 below — do not just reproduce the template's fixed example list regardless of the JD.\n"
        "7. Make the resume balanced: show both technical depth and leadership/soft skills supported by the profile.\n"
        "8. Every achievement bullet must be written in SMART form (Specific, Measurable, Achievable, Relevant, Time-bound). Wrap EVERY distinct quantifiable figure in the bullet in double asterisks — percentages, dollar/HK$ amounts, headcounts, device/user counts, time savings, etc. — no cap on how many per bullet; bold every one that appears, but never wrap overlapping or adjacent text as a single run. The reference template shows this convention in a few of its bullets; follow that pattern for every bullet you write.\n"
        "9. Within each role, order bullets by impact — lead with the most quantified, highest-impact achievements first, descending to more routine/supporting bullets last.\n"
        "10. When choosing which achievements fill a company's fixed bullet budget, JD relevance is the primary filter — but when multiple candidate highlights are comparably relevant to the JD, prefer the ones with larger quantified impact (bigger $ savings/revenue, larger user/team/device counts, wider organizational scope) over more routine ones. Never include a highly-quantified but JD-irrelevant bullet ahead of a genuinely JD-relevant one just because its number is bigger — relevance always wins ties in the other direction too.\n"
        "11. Preserve the exact company order from the candidate profile's `professional_experience` array (most recent role first, exactly as listed) — never reorder, merge, or resequence companies relative to that array.\n"
        "12. Do not state the exact computed years-of-experience figure (e.g. '27 years' or '27.33 years'). A rounded, approximate figure like '25+ years' is acceptable if it strengthens positioning, but default to 'extensive years'/'extensive experience' when no rounding is natural. For any SINGLE role's tenure (e.g. one company), never use a precise decimal duration (e.g. '9.5 years') — use a rounded phrase like 'nearly a decade' or a whole-number-plus figure instead.\n"
        "13. Weave visible people-management evidence into the most recent 2-3 roles specifically (not just a generic soft-skills line) — team leadership, mentoring/coaching, hiring or onboarding, performance management, career development, org design — using only what the profile actually documents (e.g. leading a team of 50+, coaching teams across multiple countries, mentorship and team development).\n"
        "14. Output ONLY the resume text — no preamble, no explanation, no markdown code fences (the ** bold markers from rule 8 are the one exception — those are expected).\n"
        "15. Professional Summary company references: (a) when naming companies in a list (e.g. \"...financial institutions including X, Y, Z...\"), list them in the exact same order as `professional_experience` in the candidate profile (most recent first) — do NOT copy the reference template's fixed company order verbatim. (b) The reference template's summary includes a dedicated 'spotlight' sentence naming one company and its flagship achievement (e.g. the Morgan Stanley VP / Workspace Virtualization / 120,000 desktops line) — do NOT default to spotlighting Morgan Stanley in every resume regardless of the JD. Instead, choose the spotlight company dynamically: pick whichever role's achievements are most relevant to THIS JD's themes, breaking ties in favor of the most recent role. Rewrite that sentence using that company's own title, scope, and achievements from the profile — only reuse the template's Morgan Stanley wording if Morgan Stanley is genuinely the best fit for this specific JD.\n"
        + ("16. A RECOMMENDATIONS SUMMARY block may be supplied below (pre-computed, real counts — never invent a different number or theme). If present, weave ONE brief sentence of social proof into the Professional Summary (2nd or 3rd sentence) citing the exact total-recommendations count and choosing whichever 2-3 listed themes are most relevant to THIS JD's priorities — do not always pick the same fixed themes across different JDs, and do not name individual recommenders or quote a specific recommendation verbatim unless the block explicitly provides a quotable line.\n" if recommendations_summary else "")
        + ("17. Recruiter resume-adjustment guidance (below) may be supplied — it comes from this JD's own Match Scorecard. Apply it ONLY to wording, emphasis, section framing, and which existing facts get foregrounded. It must NEVER be used to introduce a fact, figure, project, or claim that is not already present in the candidate profile data — the guidance changes how real facts are presented, never what the facts are. Never print the guidance verbatim or add a visible \"Resume Adjustments\" heading.\n" if resume_adjustment else "")
    )

    resume_adjustment_block = (
        f"""
=== RECRUITER RESUME-ADJUSTMENT GUIDANCE (from this JD's own Match Scorecard — apply per rule 17; never invent a fact to satisfy it) ===
{resume_adjustments_text}
"""
        if resume_adjustments_text else ""
    )

    resume_recommendations_summary_block = (
        f"""
=== RECOMMENDATIONS SUMMARY (pre-computed from profile.linkedin_recommendations — apply per rule 16; never invent a different count or theme) ===
{recommendations_summary}
"""
        if recommendations_summary else ""
    )

    RESUME_USER_PREFIX = f"""
{RESUME_TASK_RULES}
{resume_recommendations_summary_block}
=== REFERENCE RESUME TEMPLATE (layout/format exemplar ONLY — headers, separators, bullet style; never a source of facts or which achievements to include, see rule 6 for the one section this is called out on explicitly) ===
{template_txt}
"""

    RESUME_USER = f"""{resume_adjustment_block}
=== INSTRUCTIONS ===
Write a complete tailored resume for {name} targeting the role above.

Layout rules:
- Keep the identical header format, ________ separators, and bullet style from the reference template
- Respect and preserve any existing soft-skills section already present in the reference template, while tailoring it to this JD using only profile facts
- The `AI & AUTOMATION HIGHLIGHTS` section is mandatory and must never be omitted — populate it from `profile.ai_projects` selected/ordered for JD relevance, not from the template's own example bullets (see rule 6)
- First infer the 5-8 most important themes from THIS JD (technical, leadership, business, industry, and interpersonal)
- Replace the Professional Summary to foreground those JD-specific themes using only evidence from the profile
- Follow rule 15 for the summary's company list order and spotlight-company selection — do not blindly copy the reference template's Morgan Stanley-centric company order or spotlight sentence
{"- If a RECOMMENDATIONS SUMMARY block was supplied above, weave one brief social-proof sentence into the Professional Summary using its real count and whichever themes best fit this JD (see rule 16)" if recommendations_summary else ""}
- Add a dynamic bridging section immediately before PROFESSIONAL EXPERIENCE. The section title and content must fit the JD nature
  (for example: "HOSPITALITY RELEVANCE", "ROLE RELEVANCE", or "ENTERPRISE ARCHITECTURE RELEVANCE") and should never be hardcoded
- BULLET COUNT RULES (non-negotiable):
    * 1st company (most recent):  exactly 12 bullets
    * 2nd company:                exactly 12 bullets
    * 3rd company:                exactly 10 bullets
    * 4th company:                exactly 10 bullets
    * 5th company:                exactly 8 bullets
    * Any earlier roles:          3-4 bullets combined
- Company order must match `professional_experience` exactly, most recent first — never resequence (see rule 11)
- Choose the bullets most relevant to the JD — reword to echo JD language without distorting facts. Among comparably JD-relevant candidates within a company, prefer larger quantified impact ($ savings/revenue, user/team/device counts, scope) over routine ones (see rule 10)
- Within each company's bullets, lead with the most quantified/highest-impact achievements first (see rule 9)
- Use semantic matching to connect related evidence from the candidate profile data to the JD. Example: cybersecurity ↔ IT security / audit / risk / DR / compliance;
  stakeholder management ↔ executive presentations / business partnering; operational excellence ↔ SLA, uptime, automation, service quality improvements
- Do NOT convert adjacent evidence into unsupported exact claims. Example: do not say John has direct `PCI-DSS`, `Opera PMS`, or hotel `POS/CRM` experience unless the profile explicitly says so
- Soft skills must be visible across the resume, especially in the Professional Summary, the dynamic relevance bridge,
  and selected experience bullets: stakeholder management, executive communication, cross-functional collaboration,
  mentoring, vendor management, team leadership, problem-solving, and change leadership — only where supported by profile facts
- The goal is to maximise JD match through truthful semantic alignment, not keyword stuffing and not hallucination
- If the template includes a competencies/skills section, ensure it reflects a mix of technical and soft skills
- Keep EDUCATION & CERTIFICATIONS, LANGUAGES, and AVAILABILITY sections unchanged from profile data
- Total length: equivalent to the reference template (~2 pages of content)
- Every achievement bullet: SMART form, every distinct quantifiable figure **bolded** (see rule 8)
- Weave people-management evidence into the most recent roles (see rule 13)
- Do not state an exact years-of-experience number anywhere (see rule 12)
- Keep spacing tight: a single blank line between sections is enough, no blank line directly under a section header before its content, and NO blank lines at all between the name/address/LinkedIn/website lines at the very top of the resume
{"- If recruiter resume-adjustment guidance was supplied above, weave it in naturally (headline framing, which bullets/section lead) — do not quote it or add a visible heading for it (see rule 15)" if resume_adjustments_text else ""}
"""

    resume_text = call_llm(
        SHARED_SYSTEM_PROMPT, RESUME_USER_PREFIX + RESUME_USER, max_tokens=20000, label="Resume",
        cacheable_prefix=SHARED_CONTEXT_BLOCK,
    )
    resume_text = soften_experience_years(resume_text)
    resume_text = tighten_contact_header(resume_text)
    resume_text = normalize_blank_lines(resume_text)
    print("  ✓  Resume complete")

# ══════════════════════════════════════════════════════════════════════════════
# TASK 3 — COVER LETTER (v2: bold-markup + no "27 years")
# ══════════════════════════════════════════════════════════════════════════════
coverletter_text = ""
if run_coverletter:
    print("✉️   [3/3] Generating Cover Letter...")

    # name/location/phone/email/linkedin/website already resolved earlier
    # (hoisted profile-load block, right after CLI parsing) — reused here as-is.
    today = datetime.now().strftime("%d %B %Y")

    COVERLETTER_TASK_RULES = (
        "=== ROLE & TASK RULES (moved out of the system message so it stays a shared, cacheable\n"
        "    prefix across the Scorecard/Resume/Cover Letter calls in this run — see SHARED_SYSTEM_PROMPT) ===\n"
        "You are a professional executive cover letter writer with deep expertise in senior technology leadership roles for this task. "
        "Adapt the tone and emphasis to the specific JD and company context provided. "
        "Rules:\n"
        "1. ALL claims and achievements must come ONLY from the candidate profile data provided — no fabrication.\n"
        "2. Write in first person, confident executive tone, not generic or templated.\n"
        "3. The letter must be 5 tight paragraphs (plus sign-off) and stay concise — a hiring manager should be able to read it quickly, not a company-by-company stat dump. Cover ONLY the 5 most recent entries in `professional_experience` (skip any entries beyond the 5th — those are early-career/entry-level roles from before the candidate's senior IT career began and are not relevant to the senior leadership positions now being targeted; if the array has fewer than 5 entries, cover all of them):\n"
        "   Para 1 — Declarative opening: name the 5 most recent companies from `professional_experience`, in the exact same order as that array (most recent first) — e.g. 'Across [Company A], [Company B], [Company C], [Company D], and [Company E], I have led global infrastructure teams, stabilized mission-critical environments, and delivered measurable business impact across APAC, EMEA, and North America.' State cross-career impact directly. Do NOT open with a 'why this role at this company' framing or restate what the role/company does — the reader already knows that.\n"
        "   Para 2 — Recommendations summary ONLY (see rule 12) — a short, standalone paragraph, 1-2 sentences, immediately after the opening. No achievement detail here; it exists purely as a social-proof beat before the substantive achievement paragraphs.\n"
        "   Para 3 — The 2-3 strongest, most JD-relevant, ATS-strengthening quantified achievements in the ENTIRE letter, drawn from whichever of the 5 companies are most relevant to this JD — go deep here with real numbers; this is the letter's substantive core, not an exhaustive company-by-company stat dump.\n"
        "   Para 4 — The remaining companies from the 5, covered briefly (role/scope plus at most one light achievement each, roughly one sentence per company) so continuity and keyword coverage are preserved without repeating Para 3's depth. The oldest of the 5 (typically Merrill Lynch) may close this paragraph with a brief early notable recognition/award as a 'leadership foundation' note if the profile documents one.\n"
        "   Para 5 — Innovation/transformation/people leadership as a differentiator, then closing: call to action, enthusiasm, availability — combine into one final paragraph to keep the letter tight.\n"
        "4. Use semantic matching to connect adjacent experience from the profile to the JD without overstating direct domain experience.\n"
        "5. Never claim direct experience with named domain-specific tools, standards, or industries unless they are explicitly present in the profile data.\n"
        "6. Do NOT use generic phrases like 'I am writing to apply for'. Open with impact.\n"
        "7. Wrap EVERY distinct quantifiable figure mentioned anywhere in the letter in double asterisks — percentages, dollar amounts, headcounts, user/device counts, time savings, etc., e.g. '**delivering HK$3.5M in savings**' or '**supporting 80,000 users**' — no cap on how many, bold every one that appears, but never bold a whole sentence.\n"
        "8. Whenever a specific job title is named together with a company (e.g. 'Associate Director of Infrastructure Services at AIA', 'VP, Asia Manager at Morgan Stanley'), wrap the title phrase itself in double asterisks too — do this consistently for every company/title mention in the letter, not just one.\n"
        "9. Do not state the exact computed years-of-experience figure (e.g. '27 years' or '27.33 years'). A rounded, approximate figure like '25+ years' is acceptable if it strengthens positioning, but default to 'extensive years'/'extensive experience' when no rounding is natural. For any SINGLE company's tenure, never use a precise decimal duration (e.g. '9.5 years') — use a rounded phrase like 'nearly a decade' or a whole-number-plus figure instead.\n"
        f"10. The header block is ONLY: candidate name, location, phone | email, LinkedIn{', website' if website else ''} — one line each, no blank lines between them. Do NOT include a date line, a recipient name line ('Hiring Manager'), or a company name line anywhere before the salutation. Go directly from the header block to 'Dear Hiring Manager,'.\n"
        "11. Output ONLY the cover letter text — no preamble, no commentary (the ** bold markers from rules 7-8 are the one exception — those are expected).\n"
        + ("12. A RECOMMENDATIONS SUMMARY block may be supplied below (pre-computed, real counts — never invent a different number or theme). If present, it MUST form Para 2 on its own (see rule 3) — one brief sentence of social proof citing the exact total-recommendations count and choosing whichever 2-3 listed themes are most relevant to THIS JD's priorities — do not always pick the same fixed themes across different JDs, and do not name individual recommenders or quote a specific recommendation verbatim unless the block explicitly provides a quotable line.\n" if recommendations_summary else "")
        + ("13. Recruiter resume-adjustment guidance (below) may be supplied — it comes from this JD's own Match Scorecard. Apply it ONLY to tone, emphasis, and which existing achievements get foregrounded. It must NEVER be used to introduce a fact, figure, project, or claim that is not already present in the candidate profile data. Never print the guidance verbatim or add a visible heading for it.\n" if resume_adjustment else "")
    )

    # Derive employer display name from employer slug (e.g. MandarinOriental -> Mandarin Oriental)
    import re as _re
    employer_display = _re.sub(r'(?<=[a-z])(?=[A-Z])', ' ', employer)  # CamelCase -> spaced
    role_display     = role_slug.replace('_', ' ')                      # slug -> readable

    coverletter_adjustment_block = (
        f"""
=== RECRUITER RESUME-ADJUSTMENT GUIDANCE (from this JD's own Match Scorecard — apply per rule 13; never invent a fact to satisfy it) ===
{resume_adjustments_text}
"""
        if resume_adjustments_text else ""
    )

    recommendations_summary_block = (
        f"""
=== RECOMMENDATIONS SUMMARY (pre-computed from profile.linkedin_recommendations — apply per rule 12; never invent a different count or theme) ===
{recommendations_summary}
"""
        if recommendations_summary else ""
    )

    COVERLETTER_USER = f"""
{COVERLETTER_TASK_RULES}
{coverletter_adjustment_block}{recommendations_summary_block}
=== LETTER HEADER DETAILS ===
Name     : {name}
Location : {location}
Phone    : {phone}
Email    : {email}
LinkedIn : {linkedin}
{f"Website  : {website}" if website else ""}
(Context only, do not print as header lines — Date: {today}, Hiring Company: {employer_display})

=== INSTRUCTIONS ===
Write a formal executive cover letter for {name} applying for the {role_display} role
at {employer_display}.

Structure:
- Letter header block, exactly these lines with no blank lines between them: Name / Location / Phone | Email / LinkedIn / Website
- No date line, no recipient name line, no company name line
- Salutation: Dear Hiring Manager,
- 5 tight paragraphs as described in the rules above, plus sign-off — cover only the 5 most recent `professional_experience` entries (see rule 3), and keep the whole letter concise enough for a hiring manager to read quickly
- Professional sign-off

Key themes to hit (using only profile facts):
- The highest-priority business and technology themes from this specific JD
- Strong transferable leadership, operational, cybersecurity, transformation, vendor, and stakeholder evidence from the profile where relevant
- Budget governance and measurable savings where they support the business case
- Team building, coaching, and cross-functional partnership where relevant to the JD
- AI/automation innovation as a differentiator when it genuinely supports the role
- Immediate availability, Hong Kong-based, and Cantonese-speaking only if helpful and supported by the profile
- Use semantic alignment: connect adjacent evidence honestly rather than waiting for exact keyword matches
- Bold every distinct quantified achievement and every company/title mention (see rules 7-8)
- Do not state an exact years-of-experience number anywhere (see rule 9)
- No date/recipient-name/company lines before the salutation (see rule 10)
{"- If a RECOMMENDATIONS SUMMARY block was supplied above, it must form its own short Para 2 (see rules 3 and 12) — do not bury it inside an achievement paragraph" if recommendations_summary else ""}
{"- If recruiter resume-adjustment guidance was supplied above, weave it into tone/emphasis only — do not quote it or add a visible heading for it (see rule 13)" if resume_adjustments_text else ""}
"""

    coverletter_text = call_llm(
        SHARED_SYSTEM_PROMPT, COVERLETTER_USER, max_tokens=8000, label="Cover Letter",
        cacheable_prefix=SHARED_CONTEXT_BLOCK,
    )
    coverletter_text = soften_experience_years(coverletter_text)
    coverletter_text = clean_coverletter_header(coverletter_text, employer_display)
    coverletter_text = normalize_blank_lines(coverletter_text)
    print("  ✓  Cover Letter complete")

# ══════════════════════════════════════════════════════════════════════════════
# WRITE OUTPUT FILES
# ══════════════════════════════════════════════════════════════════════════════
# Create all output dirs (txt + docx/pdf companion folders)
for d in [OUT_DIR_SCORECARD, OUT_DIR_RESUME, OUT_DIR_COVER]:
    d.mkdir(parents=True, exist_ok=True)
# Ensure sibling docx/pdf dirs also exist for future use
for subtype in ["ScoreCard", "resume", "CoverLetter"]:
    for fmt in ["docx", "pdf"]:
        (OUT_BASE / subtype / fmt).mkdir(parents=True, exist_ok=True)

generated_at = datetime.now().strftime("%Y-%m-%d %H:%M")
print("\n💾  Writing output files...")

if run_scorecard and scorecard_text:
    # Scorecard header format is unchanged from v1 (out of scope for this pass)
    header = (
        f"JD MATCH SCORECARD — {employer.upper()} | {role_slug.upper()}\n"
        f"Generated : {generated_at}  |  Model: {MODEL}\n"
        f"Source JD : {jd_path.name}\n"
        f"Profile   : {PROFILE_PATH.relative_to(ROOT)}\n"
        f"Candidate : {name}\n"
        + "=" * 72 + "\n\n"
    )
    OUT_SCORECARD.write_text(header + scorecard_text, encoding="utf-8")
    print(f"  ✅  Scorecard   → {OUT_SCORECARD.relative_to(ROOT)}")
    maybe_convert_to_docx(OUT_SCORECARD, "Scorecard")

if run_resume and resume_text:
    # v2: no "TAILORED RESUME —" prefix, no Generated/Profile lines
    resume_header = (
        f"{employer.upper()} | {role_slug.upper()}\n"
        + "=" * 72 + "\n\n"
    )
    OUT_RESUME.write_text(resume_header + resume_text, encoding="utf-8")
    print(f"  ✅  Resume      → {OUT_RESUME.relative_to(ROOT)}")
    maybe_convert_to_docx(OUT_RESUME, "Resume")

if run_coverletter and coverletter_text:
    # v2: title prefix kept, Generated/Profile lines removed
    cl_header = (
        f"COVER LETTER — {employer.upper()} | {role_slug.upper()}\n"
        + "=" * 72 + "\n\n"
    )
    OUT_COVERLETTER.write_text(cl_header + coverletter_text, encoding="utf-8")
    print(f"  ✅  Cover Letter → {OUT_COVERLETTER.relative_to(ROOT)}")
    maybe_convert_to_docx(OUT_COVERLETTER, "Cover Letter")

print(f"\n  Output folder: data_processed/{PROFILE_NAME}/{employer}/")

print(f"\n{'='*60}")
print("  ALL DONE")
print(f"{'='*60}")
print(f"\n  How to run this script:")
print(f"  python scripts/jd_scorecard_resume_v3.py --profile=<Name>")
print(f"  python scripts/jd_scorecard_resume_v3.py --profile=<Name> \"data_raw/<Name>/jd/txt/AnotherJD.txt\"")
print(f"  python scripts/jd_scorecard_resume_v3.py --profile=<Name> --scorecard-only")
print(f"  python scripts/jd_scorecard_resume_v3.py --profile=<Name> --resume-only")
print(f"  python scripts/jd_scorecard_resume_v3.py --profile=<Name> --coverletter-only")
print(f"  python scripts/jd_scorecard_resume_v3.py --profile=<Name> --batch")
print(f"  python scripts/jd_scorecard_resume_v3.py --profile=<Name> --batch --force")
print(f"  python scripts/jd_scorecard_resume_v3.py --profile=<Name> --refresh-blueprint")
print(f"  python scripts/jd_scorecard_resume_v3.py --profile=<Name> --ResumeAdjustment  (apply this JD's scorecard 6a guidance)")
print(f"  python scripts/jd_scorecard_resume_v3.py --profile=<Name> --llm=sonnet    (default, Claude Sonnet 5)")
print(f"  python scripts/jd_scorecard_resume_v3.py --profile=<Name> --llm=deepseek  (DeepSeek R1)")
print(f"  python scripts/jd_scorecard_resume_v3.py --profile=<Name> --llm=gemini    (Gemini Flash Lite - cheaper)")
print(f"  python scripts/jd_scorecard_resume_v3.py --profile=<Name> --llm=custom --model=<id> --provider=openrouter|deepseek")
print(f"  python scripts/jd_scorecard_resume_v3.py --profile=<Name> --llm=sonnet --api-key=<key>  (bring your own key)")
