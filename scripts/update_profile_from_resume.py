#!/usr/bin/env python3
"""
Profile Backfill — Update john_profile.json from a Raw Resume File
====================================================================
Reads one or more raw, unedited resume text files under data_raw/resume/txt/
and adds whatever NEW facts/bullets they contain to the authoritative
src/data/john_profile.json — without touching or duplicating anything
already captured there. Built to fix the root cause of low JD scorecard
match numbers: john_profile.json was only ever populated once (30 Mar 2026)
by a naive regex extractor (backend/consolidation.js, since deprecated) from
a single resume, so most of John's real career history was never captured.

Designed to be reusable, not a one-off — file selection is a CLI argument so
the same logic can later be spawned by the JD Portal backend (the same way
backend/lib/pythonRunner.js already spawns scripts/jd_scorecard_resume_v2.py)
once the "Update from Resume" UI feature is built, rather than only being
runnable from a terminal.

Usage:
  python scripts/update_profile_from_resume.py <file> [<file2> ...] [options]
  python scripts/update_profile_from_resume.py data_raw/resume/txt/JohnHauResume2025-AIA.txt
  python scripts/update_profile_from_resume.py data_raw/resume/txt/JohnHauResume2025-AIA.txt --company="AIA International Ltd"
  python scripts/update_profile_from_resume.py <file> --llm=sonnet|deepseek|gemini
  python scripts/update_profile_from_resume.py <file> --dry-run

--company=<name>   Explicit employer name to match against profile.professional_experience
                    (must match an existing entry's "company" field exactly). If omitted,
                    the script tries to auto-detect it from the file content.
--llm=<name>        sonnet (default) | deepseek | gemini — same presets as jd_scorecard_resume_v2.py.
--dry-run           Run the LLM extraction and print what WOULD be added, but never
                    back up or write john_profile.json.
--create-new-employer   Requires --company="Exact New Employer Name". Allows creating a
                    brand-new profile.professional_experience entry when no existing entry
                    matches — the LLM also extracts title/period/scope from the raw text for
                    the new entry (grounded in the source, same anti-hallucination rule as
                    everything else). Without this flag, an unmatched --company is a hard error.

Safety model (soul.md-aligned):
  - Every real write is preceded by a timestamped backup written to backup/backup-*.json,
    in the exact format backend/backup.js's restoreProfile()/listBackups() already expect
    (so these backups are restorable from the not-yet-built Version History UI too).
  - Only ADDS new highlights/achievements/topics — never edits or removes existing entries.
  - Every fact the LLM outputs must be grounded in the raw file's own text (no invention) —
    same anti-hallucination rule used throughout jd_scorecard_resume_v2.py.
  - Deliberately no per-item approval gate here (unlike the planned interactive "Update
    from Resume" merge flow) — confirmed acceptable by the user for this backfill because
    the source is John's own real historical resumes, not LLM-invented content, and the
    pre-write backup + a future restore/diff view already cover the undo path.
"""

import json
import os
import re
import sys
import time
import difflib
from pathlib import Path
from datetime import datetime, timezone

import requests
from docx import Document

ROOT = Path(__file__).resolve().parent.parent
# Default (no --profile= given): John's own flat profile file, unchanged from
# before Phase 2 — every one of the user's own manual re-runs keeps working
# with zero flag changes. --profile=<Name> (added Phase 2, 11 Aug 2026)
# overrides this to src/data/<Name>/profile.json — see resolve_profile_path().
PROFILE_PATH = ROOT / "src" / "data" / "john_profile.json"


def resolve_profile_path(profile_name):
    """--profile=<Name> -> src/data/<Name>/profile.json, validated the same
    way jd_scorecard_resume_v3.py validates it (filename-safe, stays under
    src/data/) — no default here, caller decides what "no --profile" means."""
    if not re.fullmatch(r"[A-Za-z0-9_-]+", profile_name):
        sys.exit(
            f"ERROR: --profile value '{profile_name}' contains invalid characters — letters, "
            f"digits, hyphen, underscore only (no path separators or dots) to keep profile "
            f"folders isolated."
        )
    profile_dir = ROOT / "src" / "data" / profile_name
    if profile_dir.resolve().parent != (ROOT / "src" / "data").resolve():
        sys.exit(f"ERROR: resolved profile path escapes src/data/ — refusing to proceed: {profile_dir}")
    return profile_dir / "profile.json"


def read_resume_text(file_path):
    """Plain .txt read, or .docx paragraph+table text extraction (added
    Phase 2, 11 Aug 2026, to support real resume uploads like Matina Fung's
    MatinFungCV_Oct2025.docx). Document.tables only returns TOP-LEVEL tables,
    but each cell's .text already recursively includes any nested table's
    text, so walking top-level tables' cells alone captures nested-table
    content too without double-counting."""
    if file_path.suffix.lower() != ".docx":
        return file_path.read_text(encoding="utf-8")
    doc = Document(str(file_path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)
BACKUP_DIR = ROOT / "backup"


# ── Env loading (.env.local -> .env.vps -> .env), matches jd_scorecard_resume_v2.py ──
def load_env(path):
    if not path.exists():
        return False
    for line in path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if line and not line.startswith("#") and "=" in line:
            k, _, v = line.partition("=")
            os.environ[k.strip()] = v.strip()
    return True


for env_file in [ROOT / ".env.local", ROOT / ".env.vps", ROOT / ".env"]:
    load_env(env_file)

LLM_CONFIGS = {
    "sonnet":  ("anthropic/claude-sonnet-5",            "OPENROUTER_API_KEY", "https://openrouter.ai/api/v1/chat/completions"),
    # "deepseek-reasoner" is a legacy alias DeepSeek retired 24 Jul 2026 — same
    # fix already applied in jd_scorecard_resume_v2.py/v3.py on 3 Aug 2026.
    "deepseek": ("deepseek-v4-flash",                   "DEEPSEEK_API_KEY",   "https://api.deepseek.com/chat/completions"),
    "gemini":  ("google/gemini-3.1-flash-lite-preview", "OPENROUTER_API_KEY", "https://openrouter.ai/api/v1/chat/completions"),
}

LLM_MAX_ATTEMPTS = 3
LLM_RETRY_BACKOFF_SEC = 3
LLM_RETRYABLE_STATUS_CODES = {403, 408, 425, 429, 500, 502, 503, 504}


def call_llm(model, endpoint, api_key, system_prompt, user_prompt, max_tokens=6000, label=""):
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://www.askcareer-ai.com",
        "X-Title": "Profile Backfill",
    }
    payload = {
        "model": model,
        "max_tokens": max_tokens,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        "response_format": {"type": "json_object"},
    }

    resp = None
    for attempt in range(1, LLM_MAX_ATTEMPTS + 1):
        try:
            resp = requests.post(endpoint, headers=headers, json=payload, timeout=180)
        except requests.exceptions.RequestException as exc:
            if attempt < LLM_MAX_ATTEMPTS:
                wait = LLM_RETRY_BACKOFF_SEC * attempt
                print(f"  ⚠️  {label or 'unlabeled'} call failed ({exc}) (attempt {attempt}/{LLM_MAX_ATTEMPTS}) — retrying in {wait}s...")
                time.sleep(wait)
                continue
            raise
        if resp.status_code in LLM_RETRYABLE_STATUS_CODES and attempt < LLM_MAX_ATTEMPTS:
            wait = LLM_RETRY_BACKOFF_SEC * attempt
            print(f"  ⚠️  {resp.status_code} from {label or 'unlabeled'} call (attempt {attempt}/{LLM_MAX_ATTEMPTS}) — retrying in {wait}s...")
            time.sleep(wait)
            continue
        resp.raise_for_status()
        break

    data = resp.json()
    choice = data["choices"][0]
    message = choice["message"]
    content = message.get("content")
    finish_reason = choice.get("finish_reason")
    if not content:
        refusal = message.get("refusal")
        raise RuntimeError(
            f"OpenRouter returned empty content for '{label or 'unlabeled'}' call "
            f"(model={model}, finish_reason={finish_reason}, refusal={refusal!r})."
        )
    if finish_reason == "length":
        raise RuntimeError(
            f"'{label or 'unlabeled'}' call hit max_tokens ({max_tokens}) before finishing "
            f"(model={model}). Raise max_tokens and retry."
        )
    return content


def normalize(text):
    return re.sub(r"[^a-z0-9]+", " ", text.lower()).strip()


def is_near_duplicate(candidate, existing_list, threshold=0.72):
    cand_norm = normalize(candidate)
    for existing in existing_list:
        if difflib.SequenceMatcher(None, cand_norm, normalize(existing)).ratio() >= threshold:
            return True
    return False


def find_matching_experience(profile, company_hint, raw_text, allow_create=False):
    experiences = profile.get("professional_experience", [])
    if company_hint:
        for exp in experiences:
            if normalize(exp.get("company", "")) == normalize(company_hint):
                return exp
        if allow_create:
            return None  # signals process_file() to create a brand-new entry
        sys.exit(f"ERROR: --company='{company_hint}' does not exactly match any profile.professional_experience[].company. "
                  f"Known companies: {[e.get('company') for e in experiences]}. "
                  f"If this is a genuinely new employer, re-run with --create-new-employer.")

    # Auto-detect: does any known company name's first token appear in the raw text?
    text_norm = normalize(raw_text[:1500])
    candidates = []
    for exp in experiences:
        company = exp.get("company", "")
        first_token = normalize(company).split(" ")[0]
        if first_token and first_token in text_norm:
            candidates.append(exp)
    if len(candidates) == 1:
        return candidates[0]
    if len(candidates) > 1:
        sys.exit(f"ERROR: Ambiguous employer match ({[c.get('company') for c in candidates]}) — re-run with --company=\"Exact Name\".")
    sys.exit("ERROR: Could not auto-detect the employer for this file — re-run with --company=\"Exact Name\" "
              f"matching one of: {[e.get('company') for e in experiences]}")


SYSTEM_PROMPT_TEMPLATE = """You are an expert resume writer and career-data analyst helping maintain {candidate_name}'s single \
authoritative career profile (a JSON file that powers an automated JD-matching scorecard system).

You are given:
1. Raw, unedited bullet-point notes from a real resume draft for one specific employer (rough, \
sometimes duplicated, informally worded, occasional typos).
2. The content ALREADY captured about this same employer in the authoritative profile.

Your job: find facts, responsibilities, and achievements in the raw notes that are NOT already \
reflected — even loosely or differently worded — in the existing captured content, and turn ONLY \
those new items into clean, professional, resume-quality entries.

Rules:
- The raw notes may mention MULTIPLE employers (a resume often lists several jobs). Only extract \
content that is clearly about the target employer given below. Ignore/discard anything clearly about \
a different employer, even if it appears in the same file.
- Never invent, estimate, round, or embellish any fact or number not explicitly present in the raw \
notes. Every metric you output must be traceable to the raw text.
- Write each new highlight as one SMART-form bullet (Specific, Measurable, Achievable, Realistic, \
Time-bound), third person implied (no "I"), past tense for completed work.
- Wrap the single most important metric or standout phrase in each bullet in **double asterisks**.
- Deduplicate semantically against the existing captured content — do not restate something already \
covered just because the raw notes word it differently. If the raw notes themselves repeat the same \
fact more than once, output it only once.
- Do not return anything from the "already captured" list — output NEW items only.
- Use the exact employer string given for every "company" field.
- Output strict JSON only, no prose, matching exactly the schema given in the user message."""

USER_PROMPT_TEMPLATE = """Target employer for this extraction (use this exact string for every "company" \
field, and IGNORE any content in the raw notes below that is clearly about a different employer): {company}

=== Already captured in the authoritative profile for this employer ===
Existing professional_experience highlights:
{existing_highlights}

Existing major_achievements entries:
{existing_achievements}

=== Raw, unedited resume notes (may cover multiple employers — only mine content about {company}) ===
{raw_text}

=== Output schema (JSON only) ===
{{
  "new_highlights": ["SMART bullet with **bold metric**", "..."],
  "new_achievements": [
    {{"achievement": "Short Title Case Name", "company": "{company}", "<metric_field>": "value from raw text"}}
  ],
  "new_key_topics_for_qa": ["short keyword or phrase", "..."]{new_entry_schema}
}}

Guidance for new_achievements' extra fields beyond achievement/company: use whichever of these keys \
genuinely apply, only when the raw text gives a real value — metric, scale, cost_savings, impact, \
duration, fte_saved, hours_saved — or another short snake_case key if none fit. Match the style of \
the existing major_achievements entries shown above. Only propose an achievement entry for something \
standout/quantified — routine responsibilities belong in new_highlights instead.{new_entry_guidance}"""

NEW_ENTRY_SCHEMA_FRAGMENT = """,
  "new_entry": {"title": "Job title from raw text", "period": "Employment dates from raw text, e.g. 'Feb 1995 - Mar 1997'", "scope": "One-line scope/team-size summary if the raw text gives one, else empty string"}"""

NEW_ENTRY_GUIDANCE_FRAGMENT = """ There is NO existing profile.professional_experience entry for \
{company} yet — this is a brand-new entry. Also fill in "new_entry" with the title/period/scope \
exactly as stated in the raw text (never invent a date or title not present in the source)."""


# ── --create-new-profile: whole-resume -> whole-profile extraction ────────────
# Added Phase 2 (11 Aug 2026) for onboarding a brand-new profile that has no
# existing profile.json to diff against — a fundamentally different job from
# the employer-scoped diff above (find_matching_experience/process_file only
# ever add ONE employer's worth of delta content to something that already
# exists). This extracts an ENTIRE profile in one holistic call instead.
#
# Target shape mirrors shared/profileSchema.js's SECTION_KEYS/SECTION_SCHEMAS
# field-for-field — keep both in sync if that schema ever changes.
CREATE_SYSTEM_PROMPT = """You are an expert resume writer and career-data analyst building a brand-new \
candidate career profile (a JSON file that powers an automated JD-matching scorecard system) from a \
single raw resume document.

Rules:
- Every fact, figure, date, and claim must come ONLY from the raw resume text given below — never \
invent, estimate, round, or embellish anything not explicitly present.
- Extract the candidate's contact/header details (name, title, location, email, phone, LinkedIn) \
exactly as they appear in the resume header. If a field is genuinely not present in the source, \
output an empty string for it — never guess or invent a placeholder value.
- Write each professional_experience highlight as one SMART-form bullet (Specific, Measurable, \
Achievable, Realistic, Time-bound), third person implied (no "I"), past tense for completed work, \
with the single most important metric or standout phrase wrapped in **double asterisks**.
- List professional_experience entries in the same order as the resume (most recent role first).
- Only propose a major_achievements entry for something standout/quantified — routine \
responsibilities belong in a professional_experience highlight instead.
- If the resume doesn't clearly support a given section (e.g. no LinkedIn recommendations quoted \
in the document), output an empty array for it rather than inventing content.
- "ai_projects" means software/AI systems the candidate personally designed or built (e.g. an app, \
a trading platform, an internal tool they engineered) — NOT projects or initiatives where they used \
an existing AI-powered tool/vendor product as part of a non-technical role (e.g. an HR leader using \
an "AI-enabled" recruitment assessment tool is NOT an AI project). Leave this array empty unless the \
resume clearly describes the candidate building/engineering the system themselves — most resumes \
should leave this empty, and that is the expected, normal outcome.
- Output strict JSON only, no prose, matching exactly the schema given in the user message."""

CREATE_USER_PROMPT_TEMPLATE = """=== Raw resume text (may include multiple employers/sections in any order) ===
{raw_text}

=== Output schema (JSON only) — every key is required, use [] / "" / {{}} for anything not grounded in the source ===
{{
  "metadata": {{
    "name": "Full name from the resume header",
    "title": "Current or most recent job title",
    "location": "City/region from the resume header",
    "email": "Email address",
    "phone": "Phone number",
    "linkedin": "LinkedIn URL or handle",
    "availability": "",
    "years_experience": 0
  }},
  "summary": "2-4 sentence professional summary synthesized ONLY from the resume's own content",
  "professional_experience": [
    {{
      "company": "Company name",
      "title": "Job title",
      "period": "Employment dates as stated in the resume, e.g. 'Jan 2020 - Present'",
      "scope": "One-line scope/team-size/remit summary if the resume gives one, else empty string",
      "highlights": ["SMART bullet with **bold metric**", "..."]
    }}
  ],
  "major_achievements": [
    {{"achievement": "Short Title Case Name", "company": "Company name", "<metric_field>": "value from the resume"}}
  ],
  "core_competencies": {{"category_name": ["competency", "..."]}},
  "technical_skills": ["skill or skill-category string", "..."],
  "education_certifications": [
    {{"credential": "Credential/degree name", "issuer": "Institution/issuer", "year": "", "type": "", "details": ""}}
  ],
  "soft_skills": ["soft skill", "..."],
  "languages": [{{"language": "Language name", "proficiency": "Fluent/Professional/etc."}}],
  "languages_spoken": ["Language (Proficiency)", "..."],
  "key_topics_for_qa": ["short keyword or phrase for JD-matching", "..."],
  "ai_projects": [],
  "linkedin_recommendations": []
}}

Guidance for major_achievements' extra fields beyond achievement/company: use whichever of these keys \
genuinely apply, only when the resume gives a real value — metric, scale, cost_savings, impact, \
duration, fte_saved, hours_saved — or another short snake_case key if none fit."""

REQUIRED_NEW_PROFILE_METADATA_FIELDS = ["name", "location", "phone", "email", "linkedin"]


def create_new_profile(file_paths, profile_name, llm_choice, dry_run):
    """--create-new-profile: bootstrap src/data/<profile_name>/profile.json from
    one or more raw resume files (combined into a single extraction call)."""
    model, api_key_env, endpoint = LLM_CONFIGS[llm_choice]
    api_key = os.environ.get(api_key_env, "")
    if not api_key:
        sys.exit(f"ERROR: {api_key_env} not found in .env.local / .env.vps / .env")

    raw_text = "\n\n=== NEXT SOURCE FILE ===\n\n".join(read_resume_text(p) for p in file_paths)

    print(f"\n{'='*60}")
    print("  CREATE NEW PROFILE FROM RESUME")
    print(f"{'='*60}")
    print(f"  Source file(s) : {', '.join(str(p.relative_to(ROOT)) for p in file_paths)}")
    print(f"  Target profile : {PROFILE_PATH.relative_to(ROOT)}")
    print(f"  Model          : {model}")
    print(f"  Mode           : {'DRY RUN (no write)' if dry_run else 'LIVE (will write)'}")
    print(f"{'='*60}\n")

    user_prompt = CREATE_USER_PROMPT_TEMPLATE.format(raw_text=raw_text)

    print("  ↳ Calling LLM to extract the full profile...")
    content = call_llm(
        model, endpoint, api_key, CREATE_SYSTEM_PROMPT, user_prompt,
        max_tokens=32000, label=f"CreateProfile:{profile_name}",
    )
    stripped = content.strip()
    fence_match = re.match(r"^```(?:json)?\s*(.*?)\s*```$", stripped, re.DOTALL)
    if fence_match:
        stripped = fence_match.group(1)
    try:
        new_profile = json.loads(stripped)
    except json.JSONDecodeError as exc:
        sys.exit(f"ERROR: LLM did not return valid JSON: {exc}\n---\n{content}\n---")

    meta = new_profile.get("metadata", {})
    missing = [f for f in REQUIRED_NEW_PROFILE_METADATA_FIELDS if not meta.get(f)]
    if missing:
        sys.exit(
            f"ERROR: extraction did not produce required metadata field(s): {', '.join(missing)} — "
            f"refusing to write a profile with placeholder/missing contact details. Check that the "
            f"source resume actually contains this information, or fill it in manually after a "
            f"--dry-run extraction."
        )

    exp_count = len(new_profile.get("professional_experience", []))
    print(f"\n  Candidate      : {meta.get('name')}")
    print(f"  professional_experience entries : {exp_count}")
    print(f"  major_achievements entries      : {len(new_profile.get('major_achievements', []))}")

    if dry_run:
        print("\n  --dry-run set — no file written.")
        print(json.dumps(new_profile, indent=2, ensure_ascii=False))
        return

    PROFILE_PATH.parent.mkdir(parents=True, exist_ok=True)
    envelope = {
        "timestamp": datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
        "sourceFile": str(PROFILE_PATH.relative_to(ROOT)),
        "profile": new_profile,
    }
    PROFILE_PATH.write_text(json.dumps(envelope, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"\n  ✅ Created {PROFILE_PATH.relative_to(ROOT)} ({exp_count} professional_experience entries)")


def backup_profile():
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    profile_content = PROFILE_PATH.read_text(encoding="utf-8")
    profile_whole = json.loads(profile_content)

    now = datetime.now(timezone.utc)
    ts = now.isoformat().replace("+00:00", "").replace(":", "-").replace(".", "-")
    backup_filename = f"backup-{ts}.json"
    backup_path = BACKUP_DIR / backup_filename

    backup_data = {
        "timestamp": now.isoformat().replace("+00:00", "Z"),
        "sourceFile": str(PROFILE_PATH),
        "fileSize": len(profile_content),
        "profile": profile_whole,  # whole original envelope, matching backend/backup.js's restoreProfile() expectations
    }
    backup_path.write_text(json.dumps(backup_data, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"  Backup written: backup/{backup_filename}")
    return backup_path


def process_file(file_path, company_override, llm_choice, dry_run, allow_create=False):
    model, api_key_env, endpoint = LLM_CONFIGS[llm_choice]
    api_key = os.environ.get(api_key_env, "")
    if not api_key:
        sys.exit(f"ERROR: {api_key_env} not found in .env.local / .env.vps / .env")

    raw_text = read_resume_text(file_path)
    profile_envelope = json.loads(PROFILE_PATH.read_text(encoding="utf-8"))
    profile = profile_envelope["profile"]
    candidate_name = profile.get("metadata", {}).get("name", "the candidate")

    matched_exp = find_matching_experience(profile, company_override, raw_text, allow_create=allow_create)
    is_new_entry = matched_exp is None
    company = company_override if is_new_entry else matched_exp["company"]
    existing_highlights = [] if is_new_entry else matched_exp.get("highlights", [])
    existing_achievements = [] if is_new_entry else [a for a in profile.get("major_achievements", []) if company.split(" ")[0].lower() in a.get("company", "").lower()]

    print(f"\n{'='*60}")
    print("  PROFILE BACKFILL FROM RESUME")
    print(f"{'='*60}")
    print(f"  Source file : {file_path.relative_to(ROOT)}")
    print(f"  Employer    : {company}{' (NEW ENTRY)' if is_new_entry else ''}")
    print(f"  Model       : {model}")
    print(f"  Existing highlights   : {len(existing_highlights)}")
    print(f"  Existing achievements : {len(existing_achievements)}")
    print(f"  Mode        : {'DRY RUN (no write)' if dry_run else 'LIVE (will back up + write)'}")
    print(f"{'='*60}\n")

    user_prompt = USER_PROMPT_TEMPLATE.format(
        company=company,
        existing_highlights=json.dumps(existing_highlights, indent=2, ensure_ascii=False),
        existing_achievements=json.dumps(existing_achievements, indent=2, ensure_ascii=False),
        raw_text=raw_text,
        new_entry_schema=NEW_ENTRY_SCHEMA_FRAGMENT if is_new_entry else "",
        new_entry_guidance=NEW_ENTRY_GUIDANCE_FRAGMENT.format(company=company) if is_new_entry else "",
    )

    print("  ↳ Calling LLM to extract new content...")
    system_prompt = SYSTEM_PROMPT_TEMPLATE.format(candidate_name=candidate_name)
    content = call_llm(model, endpoint, api_key, system_prompt, user_prompt, max_tokens=12000, label=f"Backfill:{company}")
    # Despite response_format=json_object, models occasionally still wrap the
    # payload in a ```json fence — strip it before parsing.
    stripped = content.strip()
    fence_match = re.match(r"^```(?:json)?\s*(.*?)\s*```$", stripped, re.DOTALL)
    if fence_match:
        stripped = fence_match.group(1)
    try:
        result = json.loads(stripped)
    except json.JSONDecodeError as exc:
        sys.exit(f"ERROR: LLM did not return valid JSON: {exc}\n---\n{content}\n---")

    new_highlights = result.get("new_highlights", [])
    new_achievements = result.get("new_achievements", [])
    new_topics = result.get("new_key_topics_for_qa", [])

    # Safety-net dedupe against existing content (the model is instructed to do this,
    # this is a second pass in case it misses a near-duplicate).
    existing_achievement_names = [a.get("achievement", "") for a in existing_achievements]
    existing_topics = profile.get("key_topics_for_qa", [])

    kept_highlights = [h for h in new_highlights if not is_near_duplicate(h, existing_highlights)]
    kept_achievements = [a for a in new_achievements if not is_near_duplicate(a.get("achievement", ""), existing_achievement_names)]
    kept_topics = [t for t in new_topics if not is_near_duplicate(t, existing_topics)]

    dropped = (len(new_highlights) - len(kept_highlights)) + (len(new_achievements) - len(kept_achievements)) + (len(new_topics) - len(kept_topics))
    if dropped:
        print(f"  ↳ Dropped {dropped} near-duplicate item(s) caught by the safety-net dedupe pass")

    print(f"\n  New highlights to add    : {len(kept_highlights)}")
    for h in kept_highlights:
        print(f"    + {h}")
    print(f"\n  New achievements to add  : {len(kept_achievements)}")
    for a in kept_achievements:
        extras = {k: v for k, v in a.items() if k not in ("achievement", "company")}
        print(f"    + {a.get('achievement')} ({extras})")
    print(f"\n  New key_topics_for_qa to add : {len(kept_topics)}")
    for t in kept_topics:
        print(f"    + {t}")

    new_entry_info = result.get("new_entry") if is_new_entry else None
    if is_new_entry:
        if not new_entry_info or not new_entry_info.get("title") or not new_entry_info.get("period"):
            sys.exit(f"ERROR: --create-new-employer set but the LLM did not return a usable "
                      f"new_entry.title/period for '{company}'. Raw new_entry: {new_entry_info}")
        print(f"\n  New professional_experience entry: {company} — {new_entry_info['title']} ({new_entry_info['period']})")

    if not (kept_highlights or kept_achievements or kept_topics or is_new_entry):
        print("\n  Nothing new to add — profile already covers this file's content.")
        return

    if dry_run:
        print("\n  --dry-run set — no backup taken, no file written.")
        return

    backup_profile()

    if is_new_entry:
        matched_exp = {
            "company": company,
            "title": new_entry_info["title"],
            "period": new_entry_info["period"],
            "scope": new_entry_info.get("scope", ""),
            "highlights": kept_highlights,
        }
        profile.setdefault("professional_experience", []).append(matched_exp)
    else:
        matched_exp["highlights"] = existing_highlights + kept_highlights
    profile.setdefault("major_achievements", []).extend(kept_achievements)
    profile.setdefault("key_topics_for_qa", []).extend(kept_topics)

    metadata = profile.setdefault("metadata", {})
    sources = metadata.get("resume_sources", [])
    if file_path.name not in sources:
        sources.append(file_path.name)
    metadata["resume_sources"] = sources
    metadata["last_backfilled"] = datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")

    profile_envelope["profile"] = profile
    profile_envelope["timestamp"] = datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")

    PROFILE_PATH.write_text(json.dumps(profile_envelope, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"\n  ✅ {PROFILE_PATH.relative_to(ROOT)} updated ({file_path.name} → {company})")


def main():
    global PROFILE_PATH

    args = sys.argv[1:]
    flags = [a for a in args if a.startswith("--")]
    files = [Path(a) for a in args if not a.startswith("--")]

    if not files:
        sys.exit("ERROR: Provide at least one resume file path, e.g.\n"
                  "  python scripts/update_profile_from_resume.py data_raw/resume/txt/JohnHauResume2025-AIA.txt")

    llm_flag = next((f for f in flags if f.startswith("--llm=")), "--llm=sonnet")
    llm_choice = llm_flag.split("=", 1)[1].lower()
    if llm_choice not in LLM_CONFIGS:
        sys.exit(f"ERROR: Unknown --llm value '{llm_choice}'. Valid: {list(LLM_CONFIGS)}")

    company_flag = next((f for f in flags if f.startswith("--company=")), None)
    company_override = company_flag.split("=", 1)[1] if company_flag else None

    dry_run = "--dry-run" in flags
    allow_create = "--create-new-employer" in flags
    if allow_create and not company_override:
        sys.exit("ERROR: --create-new-employer requires --company=\"Exact New Employer Name\"")

    # --profile=<Name> (added Phase 2, 11 Aug 2026): optional, defaults to
    # John's own flat profile file (unchanged behavior) when omitted.
    profile_flag = next((f for f in flags if f.startswith("--profile=")), None)
    profile_name = profile_flag.split("=", 1)[1] if profile_flag else None
    if profile_name:
        PROFILE_PATH = resolve_profile_path(profile_name)

    create_new_profile_flag = "--create-new-profile" in flags
    if create_new_profile_flag:
        if not profile_name:
            sys.exit("ERROR: --create-new-profile requires --profile=<Name>")
        if company_override or allow_create:
            sys.exit(
                "ERROR: --create-new-profile cannot be combined with --company=/--create-new-employer "
                "(those are for updating an EXISTING profile via the employer-diff mode)."
            )
        if PROFILE_PATH.exists():
            sys.exit(
                f"ERROR: --create-new-profile requires the target profile NOT already exist, but "
                f"{PROFILE_PATH.relative_to(ROOT)} already exists. Omit --create-new-profile to run "
                f"an update instead."
            )
        file_paths = []
        for file_arg in files:
            file_path = file_arg if file_arg.is_absolute() else ROOT / file_arg
            if not file_path.exists():
                sys.exit(f"ERROR: File not found: {file_path}")
            file_paths.append(file_path)
        create_new_profile(file_paths, profile_name, llm_choice, dry_run)
        return

    for file_arg in files:
        file_path = file_arg if file_arg.is_absolute() else ROOT / file_arg
        if not file_path.exists():
            sys.exit(f"ERROR: File not found: {file_path}")
        process_file(file_path, company_override, llm_choice, dry_run, allow_create=allow_create)


if __name__ == "__main__":
    main()
