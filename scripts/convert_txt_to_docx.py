#!/usr/bin/env python3
"""
Convert plain-text resume / cover-letter files into .docx format.

Usage:
  python scripts/convert_txt_to_docx.py path/to/resume.txt
  python scripts/convert_txt_to_docx.py path/to/resume.txt path/to/coverletter.txt
  python scripts/convert_txt_to_docx.py --output-dir data_processed/docx path/to/file.txt
"""

from __future__ import annotations

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Inches, Pt
except ImportError:  # pragma: no cover
    sys.exit(
        "ERROR: python-docx is not installed. Run `pip install python-docx` or use the project environment installer first."
    )


def parse_args(argv: list[str]) -> tuple[Path | None, list[Path]]:
    output_dir = None
    paths: list[Path] = []
    i = 0
    while i < len(argv):
        arg = argv[i]
        if arg == "--output-dir":
            i += 1
            if i >= len(argv):
                sys.exit("ERROR: --output-dir requires a directory path")
            output_dir = Path(argv[i])
        else:
            paths.append(Path(arg))
        i += 1

    if not paths:
        sys.exit(
            "Usage: python scripts/convert_txt_to_docx.py [--output-dir DIR] <file1.txt> [file2.txt ...]"
        )
    return output_dir, paths


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def add_text_block(doc: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        doc.add_paragraph()
        return

    if set(stripped) <= {"=", "_", "-"} and len(stripped) >= 8:
        return

    if stripped.startswith(("•", "* ", "- ")):
        bullet_text = stripped.lstrip("•*- ").strip()
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.left_indent = Inches(0.25)
        para.add_run(bullet_text)
        return

    if stripped.isupper() and len(stripped) <= 80:
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = para.add_run(stripped)
        run.bold = True
        run.font.size = Pt(12)
        return

    if " — " in stripped and not stripped.endswith(":"):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(3)
        left, right = stripped.split(" — ", 1)
        run1 = para.add_run(left + " — ")
        run1.bold = True
        para.add_run(right)
        return

    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    para.add_run(stripped)


def convert_file(txt_path: Path, output_dir: Path | None = None) -> Path:
    if not txt_path.exists():
        raise FileNotFoundError(f"Input file not found: {txt_path}")

    doc = Document()
    style_document(doc)

    for line in txt_path.read_text(encoding="utf-8").splitlines():
        add_text_block(doc, line)

    target_dir = output_dir or txt_path.parent.parent / "docx"
    target_dir.mkdir(parents=True, exist_ok=True)
    out_path = target_dir / f"{txt_path.stem}.docx"
    doc.save(out_path)
    return out_path


def main() -> None:
    output_dir, input_paths = parse_args(sys.argv[1:])
    print("Converting text files to DOCX...")
    for raw_path in input_paths:
        path = raw_path if raw_path.is_absolute() else Path.cwd() / raw_path
        out_path = convert_file(path, output_dir)
        print(f"✅ {path.name} -> {out_path}")


if __name__ == "__main__":
    main()
