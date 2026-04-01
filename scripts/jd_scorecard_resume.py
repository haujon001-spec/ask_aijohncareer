#!/usr/bin/env python3
"""
All-in-One JD Application Generator
=====================================
Generates three documents from a Job Description + john_profile.json:
  1. JD Match Scorecard  — detailed scoring with gaps, strengths, verdict
  2. Tailored Resume     — ATS-optimised, formatted like MorganStanley template
                          (first 3 companies: 10 bullets each;
                           4th and 5th companies: 8 bullets each)
  3. Cover Letter        — executive-level, role-specific, personalised

Usage:
  python scripts/jd_scorecard_resume.py
      -> uses the most recently updated JD under `data_raw/jd/txt/`
         (or `JD_DEFAULT_PATH` if set in the environment)

  python scripts/jd_scorecard_resume.py "path/to/AnotherJD.txt"
      -> uses specified JD file

  python scripts/jd_scorecard_resume.py --scorecard-only
  python scripts/jd_scorecard_resume.py --resume-only
  python scripts/jd_scorecard_resume.py --coverletter-only
  python scripts/jd_scorecard_resume.py --refresh-blueprint
  python scripts/jd_scorecard_resume.py --batch
  python scripts/jd_scorecard_resume.py --batch --force

LLM selection (add flag, default is sonnet):
  python scripts/jd_scorecard_resume.py --llm=sonnet        (Claude Sonnet 4.6 - smart, OpenRouter)
  python scripts/jd_scorecard_resume.py --llm=deepseek      (DeepSeek R1 - reasoning, DeepSeek API)
  python scripts/jd_scorecard_resume.py --llm=gemini        (Gemini 3.1 Flash Lite - fast/cheap, OpenRouter)

Outputs go to: data_processed/<Employer>/
  ScoreCard/txt/JD_SCORECARD_<Employer>_<Role>_<DATE>.txt
  ScoreCard/docx/JD_SCORECARD_<Employer>_<Role>_<DATE>.docx
  resume/txt/JohnHauResume<YEAR>_<Employer>_<Role>_<DATE>.txt
  resume/docx/JohnHauResume<YEAR>_<Employer>_<Role>_<DATE>.docx
  CoverLetter/txt/JohnHauCoverLetter_<Employer>_<Role>_<DATE>.txt
  CoverLetter/docx/JohnHauCoverLetter_<Employer>_<Role>_<DATE>.docx

Source: src/data/john_profile.json  (single source of truth — no hallucination)
"""

import json
import os
import re
import subprocess
import sys
import requests
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Inches, Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ── Project root (one level up from scripts/) ─────────────────────────────────
ROOT = Path(__file__).resolve().parent.parent

# ── Fixed paths ────────────────────────────────────────────────────────────────
PROFILE_PATH      = ROOT / "src/data/john_profile.json"
DEFAULT_JD_DIR    = ROOT / "data_raw/jd/txt"
JD_BLUEPRINT_DIR  = ROOT / "src/data/jd"
TEMPLATE_CANDIDATES = [
    ROOT / "data_raw/resume/txt/JohnHauResume2026_MorganStanley.md",
    ROOT / "data_raw/resume/txt/JohnHauResume2026_MorganStanley.txt",
]

DATE_STAMP    = datetime.now().strftime("%d%b%Y").upper()   # e.g. 31MAR2026
RESUME_YEAR   = datetime.now().strftime("%Y")               # e.g. 2026

def resolve_resume_template():
    """Prefer the new Markdown master resume template, with txt as fallback."""
    for candidate in TEMPLATE_CANDIDATES:
        if candidate.exists():
            return candidate
    sys.exit(
        "ERROR: Resume template not found. Expected one of: "
        + ", ".join(str(p) for p in TEMPLATE_CANDIDATES)
    )

TEMPLATE_PATH = resolve_resume_template()

def resolve_default_jd():
    """Pick the JD file explicitly provided via env, otherwise the most recently updated JD in data_raw/jd/txt/."""
    env_jd = os.environ.get("JD_DEFAULT_PATH", "").strip()
    if env_jd:
        env_path = Path(env_jd)
        if not env_path.is_absolute():
            env_path = ROOT / env_path
        if env_path.exists():
            return env_path

    jd_files = sorted(DEFAULT_JD_DIR.glob("JD_*.txt"), key=lambda p: p.stat().st_mtime, reverse=True)
    if jd_files:
        return jd_files[0]

    fallback_files = sorted(DEFAULT_JD_DIR.glob("*.txt"), key=lambda p: p.stat().st_mtime, reverse=True)
    if fallback_files:
        return fallback_files[0]

    sys.exit(f"ERROR: No JD text files found under {DEFAULT_JD_DIR}")

def load_jd_blueprint(jd_stem):
    """Optionally load a structured JD blueprint from src/data/jd/<JD_STEM>.json."""
    blueprint_path = JD_BLUEPRINT_DIR / f"{jd_stem}.json"
    if not blueprint_path.exists():
        return None, None

    try:
        blueprint = json.loads(blueprint_path.read_text(encoding="utf-8"))
        return blueprint_path, blueprint
    except Exception as exc:
        print(f"⚠️  Warning: Could not parse JD blueprint {blueprint_path.name}: {exc}")
        return blueprint_path, None


def derive_jd_metadata(path):
    jd_stem_local = path.stem
    jd_parts_local = jd_stem_local.replace("JD_", "").split("_")
    employer_local = jd_parts_local[0] if jd_parts_local else "Employer"
    role_slug_local = "_".join(jd_parts_local[1:]) if len(jd_parts_local) > 1 else "Role"
    role_tag_local = f"_{role_slug_local}" if role_slug_local and role_slug_local != "Role" else ""
    return jd_stem_local, employer_local, role_slug_local, role_tag_local


def build_output_targets(path):
    jd_stem_local, employer_local, role_slug_local, role_tag_local = derive_jd_metadata(path)
    out_base = ROOT / "data_processed" / employer_local
    return {
        "blueprint": JD_BLUEPRINT_DIR / f"{jd_stem_local}.json",
        "scorecard_dir": out_base / "ScoreCard" / "txt",
        "resume_dir": out_base / "resume" / "txt",
        "cover_dir": out_base / "CoverLetter" / "txt",
        "scorecard_pattern": f"JD_SCORECARD_{employer_local}{role_tag_local}_*.txt",
        "resume_pattern": f"JohnHauResume*_{employer_local}{role_tag_local}_*.txt",
        "cover_pattern": f"JohnHauCoverLetter_{employer_local}{role_tag_local}_*.txt",
    }


def requested_outputs_exist(path, include_scorecard=True, include_resume=True, include_coverletter=True):
    targets = build_output_targets(path)
    checks = []
    if include_scorecard:
        checks.append(any(targets["scorecard_dir"].glob(targets["scorecard_pattern"])))
    if include_resume:
        checks.append(any(targets["resume_dir"].glob(targets["resume_pattern"])))
    if include_coverletter:
        checks.append(any(targets["cover_dir"].glob(targets["cover_pattern"])))
    return bool(checks) and all(checks)

# ── Parse CLI arguments ────────────────────────────────────────────────────────
args = [a for a in sys.argv[1:] if not a.startswith("-")]
flags = [a for a in sys.argv[1:] if a.startswith("-")]

jd_path = Path(args[0]) if args else resolve_default_jd()
if not jd_path.is_absolute():
    jd_path = ROOT / jd_path
if not jd_path.exists():
    sys.exit(f"ERROR: JD file not found: {jd_path}")

run_scorecard     = "--resume-only" not in flags and "--coverletter-only" not in flags
run_resume        = "--scorecard-only" not in flags and "--coverletter-only" not in flags
run_coverletter   = "--scorecard-only" not in flags and "--resume-only" not in flags
refresh_blueprint = "--refresh-blueprint" in flags
batch_mode        = "--batch" in flags
force_run         = "--force" in flags
generate_docx     = "--no-docx" not in flags

# ── LLM selection via --llm=<name> flag ───────────────────────────────────────
llm_flag = next((f for f in flags if f.startswith("--llm=")), "--llm=sonnet")
llm_choice = llm_flag.split("=", 1)[1].lower()   # sonnet | deepseek | gemini

# ── Derive employer slug from JD filename for output naming ───────────────────
jd_stem    = jd_path.stem                               # e.g. JD_MandarinOriental_ClusterDirectorOfIT
jd_parts   = jd_stem.replace("JD_", "").split("_")
employer   = jd_parts[0] if jd_parts else "Employer"   # e.g. MandarinOriental
role_slug  = "_".join(jd_parts[1:]) if len(jd_parts) > 1 else "Role"
jd_blueprint_path, jd_blueprint = load_jd_blueprint(jd_stem)

# ── Output directories: data_processed/<Employer>/<Type>/txt/ ─────────────────
OUT_BASE          = ROOT / "data_processed" / employer
OUT_DIR_SCORECARD = OUT_BASE / "ScoreCard" / "txt"
OUT_DIR_RESUME    = OUT_BASE / "resume"    / "txt"
OUT_DIR_COVER     = OUT_BASE / "CoverLetter" / "txt"
_role_tag         = f"_{role_slug}" if role_slug and role_slug != "Role" else ""
OUT_SCORECARD     = OUT_DIR_SCORECARD / f"JD_SCORECARD_{employer}{_role_tag}_{DATE_STAMP}.txt"
OUT_RESUME        = OUT_DIR_RESUME    / f"JohnHauResume{RESUME_YEAR}_{employer}{_role_tag}_{DATE_STAMP}.txt"
OUT_COVERLETTER   = OUT_DIR_COVER     / f"JohnHauCoverLetter_{employer}{_role_tag}_{DATE_STAMP}.txt"

# ── Load env (.env.local → .env.vps → .env) — file always wins over stale shell ──
def load_env(path):
    if not path.exists():
        return False
    for line in path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if line and not line.startswith("#") and "=" in line:
            k, _, v = line.partition("=")
            os.environ[k.strip()] = v.strip()
    return True

for env_file in [ROOT / ".env.local", ROOT / ".env.vps", ROOT / ".env"]:
    if load_env(env_file):
        break

# ── Resolve API key and model based on LLM selection ─────────────────────────
LLM_CONFIGS = {
    # name          model_id                                     api_key_env            base_url
    "sonnet":  ("anthropic/claude-sonnet-4.6",              "OPENROUTER_API_KEY",  "https://openrouter.ai/api/v1/chat/completions"),
    "deepseek":("deepseek-reasoner",                        "DEEPSEEK_API_KEY",    "https://api.deepseek.com/chat/completions"),
    "gemini":  ("google/gemini-3.1-flash-lite-preview",     "OPENROUTER_API_KEY",  "https://openrouter.ai/api/v1/chat/completions"),
}

if llm_choice not in LLM_CONFIGS:
    sys.exit(f"ERROR: Unknown --llm value '{llm_choice}'. Valid: sonnet, deepseek, gemini")

MODEL, api_key_env, LLM_ENDPOINT = LLM_CONFIGS[llm_choice]

API_KEY = os.environ.get(api_key_env, "")
if not API_KEY:
    sys.exit(f"ERROR: {api_key_env} not found in .env.local / .env.vps / .env")

HEADERS = {
    "Authorization": f"Bearer {API_KEY}",
    "Content-Type": "application/json",
    "HTTP-Referer": "https://www.askcareer-ai.com",
    "X-Title": "JD Application Generator - John Hau",
}

if batch_mode:
    jd_files = sorted(DEFAULT_JD_DIR.glob("JD_*.txt"))
    if not jd_files:
        sys.exit(f"ERROR: No JD text files found under {DEFAULT_JD_DIR}")

    print(f"\n{'='*60}")
    print("  JD APPLICATION GENERATOR — BATCH MODE")
    print(f"{'='*60}")
    print(f"  JD Folder   : {DEFAULT_JD_DIR.relative_to(ROOT)}")
    print(f"  File Count  : {len(jd_files)}")
    print(f"  LLM         : {llm_choice} ({MODEL})")
    print(f"  Skip Existing: {'Yes' if not force_run else 'No (force enabled)'}")
    print(f"{'='*60}\n")

    processed = 0
    skipped = 0
    failed = 0
    forwarded_flags = [f for f in flags if f not in {"--batch", "--force"}]

    for batch_jd in jd_files:
        print(f"\n▶ Processing {batch_jd.name}")
        if not force_run and requested_outputs_exist(batch_jd, run_scorecard, run_resume, run_coverletter):
            print("  ↷ Skipped — requested outputs already exist")
            skipped += 1
            continue

        cmd = [sys.executable, str(Path(__file__).resolve()), str(batch_jd)] + forwarded_flags
        result = subprocess.run(cmd, cwd=str(ROOT))
        if result.returncode == 0:
            processed += 1
        else:
            failed += 1
            print(f"  ✗ Failed — exit code {result.returncode}")

    print(f"\n{'='*60}")
    print("  BATCH SUMMARY")
    print(f"{'='*60}")
    print(f"  Processed : {processed}")
    print(f"  Skipped   : {skipped}")
    print(f"  Failed    : {failed}")
    sys.exit(0 if failed == 0 else 1)

# ── Read source files ──────────────────────────────────────────────────────────
print(f"\n{'='*60}")
print("  JD APPLICATION GENERATOR — John Hau")
print(f"{'='*60}")
print(f"  JD File     : {jd_path.name}")
print(f"  JD Blueprint: {jd_blueprint_path.relative_to(ROOT) if jd_blueprint_path and jd_blueprint else 'dynamic from JD text'}")
print(f"  Template    : {TEMPLATE_PATH.relative_to(ROOT)}")
print(f"  Profile     : src/data/john_profile.json")
print(f"  Model       : {MODEL}")
print(f"  API Key     : {API_KEY[:25]}...")
print(f"  Date        : {DATE_STAMP}")
print(f"  LLM         : {llm_choice} ({MODEL})")
print(f"  Generating  : {'Scorecard ' if run_scorecard else ''}{'Resume ' if run_resume else ''}{'CoverLetter' if run_coverletter else ''}")
print(f"{'='*60}\n")

print("📂  Reading source files...")
jd_text      = jd_path.read_text(encoding="utf-8")
profile_raw  = json.loads(PROFILE_PATH.read_text(encoding="utf-8"))
template_txt = TEMPLATE_PATH.read_text(encoding="utf-8")

# Unwrap nested "profile" key if present
profile = profile_raw.get("profile", profile_raw)

# Build a compact profile context — metadata + summary + experience + achievements
# Stays within ~12K chars to avoid token overrun while keeping all key facts
def build_profile_context(profile):
    sections = {}
    for key in ["metadata", "summary", "professional_experience", "major_achievements",
                "ai_projects", "core_competencies", "technical_skills",
                "education_certifications", "languages_spoken"]:
        if key in profile:
            sections[key] = profile[key]
    return json.dumps(sections, indent=2, ensure_ascii=False)[:14000]

profile_context = build_profile_context(profile)

def build_blueprint_context(blueprint):
    return (
        json.dumps(blueprint, indent=2, ensure_ascii=False)[:6000]
        if blueprint
        else "No JD JSON blueprint found. Infer the priority criteria, must-have skills, soft skills, and weighting dynamically from the raw JD text."
    )

jd_blueprint_context = build_blueprint_context(jd_blueprint)

# ── LLM call helper ────────────────────────────────────────────────────────────
def call_llm(system_prompt, user_prompt, max_tokens=6000, label=""):
    if label:
        print(f"  ↳ Calling OpenRouter ({MODEL}) — {label}")
    payload = {
        "model": MODEL,
        "max_tokens": max_tokens,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_prompt},
        ],
    }
    resp = requests.post(
        LLM_ENDPOINT,
        headers=HEADERS,
        json=payload,
        timeout=180,
    )
    resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"]

def extract_json_object(raw_text):
    """Extract a JSON object from an LLM response, allowing for code fences."""
    cleaned = raw_text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
        cleaned = re.sub(r"\s*```$", "", cleaned)

    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("LLM response did not contain a valid JSON object")

    return json.loads(cleaned[start:end + 1])


def style_docx_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def add_docx_text_block(doc, line):
    stripped = line.strip()
    if not stripped:
        doc.add_paragraph()
        return

    if set(stripped) <= {"=", "_", "-"} and len(stripped) >= 8:
        return

    if stripped.startswith(("•", "* ", "- ")):
        bullet_text = stripped.lstrip("•*- ").strip()
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.left_indent = Inches(0.25)
        para.add_run(bullet_text)
        return

    if stripped.isupper() and len(stripped) <= 80:
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = para.add_run(stripped)
        run.bold = True
        run.font.size = Pt(12)
        return

    if " — " in stripped and not stripped.endswith(":"):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(3)
        left, right = stripped.split(" — ", 1)
        run1 = para.add_run(left + " — ")
        run1.bold = True
        para.add_run(right)
        return

    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    para.add_run(stripped)


def convert_text_file_to_docx(txt_path):
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed")

    doc = Document()
    style_docx_document(doc)
    for line in txt_path.read_text(encoding="utf-8").splitlines():
        add_docx_text_block(doc, line)

    target_dir = txt_path.parent.parent / "docx"
    target_dir.mkdir(parents=True, exist_ok=True)
    out_path = target_dir / f"{txt_path.stem}.docx"
    doc.save(out_path)
    return out_path


def maybe_convert_to_docx(txt_path, label):
    if not generate_docx:
        return None
    if not DOCX_AVAILABLE:
        print(f"  ⚠️  {label} DOCX skipped — python-docx not installed")
        return None
    try:
        out_path = convert_text_file_to_docx(txt_path)
        print(f"  ✅  {label} DOCX → {out_path.relative_to(ROOT)}")
        return out_path
    except Exception as exc:
        print(f"  ⚠️  {label} DOCX conversion failed: {exc}")
        return None


def generate_jd_blueprint(jd_text, jd_path, employer, role_slug):
    """Auto-create a structured JD blueprint for any new JD text file."""
    blueprint_path = JD_BLUEPRINT_DIR / f"{jd_stem}.json"
    JD_BLUEPRINT_DIR.mkdir(parents=True, exist_ok=True)

    BLUEPRINT_SYS = (
        "You are an expert job architecture analyst and recruiter. "
        "Read the JD and return VALID JSON ONLY for downstream automation. "
        "Do not use markdown code fences. Do not include commentary. "
        "Make the criteria practical, recruiter-style, and specific to this JD."
    )

    BLUEPRINT_USER = f"""
=== JOB DESCRIPTION ===
{jd_text}

=== INSTRUCTIONS ===
Create a JSON blueprint for this JD using the exact top-level keys below:
{{
  "jd_file": "...",
  "company": "...",
  "company_slug": "...",
  "job_title": "...",
  "role_slug": "...",
  "industry": "...",
  "role_summary": "...",
  "must_have_requirements": ["..."],
  "nice_to_have_requirements": ["..."],
  "soft_skills_priority": ["..."],
  "semantic_match_guidance": {{
    "category_name": ["related resume/profile evidence terms"]
  }},
  "scoring_criteria": [
    {{
      "name": "...",
      "weight": 0,
      "why_it_matters": "...",
      "evidence_type_to_look_for": ["..."],
      "gap_expected": "optional"
    }}
  ],
  "resume_tailoring_guidance": {{
    "mandatory_sections": ["PROFESSIONAL SUMMARY", "AI & AUTOMATION HIGHLIGHTS", "CORE COMPETENCIES", "SOFT SKILLS"],
    "summary_focus": ["..."],
    "do_not_overclaim": ["..."]
  }},
  "cover_letter_guidance": {{
    "tone": "...",
    "positioning": "...",
    "bridge_gaps_honestly": ["..."]
  }}
}}

Rules:
- Base the blueprint on the JD text only
- Include 8-10 scoring criteria
- Make all scoring weights numeric and total exactly 100
- Keep wording concise and usable for resume / scorecard / cover letter generation
- Include semantic matching hints so adjacent evidence can be matched truthfully without hallucination
"""

    raw = call_llm(BLUEPRINT_SYS, BLUEPRINT_USER, max_tokens=2600, label="JD Blueprint")
    blueprint = extract_json_object(raw)

    blueprint.setdefault("jd_file", str(jd_path.relative_to(ROOT)).replace("\\", "/"))
    blueprint.setdefault("company", employer)
    blueprint.setdefault("company_slug", employer)
    blueprint.setdefault("job_title", role_slug.replace("_", " "))
    blueprint.setdefault("role_slug", role_slug)
    blueprint.setdefault("resume_tailoring_guidance", {})
    blueprint["resume_tailoring_guidance"].setdefault(
        "mandatory_sections",
        ["PROFESSIONAL SUMMARY", "AI & AUTOMATION HIGHLIGHTS", "CORE COMPETENCIES", "SOFT SKILLS"],
    )

    blueprint_path.write_text(json.dumps(blueprint, indent=2, ensure_ascii=False), encoding="utf-8")
    return blueprint_path, blueprint

if refresh_blueprint or not jd_blueprint:
    action = "Refreshing" if refresh_blueprint and jd_blueprint_path else "Generating"
    print(f"🧭  [0/3] {action} JD blueprint...")
    try:
        jd_blueprint_path, jd_blueprint = generate_jd_blueprint(jd_text, jd_path, employer, role_slug)
        jd_blueprint_context = build_blueprint_context(jd_blueprint)
        print(f"  ✓  JD Blueprint ready -> {jd_blueprint_path.relative_to(ROOT)}")
    except Exception as exc:
        print(f"  ⚠️  JD blueprint generation skipped: {exc}")
        jd_blueprint_context = build_blueprint_context(jd_blueprint)

# ══════════════════════════════════════════════════════════════════════════════
# TASK 1 — JD SCORECARD
# ══════════════════════════════════════════════════════════════════════════════
scorecard_text = ""
if run_scorecard:
    print("🔍  [1/3] Generating JD Scorecard...")

    SCORECARD_SYS = (
        "You are an expert executive recruiter and talent analyst. "
        "Produce a detailed, honest, structured scorecard comparing the candidate to the JD. "
        "Be evidence-based and professional. Never invent facts not present in the profile data. "
        "Use plain text with clear section headers. Derive the scoring criteria dynamically from the specific JD, "
        "focus only on materially relevant requirements, and compute a weighted overall score. "
        "When evidence is adjacent or transferable rather than exact, say so explicitly and do not present it as direct hands-on experience."
    )

    SCORECARD_USER = f"""
=== JOB DESCRIPTION ===
{jd_text}

=== OPTIONAL JD BLUEPRINT (use this if present, otherwise infer from JD text) ===
{jd_blueprint_context}

=== CANDIDATE PROFILE (source of truth — do not invent facts outside this data) ===
{profile_context}

=== INSTRUCTIONS ===
Produce a comprehensive JD Match Scorecard with ALL of these sections:

1. ROLE & COMPANY SUMMARY
   Brief overview of the role and what the employer is seeking (3-4 sentences).

2. OVERALL MATCH SCORE
   Format:  XX/100 — [STRONG MATCH / GOOD MATCH / PARTIAL MATCH / WEAK MATCH]
   Followed by one honest paragraph verdict.
   IMPORTANT: compute this as a weighted result from the JD-specific criteria below.
   Do NOT penalise the candidate for criteria that are not actually important in this JD.

3. DYNAMIC SCORING CRITERIA
   - First identify the 8-12 MOST IMPORTANT criteria directly from THIS JD.
   - These criteria MUST change depending on the role. Do not reuse a fixed checklist.
   - Use role-relevant categories only, for example when applicable:
     enterprise architecture, solution/integration design, stakeholder engagement,
     AI innovation, hospitality systems, cyber/compliance, vendor leadership,
     business partnership, programme delivery, people leadership, communication.
   - If the optional JD JSON blueprint exists, use its priorities/criteria/weights as the starting point.
   - If no JSON exists, infer the criteria directly from the raw JD.
   - Use semantic matching, not just exact keyword matching. Example: if the JD asks for cybersecurity,
     and the profile shows IT security, risk, audit, DR, patching, compliance, or resilience work,
     connect those dots clearly as relevant evidence.
   - Reward strong transferable evidence when it is meaningfully adjacent to the JD, but do NOT overstate or invent direct experience.
   - For each criterion provide:
     * Criterion Name
     * Weight (%) based on JD importance — all weights must total 100
     * Score: X/10
     * Evidence: specific proof from john_profile.json
     * Gap: what is missing or weaker versus the JD

4. KEY STRENGTHS
   Bullet list — what clearly matches or exceeds the JD requirements.

5. GAPS & RISKS
   Honest, specific gaps between the candidate profile and JD expectations.

6. TAILORING RECOMMENDATIONS
   a) Resume adjustments — specific wording and framing suggestions
   b) Interview preparation — topics to research, answers to rehearse
   c) Certifications to pursue — quick wins before applying

7. VERDICT
   Should the candidate apply? Competitive positioning. Recommended approach.
"""

    scorecard_text = call_llm(SCORECARD_SYS, SCORECARD_USER, max_tokens=6000, label="Scorecard")
    print("  ✓  Scorecard complete")

# ══════════════════════════════════════════════════════════════════════════════
# TASK 2 — TAILORED RESUME
# ══════════════════════════════════════════════════════════════════════════════
resume_text = ""
if run_resume:
    print("📝  [2/3] Generating Tailored Resume...")

    RESUME_SYS = (
        "You are a professional executive resume writer specialising in senior IT leadership roles. "
        "Rules you MUST follow:\n"
        "1. ALL facts, figures, dates, and claims come ONLY from the candidate profile — never invent or embellish.\n"
        "2. Mirror the reference template layout EXACTLY: same section headers, same ___ separators, same bullet style.\n"
        "3. Tailor bullet wording to echo the JD language without distorting facts.\n"
        "4. Use semantic alignment for related skills, but NEVER claim direct hands-on experience with named systems, standards, tools, or industries unless they appear explicitly in the profile data.\n"
        "5. If the JD asks for hospitality/PMS/POS/PCI or other domain-specific items not explicitly in the profile, position John's background as transferable or adjacent experience only.\n"
        "6. The `AI & AUTOMATION HIGHLIGHTS` section is mandatory in every resume output. Preserve it and populate it with vivid, concrete examples from the profile/template only.\n"
        "7. Make the resume balanced: show both technical depth and leadership/soft skills supported by the profile.\n"
        "8. Output ONLY the resume text — no preamble, no explanation, no markdown code fences."
    )

    RESUME_USER = f"""
=== REFERENCE RESUME TEMPLATE (replicate this layout exactly) ===
{template_txt}

=== JOB DESCRIPTION (tailor content for this role) ===
{jd_text}

=== OPTIONAL JD BLUEPRINT (use this if present, otherwise infer from JD text) ===
{jd_blueprint_context}

=== CANDIDATE PROFILE (single source of truth — all facts from here only) ===
{profile_context}

=== INSTRUCTIONS ===
Write a complete tailored resume for John Hau targeting the role above.

Layout rules:
- Keep the identical header format, ________ separators, and bullet style from the reference template
- Respect and preserve any existing soft-skills section already present in the reference template, while tailoring it to this JD using only profile facts
- The `AI & AUTOMATION HIGHLIGHTS` section from the template is mandatory and must never be omitted, because it demonstrates current innovation capability with vivid examples
- First infer the 5-8 most important themes from THIS JD (technical, leadership, business, industry, and interpersonal)
- Replace the Professional Summary to foreground those JD-specific themes using only evidence from the profile
- Add a dynamic bridging section immediately before PROFESSIONAL EXPERIENCE. The section title and content must fit the JD nature
  (for example: "HOSPITALITY RELEVANCE", "ROLE RELEVANCE", or "ENTERPRISE ARCHITECTURE RELEVANCE") and should never be hardcoded
- BULLET COUNT RULES (non-negotiable):
    * 1st company (most recent):  exactly 12 bullets
    * 2nd company:                exactly 12 bullets
    * 3rd company:                exactly 10 bullets
    * 4th company:                exactly 10 bullets
    * 5th company:                exactly 8 bullets
    * Any earlier roles:          3-4 bullets combined
- Choose the bullets most relevant to the JD — reword to echo JD language without distorting facts
- Use semantic matching to connect related evidence from `john_profile.json` to the JD. Example: cybersecurity ↔ IT security / audit / risk / DR / compliance;
  stakeholder management ↔ executive presentations / business partnering; operational excellence ↔ SLA, uptime, automation, service quality improvements
- Do NOT convert adjacent evidence into unsupported exact claims. Example: do not say John has direct `PCI-DSS`, `Opera PMS`, or hotel `POS/CRM` experience unless the profile explicitly says so
- Soft skills must be visible across the resume, especially in the Professional Summary, the dynamic relevance bridge,
  and selected experience bullets: stakeholder management, executive communication, cross-functional collaboration,
  mentoring, vendor management, team leadership, problem-solving, and change leadership — only where supported by profile facts
- The goal is to maximise JD match through truthful semantic alignment, not keyword stuffing and not hallucination
- If the template includes a competencies/skills section, ensure it reflects a mix of technical and soft skills
- Keep EDUCATION & CERTIFICATIONS, LANGUAGES, and AVAILABILITY sections unchanged from profile data
- Total length: equivalent to the reference template (~2 pages of content)
"""

    resume_text = call_llm(RESUME_SYS, RESUME_USER, max_tokens=4500, label="Resume")
    print("  ✓  Resume complete")

# ══════════════════════════════════════════════════════════════════════════════
# TASK 3 — COVER LETTER
# ══════════════════════════════════════════════════════════════════════════════
coverletter_text = ""
if run_coverletter:
    print("✉️   [3/3] Generating Cover Letter...")

    # Extract contact details from profile metadata for the letter header
    meta     = profile.get("metadata", {})
    name     = meta.get("name", "John Hau")
    location = meta.get("location", "Hong Kong SAR")
    phone    = meta.get("phone", "+852 5722 2007")
    email    = meta.get("email", "haujon001@gmail.com")
    linkedin = meta.get("linkedin", "linkedin.com/in/johnhau")
    today    = datetime.now().strftime("%d %B %Y")

    COVERLETTER_SYS = (
        "You are a professional executive cover letter writer with deep expertise in senior technology leadership roles. "
        "Adapt the tone and emphasis to the specific JD and company context provided. "
        "Rules:\n"
        "1. ALL claims and achievements must come ONLY from the candidate profile data provided — no fabrication.\n"
        "2. Write in first person, confident executive tone, not generic or templated.\n"
        "3. The letter must be 4-5 substantive paragraphs:\n"
        "   Para 1 — Why this role at this company based on the JD/company context provided.\n"
        "   Para 2 — Most relevant leadership & operational experience (quantified from profile).\n"
        "   Para 3 — How the candidate's enterprise background translates to this specific role context.\n"
        "   Para 4 — Innovation, transformation, and people leadership as differentiators.\n"
        "   Para 5 — Closing: call to action, enthusiasm, availability.\n"
        "4. Use semantic matching to connect adjacent experience from the profile to the JD without overstating direct domain experience.\n"
        "5. Never claim direct experience with named domain-specific tools, standards, or industries unless they are explicitly present in the profile data.\n"
        "6. Do NOT use generic phrases like 'I am writing to apply for'. Open with impact.\n"
        "7. Output ONLY the cover letter text — no preamble, no commentary."
    )

    # Derive employer display name from employer slug (e.g. MandarinOriental -> Mandarin Oriental)
    import re as _re
    employer_display = _re.sub(r'(?<=[a-z])(?=[A-Z])', ' ', employer)  # CamelCase -> spaced
    role_display     = role_slug.replace('_', ' ')                      # slug -> readable

    COVERLETTER_USER = f"""
=== JOB DESCRIPTION ===
{jd_text}

=== OPTIONAL JD BLUEPRINT (use this if present, otherwise infer from JD text) ===
{jd_blueprint_context}

=== CANDIDATE PROFILE (single source of truth — all facts from here only) ===
{profile_context}

=== LETTER HEADER DETAILS ===
Name     : {name}
Location : {location}
Phone    : {phone}
Email    : {email}
LinkedIn : {linkedin}
Date     : {today}
Hiring Company : {employer_display}
Role     : {role_display}

=== INSTRUCTIONS ===
Write a formal executive cover letter for John Hau applying for the {role_display} role
at {employer_display}.

Structure:
- Letter header block (candidate contact details, date, company)
- Salutation: Dear Hiring Manager,
- 4-5 focused paragraphs as described in the system instructions
- Professional sign-off

Key themes to hit (using only profile facts):
- The highest-priority business and technology themes from this specific JD
- Strong transferable leadership, operational, cybersecurity, transformation, vendor, and stakeholder evidence from the profile where relevant
- Budget governance and measurable savings where they support the business case
- Team building, coaching, and cross-functional partnership where relevant to the JD
- AI/automation innovation as a differentiator when it genuinely supports the role
- Immediate availability, Hong Kong-based, and Cantonese-speaking only if helpful and supported by the profile
- Use semantic alignment: connect adjacent evidence honestly rather than waiting for exact keyword matches
"""

    coverletter_text = call_llm(COVERLETTER_SYS, COVERLETTER_USER, max_tokens=2500, label="Cover Letter")
    print("  ✓  Cover Letter complete")

# ══════════════════════════════════════════════════════════════════════════════
# WRITE OUTPUT FILES
# ══════════════════════════════════════════════════════════════════════════════
# Create all output dirs (txt + docx/pdf companion folders)
for d in [OUT_DIR_SCORECARD, OUT_DIR_RESUME, OUT_DIR_COVER]:
    d.mkdir(parents=True, exist_ok=True)
# Ensure sibling docx/pdf dirs also exist for future use
for subtype in ["ScoreCard", "resume", "CoverLetter"]:
    for fmt in ["docx", "pdf"]:
        (OUT_BASE / subtype / fmt).mkdir(parents=True, exist_ok=True)

generated_at = datetime.now().strftime("%Y-%m-%d %H:%M")
print("\n💾  Writing output files...")

if run_scorecard and scorecard_text:
    header = (
        f"JD MATCH SCORECARD — {employer.upper()} | {role_slug.upper()}\n"
        f"Generated : {generated_at}  |  Model: {MODEL}\n"
        f"Source JD : {jd_path.name}\n"
        f"Profile   : src/data/john_profile.json\n"
        f"Candidate : John Hau\n"
        + "=" * 72 + "\n\n"
    )
    OUT_SCORECARD.write_text(header + scorecard_text, encoding="utf-8")
    print(f"  ✅  Scorecard   → {OUT_SCORECARD.relative_to(ROOT)}")
    maybe_convert_to_docx(OUT_SCORECARD, "Scorecard")

if run_resume and resume_text:
    resume_header = (
        f"TAILORED RESUME — {employer.upper()} | {role_slug.upper()}\n"
        f"Generated : {generated_at}  |  Model: {MODEL}\n"
        f"Profile   : src/data/john_profile.json\n"
        + "=" * 72 + "\n\n"
    )
    OUT_RESUME.write_text(resume_header + resume_text, encoding="utf-8")
    print(f"  ✅  Resume      → {OUT_RESUME.relative_to(ROOT)}")
    maybe_convert_to_docx(OUT_RESUME, "Resume")

if run_coverletter and coverletter_text:
    cl_header = (
        f"COVER LETTER — {employer.upper()} | {role_slug.upper()}\n"
        f"Generated : {generated_at}  |  Model: {MODEL}\n"
        f"Profile   : src/data/john_profile.json\n"
        + "=" * 72 + "\n\n"
    )
    OUT_COVERLETTER.write_text(cl_header + coverletter_text, encoding="utf-8")
    print(f"  ✅  Cover Letter → {OUT_COVERLETTER.relative_to(ROOT)}")
    maybe_convert_to_docx(OUT_COVERLETTER, "Cover Letter")

print(f"\n  Output folder: data_processed/{employer}/")

print(f"\n{'='*60}")
print("  ALL DONE")
print(f"{'='*60}")
print(f"\n  How to run this script:")
print(f"  python scripts/jd_scorecard_resume.py")
print(f"  python scripts/jd_scorecard_resume.py \"data_raw/jd/txt/AnotherJD.txt\"")
print(f"  python scripts/jd_scorecard_resume.py --scorecard-only")
print(f"  python scripts/jd_scorecard_resume.py --resume-only")
print(f"  python scripts/jd_scorecard_resume.py --coverletter-only")
print(f"  python scripts/jd_scorecard_resume.py --batch")
print(f"  python scripts/jd_scorecard_resume.py --batch --force")
print(f"  python scripts/jd_scorecard_resume.py --refresh-blueprint")
print(f"  python scripts/jd_scorecard_resume.py --llm=sonnet    (default, Claude Sonnet 4.6)")
print(f"  python scripts/jd_scorecard_resume.py --llm=deepseek  (DeepSeek R1)")
print(f"  python scripts/jd_scorecard_resume.py --llm=gemini    (Gemini Flash Lite - cheaper)")
